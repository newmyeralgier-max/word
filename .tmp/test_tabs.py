import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from word_utils import latex_to_omml, add_seq_field

doc = Document()
p = doc.add_paragraph()
# left align the paragraph
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.first_line_indent = Cm(0)

# Add custom tab stops
tab_stops = p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Cm(8.25), WD_TAB_ALIGNMENT.CENTER)
tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)

# Add the leading tab to push to center
p.add_run("\t")

# Add the OMML equation
omml = latex_to_omml(r"H=\frac{1}{2}\cdot J \cdot \omega_{12}P_{н}")
p._element.append(omml)

# Add the trailing tab to push to right
run = p.add_run("\t(")
add_seq_field(run, "Формула")
p.add_run(")")

doc.save("test_equation_tabs.docx")
