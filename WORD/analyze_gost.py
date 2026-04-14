"""Анализ БР_GOST.docx — проверка проблем форматирования."""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'execution'))

from docx import Document
from docx.shared import Cm, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOC_PATH = os.path.join(os.path.dirname(__file__), '..', 'data', 'БР_GOST.docx')
doc = Document(DOC_PATH)

ALIGN_MAP = {0: 'LEFT', 1: 'CENTER', 2: 'RIGHT', 3: 'JUSTIFY', None: 'None'}

def align_name(a):
    return ALIGN_MAP.get(a, str(a))

def has_math(p):
    return len(p._element.xpath('.//m:oMath')) > 0

def has_image(p):
    return len(p._element.xpath('.//w:drawing')) > 0

# 1. Sections
print('=== SECTIONS ===')
for i, s in enumerate(doc.sections):
    pw = s.page_width
    ph = s.page_height
    print(f'  S{i}: w={pw} h={ph} L={s.left_margin} R={s.right_margin} T={s.top_margin} B={s.bottom_margin}')

# 2. Tables
print('\n=== TABLES ===')
for i, t in enumerate(doc.tables):
    tbl_el = t._element
    prev = tbl_el.getprevious()
    caption = ''
    while prev is not None:
        tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
        if tag == 'p':
            texts = prev.findall('.//' + qn('w:t'))
            caption = ' '.join((t.text or '') for t in texts).strip()
            if caption:
                break
        prev = prev.getprevious()
    
    tblPr = tbl_el.find(qn('w:tblPr'))
    jc = ind = '-'
    if tblPr is not None:
        jc_el = tblPr.find(qn('w:jc'))
        if jc_el is not None:
            jc = jc_el.get(qn('w:val'), '-')
        ind_el = tblPr.find(qn('w:tblInd'))
        if ind_el is not None:
            ind = ind_el.get(qn('w:w'), '-')
    # Count rows
    nrows = len(t.rows)
    print(f'  T{i} ({nrows}rows): cap="{caption[:70]}" jc={jc} ind={ind}')

# 3. Formulas
print('\n=== FORMULAS (first 25) ===')
fcount = 0
for p in doc.paragraphs:
    if has_math(p):
        fcount += 1
        if fcount <= 25:
            al = align_name(p.alignment)
            pf = p.paragraph_format
            runs_text = []
            for r in p.runs:
                if r.text and r.text.strip():
                    runs_text.append(r.text.strip()[:50])
            rt = ' | '.join(runs_text) if runs_text else '[OMML]'
            # Tab stops
            pPr = p._element.find(qn('w:pPr'))
            has_tabs = False
            if pPr is not None:
                tabs_el = pPr.find(qn('w:tabs'))
                if tabs_el is not None:
                    has_tabs = True
            print(f'  F{fcount}: al={al} fi={pf.first_line_indent} li={pf.left_indent} '
                  f'tabs={has_tabs} runs="{rt}"')
print(f'Total OMML: {fcount}')

# 4. Body paragraphs with unusual alignment
print('\n=== BODY WITH NON-JUSTIFY ALIGNMENT ===')
weird = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    sn = p.style.name if p.style else 'None'
    al = align_name(p.alignment)
    if any(k in sn for k in ['Heading', 'Заголовок', 'TOC', 'Содержание']):
        continue
    if text.startswith('Рис') or text.startswith('Таблица') or text.startswith('Приложение'):
        continue
    if al in ('LEFT', 'CENTER', 'RIGHT', 'None'):
        weird += 1
        if weird <= 30:
            pf = p.paragraph_format
            print(f'  P{i}: al={al} sn="{sn}" li={pf.left_indent} fi={pf.first_line_indent} '
                  f'text="{text[:80]}"')
print(f'Total non-JUSTIFY body: {weird}')

# 5. Heading spacing
print('\n=== HEADING SPACING ===')
hc = 0
for p in doc.paragraphs:
    sn = p.style.name if p.style else ''
    if 'Heading' in sn or 'Заголовок' in sn:
        hc += 1
        if hc <= 25:
            pf = p.paragraph_format
            al = align_name(p.alignment)
            print(f'  H{hc}: sn="{sn}" al={al} sb={pf.space_before} sa={pf.space_after} '
                  f'ls={pf.line_spacing} pbb={pf.page_break_before} '
                  f'text="{p.text.strip()[:70]}"')
print(f'Total headings: {hc}')

# 6. Title zone first 30
print('\n=== FIRST 30 PARAGRAPHS (title zone) ===')
for i, p in enumerate(doc.paragraphs[:30]):
    text = p.text.strip()
    sn = p.style.name if p.style else 'None'
    al = align_name(p.alignment)
    pf = p.paragraph_format
    flags = ''
    if has_math(p): flags += ' [MATH]'
    if has_image(p): flags += ' [IMG]'
    if not text and not has_math(p) and not has_image(p): flags += ' [EMPTY]'
    txt = text[:80] if text else ''
    print(f'  P{i}: sn="{sn}" al={al} li={pf.left_indent} fi={pf.first_line_indent} '
          f'sb={pf.space_before} sa={pf.space_after}{flags} text="{txt}"')

# 7. Check if "Таблица 2.5" exists
print('\n=== TABLE CAPTION SEARCH ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Таблица 2.5' in text or 'Таблица 4.4' in text or 'Таблица 6.5' in text:
        sn = p.style.name if p.style else 'None'
        al = align_name(p.alignment)
        print(f'  P{i}: sn="{sn}" al={al} text="{text[:100]}"')
