"""Сравнение оригинала и GOST — ключевые отличия."""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'execution'))

from docx import Document
from docx.shared import Cm, Pt, Mm
from docx.oxml.ns import qn

ORIG = os.path.join(os.path.dirname(__file__), '..', 'data', 'БР.docx')
doc = Document(ORIG)

ALIGN_MAP = {0: 'LEFT', 1: 'CENTER', 2: 'RIGHT', 3: 'JUSTIFY', None: 'None'}
def aln(a): return ALIGN_MAP.get(a, str(a))

# 1. Sections
print('=== ORIGINAL SECTIONS ===')
for i, s in enumerate(doc.sections):
    print(f'  S{i}: w={s.page_width} h={s.page_height} L={s.left_margin} R={s.right_margin} T={s.top_margin} B={s.bottom_margin}')

# 2. Check tables
print('\n=== ORIGINAL TABLES ===')
for i, t in enumerate(doc.tables):
    tbl_el = t._element
    prev = tbl_el.getprevious()
    caption = ''
    while prev is not None:
        tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
        if tag == 'p':
            texts = prev.findall('.//' + qn('w:t'))
            caption = ' '.join((t.text or '') for t in texts).strip()
            if caption:
                break
        prev = prev.getprevious()
    tblPr = tbl_el.find(qn('w:tblPr'))
    jc = ind = '-'
    if tblPr is not None:
        jc_el = tblPr.find(qn('w:jc'))
        if jc_el is not None:
            jc = jc_el.get(qn('w:val'), '-')
        ind_el = tblPr.find(qn('w:tblInd'))
        if ind_el is not None:
            ind = ind_el.get(qn('w:w'), '-')
    nrows = len(t.rows)
    print(f'  T{i} ({nrows}r): cap="{caption[:70]}" jc={jc} ind={ind}')

# 3. Check headings before ВВЕДЕНИЕ — what happens with ЗАДАНИЕ?
print('\n=== ORIGINAL: P30-70 (title zone) ===')
for i, p in enumerate(doc.paragraphs[30:80]):
    text = p.text.strip()
    sn = p.style.name if p.style else 'None'
    al = aln(p.alignment)
    pf = p.paragraph_format
    flags = ''
    has_math = len(p._element.xpath('.//m:oMath')) > 0
    has_img = len(p._element.xpath('.//w:drawing')) > 0
    if has_math: flags += ' [M]'
    if has_img: flags += ' [I]'
    if not text and not has_math and not has_img: flags += ' [E]'
    idx = i + 30
    print(f'  P{idx}: sn="{sn}" al={al} li={pf.left_indent} fi={pf.first_line_indent}{flags} text="{text[:80] if text else ""}"')

# 4. Find table captions 2.5, 4.4, 6.5
print('\n=== ORIGINAL: TABLE CAPTION SEARCH ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Таблица 2.5' in text or 'Таблица 4.4' in text or 'Таблица 6.5' in text:
        sn = p.style.name if p.style else 'None'
        al = aln(p.alignment)
        # Check next paragraph — is it a table or text?
        nxt = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        nxt_text = nxt.text.strip()[:80] if nxt else 'END'
        nxt_sn = nxt.style.name if nxt and nxt.style else 'None'
        print(f'  P{i}: sn="{sn}" al={al} text="{text[:100]}"')
        print(f'    NEXT: sn="{nxt_sn}" text="{nxt_text}"')

# 5. Compare page margins — original vs what formatter sets
print('\n=== ORIGINAL FIRST SECTION MARGINS (in mm) ===')
s0 = doc.sections[0]
print(f'  page: {Mm(s0.page_width.mm):.1f} x {Mm(s0.page_height.mm):.1f} mm')
print(f'  margins: L={Mm(s0.left_margin.mm):.1f} R={Mm(s0.right_margin.mm):.1f} T={Mm(s0.top_margin.mm):.1f} B={Mm(s0.bottom_margin.mm):.1f} mm')
