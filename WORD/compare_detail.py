"""Сравнение оригинала и GOST — детальная диагностика."""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'execution'))

from docx import Document
from docx.shared import Cm, Pt, Mm, Emu, Twips
from docx.oxml.ns import qn

ORIG = '/mnt/d/1. Project/Word/data/БР.docx'
GOST = '/mnt/d/1. Project/Word/data/БР_GOST.docx'

ALIGN_MAP = {0: 'LEFT', 1: 'CENTER', 2: 'RIGHT', 3: 'JUSTIFY', None: 'None'}
def aln(a): return ALIGN_MAP.get(a, str(a))

# Margins comparison
for label, path in [('ORIG', ORIG), ('GOST', GOST)]:
    doc = Document(path)
    s0 = doc.sections[0]
    print(f'=== {label} SECTION 0 ===')
    print(f'  page_w={s0.page_width} page_h={s0.page_height}')
    print(f'  L={s0.left_margin.cm:.2f}cm R={s0.right_margin.cm:.2f}cm T={s0.top_margin.cm:.2f}cm B={s0.bottom_margin.cm:.2f}cm')
    # All sections
    for i, s in enumerate(doc.sections):
        print(f'  S{i}: L={s.left_margin.cm:.2f} R={s.right_margin.cm:.2f} T={s.top_margin.cm:.2f} B={s.bottom_margin.cm:.2f} w={s.page_width.cm:.2f}x{s.page_height.cm:.2f}')

# Heading spacing comparison
for label, path in [('ORIG', ORIG), ('GOST', GOST)]:
    doc = Document(path)
    print(f'\n=== {label} HEADING FORMAT (body headings after ВВЕДЕНИЕ) ===')
    started = False
    count = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if 'ВВЕДЕНИЕ' in t and 'Heading' in (p.style.name or ''):
            started = True
        if not started:
            continue
        sn = p.style.name if p.style else ''
        if 'Heading' in sn and t:
            pf = p.paragraph_format
            al = aln(p.alignment)
            sb = f'{pf.space_before.pt:.1f}pt' if pf.space_before else 'inh'
            sa = f'{pf.space_after.pt:.1f}pt' if pf.space_after else 'inh'
            ls = f'{pf.line_spacing}' if pf.line_spacing else 'inh'
            li = f'{pf.left_indent.cm:.2f}cm' if pf.left_indent else 'inh'
            fi = f'{pf.first_line_indent.cm:.2f}cm' if pf.first_line_indent else 'inh'
            print(f'  P{i}: sn="{sn}" al={al} sb={sb} sa={sa} ls={ls} li={li} fi={fi} "{t[:70]}"')
            count += 1
            if count >= 25:
                break

# Table properties
for label, path in [('ORIG', ORIG), ('GOST', GOST)]:
    doc = Document(path)
    print(f'\n=== {label} TABLE PROPERTIES (body tables >2 rows) ===')
    for i, t in enumerate(doc.tables):
        nrows = len(t.rows)
        if nrows <= 2:
            continue
        tbl_el = t._element
        tblPr = tbl_el.find(qn('w:tblPr'))
        jc = ind = w = '-'
        if tblPr is not None:
            jc_el = tblPr.find(qn('w:jc'))
            if jc_el is not None: jc = jc_el.get(qn('w:val'), '-')
            ind_el = tblPr.find(qn('w:tblInd'))
            if ind_el is not None: ind = ind_el.get(qn('w:w'), '-')
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None: w = tblW.get(qn('w:w'), '-') + tblW.get(qn('w:type'), '-')
        # Caption from previous paragraph
        prev = tbl_el.getprevious()
        caption = ''
        while prev is not None:
            tag = prev.tag.split('}')[-1] if '}' in prev.tag else prev.tag
            if tag == 'p':
                texts = prev.findall('.//' + qn('w:t'))
                caption = ' '.join((tt.text or '') for tt in texts).strip()[:70]
                break
            prev = prev.getprevious()
        print(f'  T{i} ({nrows}r): jc={jc} ind={ind} w={w} cap="{caption}"')

# Formula paragraphs
for label, path in [('ORIG', ORIG), ('GOST', GOST)]:
    doc = Document(path)
    print(f'\n=== {label} FORMULA PARAGRAPHS (first 15 OMML) ===')
    count = 0
    started = False
    for i, p in enumerate(doc.paragraphs):
        if 'ВВЕДЕНИЕ' in p.text:
            started = True
        if not started:
            continue
        has_math = len(p._element.xpath('.//m:oMath')) > 0
        if has_math:
            al = aln(p.alignment)
            pf = p.paragraph_format
            li = f'{pf.left_indent.cm:.2f}cm' if pf.left_indent else 'inh'
            fi = f'{pf.first_line_indent.cm:.2f}cm' if pf.first_line_indent else 'inh'
            # Tab stops
            pPr = p._element.find(qn('w:pPr'))
            tab_info = ''
            if pPr is not None:
                tabs_el = pPr.find(qn('w:tabs'))
                if tabs_el is not None:
                    for tab in tabs_el.findall(qn('w:tab')):
                        tab_info += f' tab@{tab.get(qn("w:pos"), "?")}'
            # Text
            texts = p._element.findall('.//' + qn('w:t'))
            txt = ' '.join((t.text or '') for t in texts).strip()[:70]
            print(f'  P{i}: al={al} li={li} fi={fi} tabs=[{tab_info}] "{txt}"')
            count += 1
            if count >= 15:
                break

# Table 2.5 and 4.4 area
for label, path in [('ORIG', ORIG), ('GOST', GOST)]:
    doc = Document(path)
    print(f'\n=== {label} TABLE 2.5 / 4.4 AREA ===')
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if 'Таблица 2.5' in t or 'Таблица 4.4' in t:
            sn = p.style.name
            al = aln(p.alignment)
            print(f'  P{i}: sn="{sn}" al={al} text="{t[:100]}"')
            for j in range(1, 4):
                if i+j < len(doc.paragraphs):
                    np_ = doc.paragraphs[i+j]
                    ntxt = np_.text.strip()[:60]
                    nsn = np_.style.name
                    print(f'    +{j}: sn="{nsn}" text="{ntxt}"')
            # Check next XML sibling
            nxt = p._element.getnext()
            if nxt is not None:
                tag = nxt.tag.split('}')[-1] if '}' in nxt.tag else nxt.tag
                print(f'    NEXT_ELEM: tag={tag}')

# Text alignment — check pages 31-32 area for left-shifted text
for label, path in [('ORIG', ORIG), ('GOST', GOST)]:
    doc = Document(path)
    print(f'\n=== {label} BODY TEXT WITH NON-JUSTIFY ALIGNMENT ===')
    count = 0
    started = False
    for i, p in enumerate(doc.paragraphs):
        if 'ВВЕДЕНИЕ' in p.text:
            started = True
        if not started:
            continue
        sn = p.style.name if p.style else ''
        if 'Heading' in sn or 'Title' in sn:
            continue
        al = p.alignment
        if al is not None and al != 3:  # Not JUSTIFY and not None
            t = p.text.strip()[:60]
            if t:
                print(f'  P{i}: sn="{sn}" al={aln(al)} "{t}"')
                count += 1
                if count >= 20:
                    break
