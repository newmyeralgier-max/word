"""
build_docx.py -- Сборка ГОСТ-документа из Markdown-файла.

* Парсинг Markdown через mistune (AST-рендерер, плагины table)
* Pygments для графического выделения кода
* SEQ автонумерация таблиц, рисунков, формул
* Мультифайловая сборка
"""

import sys, os, re, glob, argparse

# ?? Путь к модулям рядом ??????????????????????????????????????????
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import word_config as cfg
import word_utils  as wu

import mistune
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from pygments import lex
from pygments.lexers import get_lexer_by_name, guess_lexer
from pygments.token import Token

#  Вспомогательные: Code Highlighting
# ??????????????????????????????????????????????????????????????????

def _apply_p_format(p, align, left_ind=Cm(0), right_ind=Cm(0), first_line_ind=Cm(0), 
                   space_before=Pt(0), space_after=Pt(0), keep_next=False, 
                   line_spacing=cfg.LINE_SPACING):
    """Универсальная настройка формата абзаца."""
    p.alignment = align
    fmt = p.paragraph_format
    fmt.left_indent = left_ind
    fmt.right_indent = right_ind
    fmt.first_line_indent = first_line_ind
    fmt.space_before = space_before
    fmt.space_after = space_after
    fmt.keep_with_next = keep_next
    fmt.line_spacing = line_spacing


def _add_source_code(doc, code, language=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    try:
        lexer = get_lexer_by_name(language) if language else guess_lexer(code)
    except Exception:
        lexer = get_lexer_by_name("text")

    for token, text in lex(code, lexer):
        run = p.add_run(text)
        run.font.name = cfg.FONT_NAME_CODE
        run.font.size = cfg.FONT_SIZE_CODE
        if token in Token.Keyword:
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.bold = True
        elif token in Token.String:
            run.font.color.rgb = RGBColor(163, 21, 21)
        elif token in Token.Comment:
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.italic = True
        elif token in Token.Number:
            run.font.color.rgb = RGBColor(9, 134, 88)
        elif token in Token.Name.Builtin:
            run.font.color.rgb = RGBColor(121, 94, 38)
        else:
            run.font.color.rgb = cfg.COLOR_BLACK

# ??????????????????????????????????????????????????????????????????
#  Вспомогательные: формулы
# ??????????????????????????????????????????????????????????????????

def _add_equation(doc, latex_str, h1_idx=0):
    """Display-формула (по центру). Авто-нумерация по правому краю."""
    p = doc.add_paragraph()
    _apply_p_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6), space_after=Pt(6))

    # Настраиваем табуляцию: центр на 8.25 см, правый край на 16.5 см
    from docx.enum.text import WD_TAB_ALIGNMENT
    tab_stops = p.paragraph_format.tab_stops
    if len(tab_stops) > 0:
        tab_stops.clear_all()
    tab_stops.add_tab_stop(Cm(8.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)

    p.add_run("\t")

    omml = wu.latex_to_omml(latex_str)
    if omml is not None:
        p._element.append(omml)
    else:
        run = p.add_run(latex_str)
        run.font.name = cfg.FONT_NAME_MATH
        run.font.size = cfg.FONT_SIZE_MAIN
        run.italic = True

    # Авто-нумерация через SEQ на правом краю
    if h1_idx > 0:
        run = p.add_run(f"\t({h1_idx}.")
    else:
        run = p.add_run("\t(")

    run.font.name = cfg.FONT_NAME
    run.font.size = cfg.FONT_SIZE_MAIN
    
    seq_name = f"Формула_{h1_idx}" if h1_idx > 0 else "Формула"
    wu.add_seq_field(run, seq_name)
    
    run_close = p.add_run(")")
    run_close.font.name = cfg.FONT_NAME
    run_close.font.size = cfg.FONT_SIZE_MAIN
    return p


def _add_inline_math(paragraph, latex_str):
    """Инлайн-формула внутри абзаца."""
    omml = wu.latex_to_omml(latex_str)
    if omml is not None:
        paragraph._element.append(omml)
    else:
        run = paragraph.add_run(latex_str)
        run.font.name = cfg.FONT_NAME_MATH
        run.font.size = cfg.FONT_SIZE_MAIN
        run.italic = True


# ??????????????????????????????????????????????????????????????????
#  Вспомогательные: таблицы
# ??????????????????????????????????????????????????????????????????

def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tcBorders>'
    ))


def _add_table_caption(doc, text, h1_idx=0):
    """Подпись таблицы (две строки: номер справа, название по центру)."""
    clean_caption = re.sub(r'^(?:\*?Таблица\s*[\d\.]*|\*?Таблица)\s*(?:--|-|—)?\s*', '', text.strip('* '))
    
    # 1. Номер (справа)
    p1 = doc.add_paragraph()
    _apply_p_format(p1, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(12), keep_next=True)
    
    if h1_idx > 0:
        run1 = p1.add_run(f"Таблица {h1_idx}.")
    else:
        run1 = p1.add_run("Таблица ")
    run1.font.name = cfg.FONT_NAME
    run1.font.size = cfg.FONT_SIZE_MAIN
    seq_name = f"Таблица_{h1_idx}" if h1_idx > 0 else "Таблица"
    wu.add_seq_field(run1, seq_name)
    
    # 2. Название (по центру)
    if clean_caption:
        p2 = doc.add_paragraph()
        _apply_p_format(p2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6), keep_next=True)
        
        run2 = p2.add_run(clean_caption)
        run2.font.name = cfg.FONT_NAME
        run2.font.size = cfg.FONT_SIZE_MAIN
        return p2
    return p1


def _add_table(doc, header_cells, data_rows, caption="", h1_idx=0):
    if caption:
        _add_table_caption(doc, caption, h1_idx)

    ncols = len(header_cells)
    table = doc.add_table(rows=len(data_rows) + 1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tl = OxmlElement('w:tblLayout')
    tl.set(qn('w:type'), 'autofit')
    table._tbl.tblPr.append(tl)

    for j, txt in enumerate(header_cells):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_rich_text_to_paragraph(p, txt, font_size=cfg.FONT_SIZE_TABLE)
        _set_cell_border(cell)

    for i, row in enumerate(data_rows):
        for j in range(ncols):
            cell = table.cell(i + 1, j)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = 0
            _add_rich_text_to_paragraph(p, row[j], font_size=cfg.FONT_SIZE_TABLE)
            _set_cell_border(cell)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after  = Pt(0)
    spacer.paragraph_format.first_line_indent = Cm(0)


# ??????????????????????????????????????????????????????????????????
#  Вспомогательные: изображения и подписи
# ??????????????????????????????????????????????????????????????????

def _add_figure_caption(doc, text, h1_idx=0):
    clean_caption = re.sub(r'^(?:\*?Рисунок\s*[\d\.]*|\*?Рисунок)\s*(?:--|-|—)?\s*', '', text.strip('* '))
    p = doc.add_paragraph()
    _apply_p_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
    
    if h1_idx > 0:
        run = p.add_run(f"Рис. {h1_idx}.")
    else:
        run = p.add_run("Рис. ")

    run.font.name = cfg.FONT_NAME
    run.font.size = cfg.FONT_SIZE_MAIN
    seq_name = f"Рисунок_{h1_idx}" if h1_idx > 0 else "Рисунок"
    wu.add_seq_field(run, seq_name)
    
    if clean_caption:
        # Убираем возможные начальные точки/тире из названия и добавляем точку в конце
        clean_text = re.sub(r'^[.\s—–-]+', '', clean_caption).strip()
        run2 = p.add_run(f". {clean_text}.")
        run2.font.name = cfg.FONT_NAME
        run2.font.size = cfg.FONT_SIZE_MAIN


def _add_image_centered(doc, image_path):
    p = doc.add_paragraph()
    _apply_p_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6), keep_next=True)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(15))


# ??????????????????????????????????????????????????????????????????
#  TOC
# ??????????????????????????????????????????????????????????????????

def _add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run('СОДЕРЖАНИЕ')
    run.bold      = False
    run.font.name = cfg.FONT_NAME
    run.font.size = cfg.FONT_SIZE_TOC_TITLE

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pt.paragraph_format.first_line_indent = Cm(0)
    rt = pt.add_run()
    for tag, attr, text in [
        ('w:fldChar', 'begin', None),
        ('w:instrText', None, 'TOC \\o "1-3" \\h \\z \\u'),
        ('w:fldChar', 'separate', None),
    ]:
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = text
        else:
            el.set(qn('w:fldCharType'), attr)
        rt._r.append(el)

    # Placeholder text for cases where COM update is skipped
    tr = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = ' [Оглавление будет обновлено автоматически] '
    tr.append(t)
    rt._r.append(tr)

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    rt._r.append(end)

    doc.add_page_break()


# ??????????????????????????????????????????????????????????????????
#  Heading
# ??????????????????????????????????????????????????????????????????

def _add_heading(doc, text, level=1):
    bold_map = {1: cfg.BOLD_H1, 2: cfg.BOLD_H2, 3: cfg.BOLD_H3}
    size_map = {1: cfg.FONT_SIZE_H1, 2: cfg.FONT_SIZE_H2, 3: cfg.FONT_SIZE_H3}

    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    ol  = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), str(level - 1))
    pPr.append(ol)

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before      = Pt(24) if level == 1 else Pt(18)
    p.paragraph_format.space_after       = Pt(12)
    p.paragraph_format.line_spacing      = cfg.LINE_SPACING
    p.paragraph_format.keep_with_next    = True

    run = p.add_run(text.rstrip('.'))
    run.bold       = bold_map.get(level, False)
    run.font.name  = cfg.FONT_NAME
    run.font.size  = size_map.get(level, cfg.FONT_SIZE_MAIN)
    run.font.color.rgb = cfg.COLOR_BLACK
    return p


# ——————————————————————————————————————————————————————————————————
#  Rich-text paragraph / Inline rendering
# ——————————————————————————————————————————————————————————————————

def _add_rich_paragraph(doc, text, indent=True, align='justify'):
    t_low = text.strip().lower()
    marker_chars = r'0-9\-\u2013\u2014\.\•'

    # Авто-детект заголовков Рисунков/Таблиц для центрирования
    if t_low.startswith('рисунок') or t_low.startswith('таблица') or t_low.startswith('рис.') or t_low.startswith('*рис.') or t_low.startswith('*рисунок') or t_low.startswith('*таблица'):
        align = 'center'
        indent = False
    
    # Авто-детект списков (если начинается с тире, точки или цифры - убираем отступ)
    if re.match(r'^[' + marker_chars + r']', text.strip()):
        indent = False

    # === СИСТЕМА РАЗДЕЛЕНИЯ СКЛЕЕННЫХ СПИСКОВ ===

    is_where = t_low.startswith('где ') or t_low.startswith('здесь ')
    
    # Рекурсивная функция для применения деления
    def _apply_split(p_text, pattern, current_indent):
        if re.search(pattern, p_text):
            sub_parts = re.split(pattern, p_text)
            for i, pt in enumerate(sub_parts):
                val = pt.strip()
                if not val: continue
                # Первая часть сохраняет отступ, остальные - нет
                _add_rich_paragraph(doc, val, indent=(current_indent if i==0 else False), align='left')
            return True
        return False

    # 1. По ";" (для "где" любые слова, для остальных - маркеры)
    if is_where:
        split_pat = r'(?<=;)\s*(?=[\w\\а-яА-ЯёЁ])'
    else:
        split_pat = r'(?<=;)\s*(?=[' + marker_chars + r'])'

    if _apply_split(text, split_pat, indent):
        return

    # 2. По ":" если за ним сразу тире/маркер
    colon_split_pat = r'(?<=[:])\s*(?=[' + marker_chars + r'])'
    if _apply_split(text, colon_split_pat, indent):
        return
    # === КОНЕЦ СИСТЕМЫ РАЗДЕЛЕНИЯ ===
    p = doc.add_paragraph()

    # Центрирование или отступ
    if align == 'center':
        _apply_p_format(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        l_align = WD_ALIGN_PARAGRAPH.JUSTIFY
        if align == 'left': l_align = WD_ALIGN_PARAGRAPH.LEFT
        
        f_indent = cfg.FIRST_LINE_INDENT if indent else Cm(0)
        _apply_p_format(p, align=l_align, first_line_ind=f_indent)

    _add_rich_text_to_paragraph(p, text)
    return p

_SUB_RE = re.compile(
    r'(?<![A-Za-zА-Яа-яЁё0-9])([A-Za-zА-Яа-яЁё\u03c9\u03c0]+_[A-Za-zА-Яа-яЁё0-9]+)(?![A-Za-zА-Яа-яЁё0-9])'
)

def _auto_wrap_subscripts(text):
    parts = re.split(r'(`[^`]+?`|\$[^$]+?\$)', text)
    for i in range(0, len(parts), 2):
        if parts[i]:
            parts[i] = _SUB_RE.sub(r'$\1$', parts[i])
    return ''.join(parts)


def _add_rich_text_to_paragraph(paragraph, text, font_size=None):
    """
    Разбор текста с поддержкой:
    - Авто-подстрочных индексов (B_л -> $B_л$)
    - Жирный (**), Курсив (*), Код (`), Формулы ($)
    """
    if not text:
        return

    if font_size is None:
        font_size = cfg.FONT_SIZE_MAIN

    text = _auto_wrap_subscripts(text)
    
    # Регулярка для деления на токены
    token_re = re.compile(r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`|\$[^$]+?\$)')
    
    for part in token_re.split(text):
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = cfg.FONT_NAME
            run.font.size = font_size
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = cfg.FONT_NAME
            run.font.size = font_size
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = cfg.FONT_NAME_CODE
            run.font.size = cfg.FONT_SIZE_CODE
        elif part.startswith('$') and part.endswith('$'):
            _add_inline_math(paragraph, part[1:-1])
        else:
            run = paragraph.add_run(part)
            run.font.name = cfg.FONT_NAME
            run.font.size = font_size


def _render_inline_children(paragraph, children):
    for child in children:
        tp = child['type']
        raw = child.get('raw', child.get('text', ''))
        
        if tp == 'text':
            _add_rich_text_to_paragraph(paragraph, raw)
        elif tp == 'strong':
            run = paragraph.add_run(raw or _children_text(child))
            run.bold      = True
            run.font.name = cfg.FONT_NAME
            run.font.size = cfg.FONT_SIZE_MAIN
        elif tp == 'emphasis':
            run = paragraph.add_run(raw or _children_text(child))
            run.italic    = True
            run.font.name = cfg.FONT_NAME
            run.font.size = cfg.FONT_SIZE_MAIN
        elif tp == 'codespan':
            run = paragraph.add_run(raw)
            run.font.name = cfg.FONT_NAME_CODE
            run.font.size = cfg.FONT_SIZE_CODE
        elif tp in ('softbreak', 'linebreak'):
            pass
        else:
            if raw:
                run = paragraph.add_run(raw)
                run.font.name = cfg.FONT_NAME
                run.font.size = cfg.FONT_SIZE_MAIN

def _children_text(node):
    if 'children' not in node:
        return node.get('raw', node.get('text', ''))
    return ''.join(_children_text(c) for c in node['children'])

def _flat_text(node):
    if isinstance(node, str): return node
    text = node.get('raw', '') or node.get('text', '')
    if 'children' in node and node['children']:
        text += ''.join(_flat_text(c) for c in node['children'])
    return text


def _read_inputs(input_path):
    """Слияние Markdown файлов из директории или чтение одного файла."""
    if os.path.isdir(input_path):
        files = sorted(glob.glob(os.path.join(input_path, "*.md")))
        if not files:
            print(f"[!] MD файлы не найдены в папке {input_path}")
            sys.exit(1)
        texts = []
        for f in files:
            with open(f, 'r', encoding='utf-8') as fid:
                texts.append(fid.read())
        print(f"+ Склеено файлов: {len(files)}")
        return "\n\n---\n\n".join(texts), files[0]
    else:
        with open(input_path, 'r', encoding='utf-8') as f:
            return f.read(), input_path

# ??????????????????????????????????????????????????????????????????
#  Главная сборка
# ??????????????????????????????????????????????????????????????????

def build_document(input_path, output=None, fast=False, append_doc_path=None):
    if not os.path.exists(input_path):
        print(f"[!] Путь не найден: {input_path}")
        return

    raw_md, sample_path = _read_inputs(input_path)

    base_dir = os.path.dirname(os.path.abspath(sample_path))
    images_dir = os.path.join(base_dir, '..', 'data', 'images')
    if not os.path.exists(images_dir):
        images_dir = base_dir

    md = mistune.create_markdown(renderer='ast', plugins=['table'])
    ast_tokens = md(raw_md)
    print(f"  + AST-нод: {len(ast_tokens)}")

    if append_doc_path and os.path.exists(append_doc_path):
        doc = Document(append_doc_path)
        print(f"  + Открыт существующий документ для дозаписи: {append_doc_path}")
        # Если добавляем текст, добавим разрыв страницы перед новыми данными
        if doc.paragraphs:
            doc.add_page_break()
    else:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = cfg.MARGIN_TOP
            section.bottom_margin = cfg.MARGIN_BOTTOM
            section.left_margin   = cfg.MARGIN_LEFT
            section.right_margin  = cfg.MARGIN_RIGHT

    wu.setup_gost_styles(doc)
    wu.add_page_numbering(doc, smart_skip=True)

    has_h1 = False
    current_h1_idx = 0
    token_iterator = enumerate(ast_tokens)

    for i, node in token_iterator:
        tp = node['type']

        if tp == 'heading':
            level = node['attrs']['level']
            text  = _flat_text(node).strip()
            if level == 1:
                if has_h1:
                    doc.add_page_break()
                has_h1 = True
                if text.upper() not in cfg.STRICT_H1:
                    current_h1_idx += 1
            _add_heading(doc, text, level=min(level, 3))
            continue

        if tp == 'paragraph':
            raw = _flat_text(node).strip()

            if raw.upper() in ('[TOC]', '[[TOC]]'):
                _add_toc(doc)
                continue

            # Подпись Рисунка (более гибкий поиск)
            if re.match(r'^\*?\s*Рисунок\s*[\d\.]+', raw, re.I):
                caption_text = raw.strip('* ')
                m = re.search(r'Рисунок\s+(\d+\.\d+)', caption_text)
                if m:
                    fig_num = m.group(1)
                    found_images = glob.glob(os.path.join(images_dir, f'fig{fig_num}.*'))
                    if found_images:
                        _add_image_centered(doc, found_images[0])
                _add_figure_caption(doc, caption_text, current_h1_idx)
                continue

            # Формулы
            dm = re.match(r'^\$\$(.+?)\$\$\s*(\(\d+\.\d+\))?\s*$', raw)
            if dm:
                _add_equation(doc, dm.group(1).strip(), current_h1_idx)
                continue

            if '$$' in raw:
                parts = re.split(r'\$\$(.+?)\$\$', raw)
                p = doc.add_paragraph()
                _apply_p_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_ind=cfg.FIRST_LINE_INDENT)
                for idx, part in enumerate(parts):
                    if idx % 2 == 0:
                        if part.strip():
                            run = p.add_run(part)
                            run.font.name = cfg.FONT_NAME
                            run.font.size = cfg.FONT_SIZE_MAIN
                    else:
                        _add_inline_math(p, part.strip())
                continue

            # Таблицы (caption + lookahead)
            if re.match(r'^Таблица\s+\d', raw) or re.match(r'^\*Таблица\s+\d', raw):
                caption_text = raw
                if i + 1 < len(ast_tokens) and ast_tokens[i + 1]['type'] == 'table':
                    _, next_node = next(token_iterator)
                    node = next_node
                    tp = 'table'
                    node['_caption'] = caption_text
                else:
                    _add_table_caption(doc, caption_text, current_h1_idx)
                    continue

            if tp == 'paragraph' and 'children' in node and node['children']:
                # === ПРОВЕРКА НА СПИСКИ (РАЗДЕЛЕНИЕ) ===
                split_marker_chars = r'0-9\-\u2013\u2014\.\•'
                is_where_clause = raw.lower().startswith('где ') or raw.lower().startswith('здесь ')
                
                needs_split = False
                if is_where_clause and ';' in raw:
                    needs_split = True
                elif re.search(r';\s*[' + split_marker_chars + r']', raw):
                    needs_split = True
                elif re.search(r'(?<=[:])\s*(?=[' + split_marker_chars + r'])', raw):
                    needs_split = True
                
                if needs_split:
                    _add_rich_paragraph(doc, raw)
                    continue
                # ========================================

                p = doc.add_paragraph()
                _apply_p_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_ind=cfg.FIRST_LINE_INDENT)
                _render_inline_children(p, node['children'])
                continue

            _add_rich_paragraph(doc, _auto_wrap_subscripts(raw))
            continue

        if tp == 'table':
            caption = node.get('_caption', '')
            header_cells = []
            data_rows    = []
            if 'children' in node:
                for child in node['children']:
                    if child['type'] == 'table_head':
                        header_cells = [_flat_text(c) for c in child.get('children', [])]
                    elif child['type'] == 'table_body':
                        for row in child.get('children', []):
                            data_rows.append([_flat_text(c) for c in row.get('children', [])])
            if header_cells:
                _add_table(doc, header_cells, data_rows, caption, current_h1_idx)
            continue

        if tp == 'list':
            ordered = node.get('attrs', {}).get('ordered', False)
            for li_idx, item in enumerate(node.get('children', [])):
                text = _flat_text(item).strip()
                prefix = f"{li_idx + 1}. " if ordered else "-- "
                _add_rich_paragraph(doc, _auto_wrap_subscripts(prefix + text))
            continue

        if tp == 'code':
            raw = node.get('raw', node.get('text', ''))
            info = node.get('attrs', {}).get('info', '')
            _add_source_code(doc, raw, info)
            continue

        if tp == 'thematic_break':
            doc.add_page_break()
            continue

        if tp in ('blank_line', 'newline'):
            continue

        raw = node.get('raw', node.get('text', ''))
        if raw and raw.strip():
            _add_rich_paragraph(doc, _auto_wrap_subscripts(raw.strip()))

    if output is None:
        out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'data')
        os.makedirs(out_dir, exist_ok=True)
        base = os.path.splitext(os.path.basename(sample_path))[0]
        v = 1
        while True:
            output = os.path.join(out_dir, f"{base}_builder_v{v}.docx")
            if not os.path.exists(output): break
            v += 1

    wu.save_document_safe(doc, output)
    print(f"   Параграфов: {len(doc.paragraphs)}")
    print(f"   Таблиц: {len(doc.tables)}")

    if not fast:
        wu.update_document_via_com(output)

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Сборка Markdown в Word GOST")
    parser.add_argument('-i', '--input', help="Путь к файлу или папке", default=None)
    parser.add_argument('-o', '--output', help="Путь для сохранения", default=None)
    parser.add_argument('-a', '--append', help="Путь к существующему Word-документу, в конец которого нужно дописать текст", default=None)
    parser.add_argument('--fast', action='store_true', help="Быстрая сборка без запуска MS Word (не обновляет оглавление и формулы)")
    
    args = parser.parse_args()
    
    input_val = args.input
    if not input_val:
        input_val = os.path.join(os.path.dirname(__file__), '..', '..', '.tmp', 'rewritten_guide_section1.md')
        
    build_document(input_val, args.output, args.fast, args.append)
