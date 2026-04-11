"""
docx_formatter.py — Форматирование Word-документов по эталону или стандарту ГОСТ.

Читает стили из эталонного документа и применяет их к рабочему,
или применяет стандартные настройки ГОСТ.

Использование:
    python docx_formatter.py --source "эталон.docx" --target "рабочий.docx"
    python docx_formatter.py --target "рабочий.docx" --standard gost
    python docx_formatter.py --target "рабочий.docx" --standard gost --save-as "результат.docx"

Зависимости: python-docx
"""

import argparse
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

sys.path.insert(0, str(Path(__file__).parent))
import word_utils as wu
import word_config as cfg


# Стандартные настройки ГОСТ для учебных/технических документов
GOST_SETTINGS = {
    "font_name": "Times New Roman",
    "font_size_body": Pt(14),
    "font_size_heading1": Pt(16),
    "font_size_heading2": Pt(14),
    "line_spacing": 1.5,
    "space_after": Pt(0),
    "space_before": Pt(0),
    "first_line_indent": Cm(1.25),
    "alignment_body": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "alignment_heading": WD_ALIGN_PARAGRAPH.CENTER,
    "margin_top": Cm(2),
    "margin_bottom": Cm(2),
    "margin_left": Cm(3),
    "margin_right": Cm(1.5),
}


def extract_styles_from_reference(ref_doc: Document) -> dict:
    """Извлекает стили из эталонного документа."""
    styles = {
        "body": {"font_name": None, "font_size": None, "line_spacing": None, "alignment": None, "first_line_indent": None},
        "headings": {},
        "page_setup": {},
    }

    # Анализ Normal-параграфов (основной текст)
    normal_paras = [p for p in ref_doc.paragraphs if p.style and p.style.name == "Normal" and p.text.strip()]
    if normal_paras:
        sample = normal_paras[0]
        if sample.runs:
            run = sample.runs[0]
            styles["body"]["font_name"] = run.font.name
            styles["body"]["font_size"] = run.font.size
        pf = sample.paragraph_format
        styles["body"]["line_spacing"] = pf.line_spacing
        styles["body"]["alignment"] = sample.alignment
        styles["body"]["first_line_indent"] = pf.first_line_indent

    # Анализ заголовков
    for para in ref_doc.paragraphs:
        if para.style and "heading" in para.style.name.lower():
            level = para.style.name
            if level not in styles["headings"]:
                heading_info = {"font_name": None, "font_size": None, "bold": None, "alignment": None}
                if para.runs:
                    run = para.runs[0]
                    heading_info["font_name"] = run.font.name
                    heading_info["font_size"] = run.font.size
                    heading_info["bold"] = run.bold
                heading_info["alignment"] = para.alignment
                styles["headings"][level] = heading_info

    # Настройки страницы
    if ref_doc.sections:
        section = ref_doc.sections[0]
        styles["page_setup"] = {
            "top_margin": section.top_margin,
            "bottom_margin": section.bottom_margin,
            "left_margin": section.left_margin,
            "right_margin": section.right_margin,
        }

    return styles


def apply_gost_formatting(doc: Document):
    """Применяет стандартные настройки ГОСТ к документу через объединенные утилиты."""
    # 1. Настройка основных стилей (Normal, Heading 1-3)
    wu.setup_gost_styles(doc)
    
    # 2. Настройки страницы
    for section in doc.sections:
        section.top_margin = cfg.MARGIN_TOP
        section.bottom_margin = cfg.MARGIN_BOTTOM
        section.left_margin = cfg.MARGIN_LEFT
        section.right_margin = cfg.MARGIN_RIGHT

    # 3. Применение специфичных фиксов из Phase 2
    wu.fix_list_indents(doc)
    wu.remove_first_page_numbering(doc)
    
    # 4. Нумерация страниц (кроме первой)
    wu.add_page_numbering(doc, smart_skip=True)



def apply_reference_formatting(doc: Document, ref_styles: dict):
    """Применяет стили из эталонного документа."""
    body = ref_styles.get("body", {})
    headings = ref_styles.get("headings", {})
    page = ref_styles.get("page_setup", {})

    # Настройки страницы
    if page:
        for section in doc.sections:
            if page.get("top_margin"):
                section.top_margin = page["top_margin"]
            if page.get("bottom_margin"):
                section.bottom_margin = page["bottom_margin"]
            if page.get("left_margin"):
                section.left_margin = page["left_margin"]
            if page.get("right_margin"):
                section.right_margin = page["right_margin"]

    # Параграфы
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"

        if style_name in headings:
            h = headings[style_name]
            if h.get("alignment") is not None:
                para.alignment = h["alignment"]
            for run in para.runs:
                if h.get("font_name"):
                    run.font.name = h["font_name"]
                if h.get("font_size"):
                    run.font.size = h["font_size"]
                if h.get("bold") is not None:
                    run.bold = h["bold"]
        else:
            # Основной текст
            if body.get("alignment") is not None:
                para.alignment = body["alignment"]
            pf = para.paragraph_format
            if body.get("line_spacing"):
                pf.line_spacing = body["line_spacing"]
            if body.get("first_line_indent"):
                pf.first_line_indent = body["first_line_indent"]
            for run in para.runs:
                if body.get("font_name"):
                    run.font.name = body["font_name"]
                if body.get("font_size"):
                    run.font.size = body["font_size"]


def main():
    parser = argparse.ArgumentParser(description="Форматирование Word-документов")
    parser.add_argument("--source", help="Путь к эталонному .docx файлу")
    parser.add_argument("--target", required=True, help="Путь к рабочему .docx файлу")
    parser.add_argument("--standard", choices=["gost"], help="Применить стандартное форматирование")
    parser.add_argument("--save-as", help="Сохранить как (не перезаписывать оригинал)")
    parser.add_argument("--no-backup", action="store_true", help="Не создавать резервную копию")
    args = parser.parse_args()

    if not args.source and not args.standard:
        print("ОШИБКА: Укажите --source (эталон) или --standard gost", file=sys.stderr)
        sys.exit(1)

    target_path = Path(args.target)
    if not target_path.exists():
        print(f"ОШИБКА: Файл не найден: {target_path}", file=sys.stderr)
        sys.exit(1)

    # Бэкап
    if not args.no_backup:
        backup = wu.create_backup(str(target_path))
        print(f"✅ Резервная копия: {backup}")

    doc = Document(str(target_path))

    if args.source:
        source_path = Path(args.source)
        if not source_path.exists():
            print(f"ОШИБКА: Эталон не найден: {source_path}", file=sys.stderr)
            sys.exit(1)

        ref_doc = Document(str(source_path))
        ref_styles = extract_styles_from_reference(ref_doc)
        apply_reference_formatting(doc, ref_styles)
        mode = f"по эталону ({source_path.name})"
    else:
        apply_gost_formatting(doc)
        mode = "по стандарту ГОСТ"

    save_path = args.save_as if args.save_as else str(target_path)
    wu.save_document_safe(doc, save_path)

    # После сохранения обновляем TOC через Word COM
    wu.update_document_via_com(save_path)

    print(f"\n✅ Документ отформатирован {mode}")
    print(f"📄 Сохранено: {save_path}")
    wu.log_operation("docx_formatter", f"Файл: {target_path.name}, Режим: {mode}")



if __name__ == "__main__":
    main()
