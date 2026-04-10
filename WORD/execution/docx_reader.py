"""
docx_reader.py — Чтение и анализ Word-документов.

Извлекает параграфы, таблицы, стили и метаданные из .docx файлов.
Результат выводится в формате text или JSON.

Использование:
    python docx_reader.py --file "путь/к/файлу.docx" --output text
    python docx_reader.py --file "путь/к/файлу.docx" --output json
    python docx_reader.py --file "путь/к/файлу.docx" --output json --styles

Зависимости: python-docx
"""

import argparse
import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu


def read_paragraphs(doc: Document) -> list[dict]:
    """Извлекает все параграфы с их стилями и форматированием."""
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        runs_data = []
        for run in para.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            }
            if run.font:
                run_info["font_name"] = run.font.name
                run_info["font_size_pt"] = run.font.size / Pt(1) if run.font.size else None
            runs_data.append(run_info)

        para_data = {
            "index": i,
            "text": para.text,
            "style": para.style.name if para.style else "Normal",
            "alignment": str(para.alignment) if para.alignment else None,
            "runs": runs_data,
        }

        # Отступы и интервалы
        pf = para.paragraph_format
        if pf:
            para_data["line_spacing"] = str(pf.line_spacing) if pf.line_spacing else None
            para_data["space_before"] = str(pf.space_before) if pf.space_before else None
            para_data["space_after"] = str(pf.space_after) if pf.space_after else None
            para_data["first_line_indent"] = str(pf.first_line_indent) if pf.first_line_indent else None

        paragraphs.append(para_data)
    return paragraphs


def read_tables(doc: Document) -> list[dict]:
    """Извлекает все таблицы из документа."""
    tables_data = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows_data.append(cells)
        tables_data.append({
            "index": i,
            "rows_count": len(table.rows),
            "cols_count": len(table.columns),
            "data": rows_data,
        })
    return tables_data


def read_styles(doc: Document) -> list[dict]:
    """Извлекает все используемые стили документа."""
    styles_data = []
    used_styles = set()
    for para in doc.paragraphs:
        if para.style and para.style.name not in used_styles:
            used_styles.add(para.style.name)
            style_info = {
                "name": para.style.name,
                "type": str(para.style.type),
                "base_style": para.style.base_style.name if para.style.base_style else None,
            }
            # Попробуем извлечь шрифт стиля
            if para.style.font:
                style_info["font_name"] = para.style.font.name
                style_info["font_size_pt"] = para.style.font.size / Pt(1) if para.style.font.size else None
                style_info["bold"] = para.style.font.bold
                style_info["italic"] = para.style.font.italic
            styles_data.append(style_info)
    return styles_data


def read_document_properties(doc: Document) -> dict:
    """Извлекает свойства документа."""
    props = doc.core_properties
    return {
        "author": props.author,
        "title": props.title,
        "subject": props.subject,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "last_modified_by": props.last_modified_by,
    }


def read_page_setup(doc: Document) -> dict:
    """Извлекает настройки страницы из первой секции."""
    section = doc.sections[0] if doc.sections else None
    if not section:
        return {}
    
    def emu_to_cm(val):
        if val is None:
            return None
        return round(val / Cm(1), 2)
    
    return {
        "page_width_cm": emu_to_cm(section.page_width),
        "page_height_cm": emu_to_cm(section.page_height),
        "top_margin_cm": emu_to_cm(section.top_margin),
        "bottom_margin_cm": emu_to_cm(section.bottom_margin),
        "left_margin_cm": emu_to_cm(section.left_margin),
        "right_margin_cm": emu_to_cm(section.right_margin),
        "orientation": str(section.orientation),
    }


def main():
    parser = argparse.ArgumentParser(description="Чтение и анализ Word-документов")
    parser.add_argument("--file", required=True, help="Путь к .docx файлу")
    parser.add_argument("--output", choices=["text", "json"], default="text", help="Формат вывода")
    parser.add_argument("--styles", action="store_true", help="Включить информацию о стилях")
    parser.add_argument("--tables", action="store_true", help="Включить таблицы")
    parser.add_argument("--properties", action="store_true", help="Включить свойства документа")
    parser.add_argument("--page-setup", action="store_true", help="Включить настройки страницы")
    parser.add_argument("--all", action="store_true", help="Включить всё")
    args = parser.parse_args()

    filepath = Path(args.file)
    if not filepath.exists():
        print(f"ОШИБКА: Файл не найден: {filepath}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(filepath))
    paragraphs = read_paragraphs(doc)

    if args.output == "json":
        result = {"file": str(filepath), "paragraphs": paragraphs}

        if args.tables or args.all:
            result["tables"] = read_tables(doc)
        if args.styles or args.all:
            result["styles"] = read_styles(doc)
        if args.properties or args.all:
            result["properties"] = read_document_properties(doc)
        if args.page_setup or args.all:
            result["page_setup"] = read_page_setup(doc)

        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        # Текстовый вывод
        print(f"=== Документ: {filepath.name} ===\n")
        
        if args.properties or args.all:
            props = read_document_properties(doc)
            print("--- Свойства ---")
            for k, v in props.items():
                if v:
                    print(f"  {k}: {v}")
            print()

        if args.page_setup or args.all:
            ps = read_page_setup(doc)
            print("--- Настройки страницы ---")
            for k, v in ps.items():
                if v:
                    print(f"  {k}: {v}")
            print()

        print("--- Параграфы ---")
        for p in paragraphs:
            if p["text"].strip():
                print(f'[{p["index"]}] ({p["style"]}) {p["text"][:300]}')

        if args.tables or args.all:
            tables = read_tables(doc)
            print(f"\n--- Таблицы ({len(tables)} шт.) ---")
            for t in tables:
                print(f'\nТаблица {t["index"]}: {t["rows_count"]} строк × {t["cols_count"]} столбцов')
                for row in t["data"][:5]:  # Первые 5 строк
                    print(f"  | {' | '.join(row)} |")
                if len(t["data"]) > 5:
                    print(f"  ... ещё {len(t['data']) - 5} строк")

        if args.styles or args.all:
            styles = read_styles(doc)
            print(f"\n--- Стили ({len(styles)} шт.) ---")
            for s in styles:
                font_info = f", шрифт: {s.get('font_name', '—')} {s.get('font_size_pt', '—')}пт"
                print(f"  {s['name']}{font_info}")


if __name__ == "__main__":
    main()
