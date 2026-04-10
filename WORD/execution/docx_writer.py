"""
docx_writer.py — Редактирование Word-документов.

Замена текста, вставка параграфов, добавление таблиц.
Перед каждым изменением автоматически создаётся резервная копия.

Использование:
    python docx_writer.py --file "файл.docx" --replace "старый текст" "новый текст"
    python docx_writer.py --file "файл.docx" --replace-at 5 "новый текст параграфа"
    python docx_writer.py --file "файл.docx" --insert-after 10 "текст нового параграфа"
    python docx_writer.py --file "файл.docx" --delete 15
    python docx_writer.py --file "файл.docx" --fix-decimals
    python docx_writer.py --file "файл.docx" --no-backup --replace "old" "new"

Зависимости: python-docx
"""

import argparse
import sys
from pathlib import Path
from docx import Document

# Добавляем путь к utils
sys.path.insert(0, str(Path(__file__).parent))
from utils.docx_utils import create_backup, fix_decimal_separators, log_operation


def replace_text_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """
    Заменяет текст в параграфе с сохранением форматирования.
    Работает на уровне runs для сохранения стилей.
    """
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Простой случай: замена в одном run
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # Сложный случай: текст разбит между runs
    # Собираем весь текст, заменяем, перезаписываем
    new_full = full_text.replace(old_text, new_text)
    if paragraph.runs:
        # Сохраняем форматирование первого run
        first_run = paragraph.runs[0]
        first_run.text = new_full
        # Очищаем остальные runs
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def replace_text_globally(doc: Document, old_text: str, new_text: str) -> int:
    """Заменяет текст во всех параграфах документа. Возвращает количество замен."""
    count = 0
    for para in doc.paragraphs:
        if replace_text_in_paragraph(para, old_text, new_text):
            count += 1
    # Также проверяем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_text_in_paragraph(para, old_text, new_text):
                        count += 1
    return count


def replace_paragraph_at(doc: Document, index: int, new_text: str) -> bool:
    """Заменяет текст параграфа по индексу."""
    if 0 <= index < len(doc.paragraphs):
        para = doc.paragraphs[index]
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_text
        return True
    return False


def insert_paragraph_after(doc: Document, index: int, text: str, style: str = "Normal") -> bool:
    """Вставляет новый параграф после указанного индекса."""
    if 0 <= index < len(doc.paragraphs):
        ref_para = doc.paragraphs[index]
        new_para = ref_para.insert_paragraph_after(text)
        try:
            new_para.style = doc.styles[style]
        except KeyError:
            pass  # Стиль не найден, оставляем по умолчанию
        return True
    return False


def delete_paragraph(doc: Document, index: int) -> bool:
    """Удаляет параграф по индексу."""
    if 0 <= index < len(doc.paragraphs):
        para = doc.paragraphs[index]
        p_element = para._element
        p_element.getparent().remove(p_element)
        return True
    return False


def fix_all_decimals(doc: Document) -> int:
    """Заменяет точки на запятые во всех десятичных числах. Возвращает количество изменённых параграфов."""
    count = 0
    for para in doc.paragraphs:
        original = para.text
        fixed = fix_decimal_separators(original)
        if fixed != original:
            for run in para.runs:
                run.text = fix_decimal_separators(run.text)
            count += 1
    return count


def main():
    parser = argparse.ArgumentParser(description="Редактирование Word-документов")
    parser.add_argument("--file", required=True, help="Путь к .docx файлу")
    parser.add_argument("--no-backup", action="store_true", help="Не создавать резервную копию")
    parser.add_argument("--save-as", help="Сохранить как новый файл (не перезаписывать оригинал)")

    # Операции (можно несколько за раз)
    parser.add_argument("--replace", nargs=2, metavar=("СТАРЫЙ", "НОВЫЙ"),
                        action="append", help="Заменить текст (можно указать несколько раз)")
    parser.add_argument("--replace-at", nargs=2, metavar=("ИНДЕКС", "ТЕКСТ"),
                        help="Заменить текст параграфа по индексу")
    parser.add_argument("--insert-after", nargs=2, metavar=("ИНДЕКС", "ТЕКСТ"),
                        help="Вставить параграф после указанного индекса")
    parser.add_argument("--delete", type=int, metavar="ИНДЕКС",
                        help="Удалить параграф по индексу")
    parser.add_argument("--fix-decimals", action="store_true",
                        help="Заменить точки на запятые в десятичных числах")

    args = parser.parse_args()

    filepath = Path(args.file)
    if not filepath.exists():
        print(f"ОШИБКА: Файл не найден: {filepath}", file=sys.stderr)
        sys.exit(1)

    # Создаем бэкап
    if not args.no_backup:
        backup_path = create_backup(str(filepath))
        print(f"✅ Резервная копия: {backup_path}")

    doc = Document(str(filepath))
    changes = []

    # Выполняем операции
    if args.replace:
        for old_text, new_text in args.replace:
            count = replace_text_globally(doc, old_text, new_text)
            changes.append(f"Замена '{old_text}' → '{new_text}': {count} замен")

    if args.replace_at:
        idx, text = int(args.replace_at[0]), args.replace_at[1]
        if replace_paragraph_at(doc, idx, text):
            changes.append(f"Параграф [{idx}] заменён")
        else:
            changes.append(f"ОШИБКА: Параграф [{idx}] не найден")

    if args.insert_after:
        idx, text = int(args.insert_after[0]), args.insert_after[1]
        if insert_paragraph_after(doc, idx, text):
            changes.append(f"Вставлен параграф после [{idx}]")
        else:
            changes.append(f"ОШИБКА: Параграф [{idx}] не найден")

    if args.delete is not None:
        if delete_paragraph(doc, args.delete):
            changes.append(f"Удалён параграф [{args.delete}]")
        else:
            changes.append(f"ОШИБКА: Параграф [{args.delete}] не найден")

    if args.fix_decimals:
        count = fix_all_decimals(doc)
        changes.append(f"Исправлены десятичные разделители: {count} параграфов")

    # Сохраняем
    save_path = args.save_as if args.save_as else str(filepath)
    doc.save(save_path)

    print(f"\n📄 Сохранено: {save_path}")
    print(f"📝 Изменения ({len(changes)}):")
    for c in changes:
        print(f"   • {c}")

    log_operation("docx_writer", f"Файл: {filepath.name}, Изменений: {len(changes)}")


if __name__ == "__main__":
    main()
