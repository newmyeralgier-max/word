"""
extract_title.py — Выкачать титульник из оригинала БР.docx в отдельный файл.
Сохраняет ВСЁ до первого структурного элемента (ВВЕДЕНИЕ, СОДЕРЖАНИЕ и т.п.)
как есть — без изменений.
"""
from docx import Document
from docx.oxml.ns import qn
import copy
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

STRUCTURAL_KEYWORDS = {
    'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ',
    'СПИСОК ЛИТЕРАТУРЫ', 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
    'СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ',
    'РЕФЕРАТ', 'АННОТАЦИЯ',
    'ОПРЕДЕЛЕНИЯ', 'ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ',
    'ПРИЛОЖЕНИЕ', 'ПРИЛОЖЕНИЯ',
    # 'ЗАДАНИЕ' убрано — задание это ЧАСТЬ титульника, не структурный элемент
}

def extract_title_section(input_path, output_path):
    """Скопировать первую секцию (титульник + задание) как есть."""
    doc = Document(input_path)
    
    # Находим границу: первый параграф с текстом из STRUCTURAL_KEYWORDS
    title_end_idx = None
    for i, p in enumerate(doc.paragraphs):
        upper = p.text.strip().upper().rstrip('.')
        if upper in STRUCTURAL_KEYWORDS:
            title_end_idx = i
            break
    
    if title_end_idx is None:
        print("[!] Не найден структурный элемент — берём всю первую секцию")
        # Берём все параграфы первой секции
        title_end_idx = len(doc.paragraphs)
    
    # Сохраняем XML-элементы титульника
    body = doc._element.find(qn('w:body'))
    all_elements = list(body)
    
    # Находим индеки элементов до title_end_idx
    # doc.paragraphs может не включать таблицы и другие элементы
    # Поэтому считаем все дочерние элементы body до нужного параграфа
    
    title_elements = []
    para_count = 0
    for elem in all_elements:
        title_elements.append(elem)
        if elem.tag == qn('w:p') or elem.tag.endswith('}p'):
            para_count += 1
            if para_count >= title_end_idx:
                break
        elif elem.tag == qn('w:tbl') or elem.tag.endswith('}tbl'):
            # Таблицы тоже включаем
            pass
    
    # Создаём новый документ с титульником
    new_doc = Document()
    new_body = new_doc._element.find(qn('w:body'))
    # Удаляем пустой параграф по умолчанию
    for child in list(new_body):
        new_body.remove(child)
    
    # Копируем элементы титульника
    for elem in title_elements:
        new_body.append(copy.deepcopy(elem))
    
    # Копируем свойства секции (поля, размер и т.д.) из первой секции оригинала
    if doc.sections and new_doc.sections:
        sec = doc.sections[0]
        new_sec = new_doc.sections[0]
        try:
            new_sec.page_width = sec.page_width
            new_sec.page_height = sec.page_height
            new_sec.orientation = sec.orientation
            new_sec.left_margin = sec.left_margin
            new_sec.right_margin = sec.right_margin
            new_sec.top_margin = sec.top_margin
            new_sec.bottom_margin = sec.bottom_margin
            new_sec.header_distance = sec.header_distance
            new_sec.footer_distance = sec.footer_distance
        except Exception as e:
            print(f"[!] Не удалось скопировать свойства секции: {e}")
    
    new_doc.save(output_path)
    print(f"[OK] Титульник сохранён: {output_path}")
    print(f"     Параграфов: {para_count}, элементов: {len(title_elements)}")


def restore_title_section(title_path, gost_path, output_path):
    """Заменить первую секцию в GOST-документе на оригинальный титульник."""
    title_doc = Document(title_path)
    gost_doc = Document(gost_path)
    
    # Находим границу основной зоны в GOST
    gost_body = gost_doc._element.find(qn('w:body'))
    
    STRUCTURAL_KEYWORDS_RESTORE = {
        'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ',
        'СПИСОК ЛИТЕРАТУРЫ', 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
        'СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ',
        'РЕФЕРАТ', 'АННОТАЦИЯ', 'ЗАДАНИЕ',
        'ОПРЕДЕЛЕНИЯ', 'ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ',
        'ПРИЛОЖЕНИЕ', 'ПРИЛОЖЕНИЯ',
    }
    
    # Находим индекс первого структурного элемента в GOST
    main_start_elem = None
    main_start_idx = None
    all_gost_elems = list(gost_body)
    para_idx = 0
    for i, elem in enumerate(all_gost_elems):
        if elem.tag == qn('w:p') or elem.tag.endswith('}p'):
            # Получаем текст параграфа
            texts = elem.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in texts).strip().upper().rstrip('.')
            if text in STRUCTURAL_KEYWORDS_RESTORE:
                main_start_elem = elem
                main_start_idx = i
                break
            para_idx += 1
    
    if main_start_idx is None:
        print("[!] Не найден структурный элемент в GOST — заменяем первую секцию целиком")
        main_start_idx = 0
    
    # Собираем элементы титульника из title_doc
    title_body = title_doc._element.find(qn('w:body'))
    title_elems = list(title_body)
    # Убираем sectPr из титульника (он свой у GOST-документа)
    title_elems_no_sect = [e for e in title_elems 
                           if e.tag != qn('w:sectPr') and not e.tag.endswith('}sectPr')]
    
    # Удаляем все элементы до main_start_idx из GOST
    elems_to_remove = all_gost_elems[:main_start_idx]
    for elem in elems_to_remove:
        gost_body.remove(elem)
    
    # Вставляем элементы титульника в начало GOST
    title_xml = copy.deepcopy(title_elems_no_sect)
    # Вставляем перед первым элементом основной зоны
    first_remaining = list(gost_body)[0]
    for i, elem in enumerate(title_xml):
        first_remaining.addprevious(elem)
    
    # Восстанавливаем поля первой секции из оригинала
    if title_doc.sections and gost_doc.sections:
        orig_sec = title_doc.sections[0]
        gost_sec = gost_doc.sections[0]
        try:
            gost_sec.page_width = orig_sec.page_width
            gost_sec.page_height = orig_sec.page_height
            gost_sec.orientation = orig_sec.orientation
            gost_sec.left_margin = orig_sec.left_margin
            gost_sec.right_margin = orig_sec.right_margin
            gost_sec.top_margin = orig_sec.top_margin
            gost_sec.bottom_margin = orig_sec.bottom_margin
            gost_sec.header_distance = orig_sec.header_distance
            gost_sec.footer_distance = orig_sec.footer_distance
        except Exception as e:
            print(f"[!] Не удалось скопировать свойства секции: {e}")
    
    # ★ Также восстанавливаем вторую секцию (содержание/альбомная) из оригинала
    # В оригинале S2 = LANDSCAPE, форматер мог её убить
    if len(title_doc.sections) > 1 and len(gost_doc.sections) > 1:
        orig_sec2 = title_doc.sections[1]
        gost_sec2 = gost_doc.sections[1]
        try:
            gost_sec2.page_width = orig_sec2.page_width
            gost_sec2.page_height = orig_sec2.page_height
            gost_sec2.orientation = orig_sec2.orientation
        except Exception as e:
            print(f"[!] Не удалось восстановить ориентацию S2: {e}")
    
    gost_doc.save(output_path)
    print(f"[OK] Титульник восстановлен: {output_path}")


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description="Выкачать/восстановить титульник БР")
    parser.add_argument('--extract', action='store_true', help="Выкачать титульник из оригинала")
    parser.add_argument('--restore', action='store_true', help="Восстановить титульник в GOST")
    parser.add_argument('-i', '--input', help="Входной файл")
    parser.add_argument('-t', '--title', default=r"D:\1. Project\Word\data\БР_титульник.docx",
                        help="Файл титульника")
    parser.add_argument('-o', '--output', help="Выходной файл")
    
    args = parser.parse_args()
    
    # Пути по умолчанию (Linux WSL формат)
    default_br = '/mnt/d/1. Project/Word/data/БР.docx'
    default_gost = '/mnt/d/1. Project/Word/data/БР_GOST.docx'
    default_title = '/mnt/d/1. Project/Word/data/БР_титульник.docx'
    
    if args.extract:
        inp = args.input or default_br
        out = args.title or default_title
        extract_title_section(inp, out)
    elif args.restore:
        inp = args.input or default_gost
        out = args.output or inp
        title = args.title or default_title
        restore_title_section(title, inp, out)
    else:
        print("Укажите --extract или --restore")
