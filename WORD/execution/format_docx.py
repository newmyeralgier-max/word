"""
format_docx.py — ГОСТ-форматер для Word-документов.
Версия 3.2 — Фиксы по скринам недочётов БР.

Принципы:
  1. НЕ разрушаем содержимое — модифицируем runs, а не перезаписываем p.text
  2. Опираемся на СТИЛИ Word, а не на хрупкие эвристики по тексту
  3. Строго следуем word_config.py — единый источник правды
  4. Разрыв страницы перед заголовками H1 — требование ГОСТ

v3.2 фиксы:
  ФИКС 7: ВВЕДЕНИЕ в стиле Title → STRUCTURAL_H1 (не иммунитет)
  ФИКС 8: "1.ТЕКСТ" без пробела — распознаётся как заголовок
  ФИКС 9: Формулы — очистка TAB в runs перед установкой tab stops
  ФИКС 10: Точка в конце заголовков — удаляется
  ФИКС 11: Пустые OMML-параграфы — удаляются
  ФИКС 12: "Таблица 4.1." — точка после номера срезается
  ФИКС 13: Стили HTML Preformatted/Normal (Web) → Body Text
  ФИКС 14: Удаление пустых страниц в хвосте (приложения)
"""

import sys, os, re, copy
from docx import Document
from docx.shared import Pt, Cm, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import word_config as cfg
import word_utils  as wu


# ──────────────────────────────────────────────────────────────────
#  Константы детекции
# ──────────────────────────────────────────────────────────────────

# Структурные элементы ГОСТ — всегда по центру, без нумерации, КАПСОМ
STRUCTURAL_KEYWORDS = {
    'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ',
    'СПИСОК ЛИТЕРАТУРЫ', 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
    'СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ',
    'РЕФЕРАТ', 'АННОТАЦИЯ', 'ЗАДАНИЕ',
    'ОПРЕДЕЛЕНИЯ', 'ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ',
    'ПРИЛОЖЕНИЕ', 'ПРИЛОЖЕНИЯ',
    'ПРЕЗЕНТАЦИОННЫЕ МАТЕРИАЛЫ',  # ★ ФИКС 33
}

# Ключевые слова, после которых начинается основная зона текста
MAIN_ZONE_TRIGGERS = STRUCTURAL_KEYWORDS

# Паттерн подписи к рисунку
_RE_FIGURE = re.compile(
    r'^(Рис\.|Рисунок)\s*(\d+[\.\d]*)\s*[.\s—–\-]*\s*(.*)', re.I
)
# Паттерн подписи к таблице
# ★ ФИКС 12: "Таблица 4.1." — точка после номера срезается
_RE_TABLE_CAP = re.compile(
    r'^(Табл\.|Таблица)\s*(\d+(?:\.\d+)?)\.?\s*[.\s—–\-]*\s*(.*)', re.I
)
# Паттерн нумерации формулы в конце строки: (1), (1.1), (2.3.4)
_RE_FORMULA_NUM = re.compile(r'\((\d+[\.\d]*)\)\s*$')
# Паттерн маркированного списка
_RE_LIST_PREFIX = re.compile(r'^[-—–•]\s+')
# Паттерн нумерованного списка: а), б), 1), 2)
_RE_NUM_LIST = re.compile(r'^([а-яёА-ЯЁ]\)|\d+\))\s+')
# Паттерн заголовка раздела: "1. ТЕКСТ ЗАГЛАВНЫМИ" или "1.1 ТЕКСТ ЗАГЛАВНЫМИ"
_RE_SECTION_HEADING = re.compile(r'^(\d+(?:\.\d+)?)\s+[\-—]?\s*([А-ЯЁ]{2,}|[А-ЯЁ].*[А-ЯЁ])')
# Паттерн «где» — расшифровка формулы
_RE_WHERE_LINE = re.compile(r'^где\s+', re.I)
# Ключевые слова начала библиографии
_BIBLIO_KEYWORDS = {
    'СПИСОК ЛИТЕРАТУРЫ', 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
    'СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ',
}

# Счётчик заголовков из List Paragraph (ФИКС 15)
_lp_heading_counter = [0]


# ──────────────────────────────────────────────────────────────────
#  Вспомогательные функции
# ──────────────────────────────────────────────────────────────────

def _set_run_font(run, font_name=cfg.FONT_NAME, font_size=cfg.FONT_SIZE_MAIN,
                  bold=None, italic=None, color=cfg.COLOR_BLACK):
    """Безопасно установить шрифт run, не трогая моноширинные (код).
    ★ ФИКС A: НЕ перезаписываем size если он None — наследование от стиля важнее.
    ★ Если font_size=None — вообще не трогаем размер run."""
    # Не трогаем шрифт кода
    if run.font.name and run.font.name.lower() in ('consolas', 'courier new'):
        return
    # ★ ГОСТ требует TNR — ставим всегда
    run.font.name  = font_name
    # ★ ФИКС A: Если font_size=None — не трогаем размер (наследуется от стиля)
    # Если font_size задан — ставим ТОЛЬКО если run.font.size уже был задан явно (не None)
    # Если run.font.size is None (наследуется) — тоже не перезаписываем, стиль знает лучше
    if font_size is not None and run.font.size is not None:
        run.font.size  = font_size
    if bold   is not None: run.bold   = bold
    if italic is not None: run.italic = italic
    if color  is not None: run.font.color.rgb = color


def _clear_indents_and_set(p, align=None, first_line_ind=None,
                           left_ind=Cm(0), right_ind=Cm(0),
                           space_before=Pt(0), space_after=Pt(0),
                           line_spacing=cfg.LINE_SPACING,
                           keep_next=False, widow=True):
    """Установить формат абзаца, сбросив лишние отступы."""
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.left_indent        = left_ind
    pf.right_indent       = right_ind
    pf.first_line_indent  = first_line_ind if first_line_ind is not None else Cm(0)
    pf.space_before       = space_before
    pf.space_after        = space_after
    pf.line_spacing       = line_spacing
    pf.keep_with_next     = keep_next
    pf.widow_control      = widow


def _has_page_break_before(p):
    """Проверить, есть ли уже разрыв страницы перед этим параграфом."""
    # ★ БАГ 13: Проверяем pageBreakBefore в свойствах текущего параграфа
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None:
        return True
    prev = p._element.getprevious()
    while prev is not None:
        # Проверяем: содержит ли предыдущий элемент разрыв страницы
        # ★ ФИКС: Используем qn() для неймспейсов вместо прямого XPath
        breaks = prev.findall('.//' + qn('w:br'))
        page_breaks = [br for br in breaks 
                       if br.get(qn('w:type')) == 'page']
        if page_breaks:
            return True
        # Если предыдущий параграф пустой — продолжаем искать
        text_el = prev.findall('.//' + qn('w:t'))
        text = ''.join(t.text or '' for t in text_el).strip()
        if not text:
            prev = prev.getprevious()
            continue
        break
    return False


def _remove_empty_paragraphs_before(p):
    """Удалить пустые параграфы непосредственно перед данным."""
    removed = 0
    while True:
        prev_el = p._element.getprevious()
        if prev_el is None:
            break
        # Это параграф?
        if prev_el.tag != qn('w:p') and not prev_el.tag.endswith('}p'):
            break
        text_el = prev_el.xpath('.//w:t')
        text = ''.join(t.text or '' for t in text_el).strip()
        has_br = prev_el.xpath('.//w:br[@w:type="page"]')
        if not text and not has_br:
            prev_el.getparent().remove(prev_el)
            removed += 1
        else:
            break
    return removed


def _insert_page_break_before(paragraph):
    """Вставить разрыв страницы перед указанным параграфом.
    ★ ФИКС 21: Используем page_break_before=True вместо нового параграфа.
    Старый метод (новый параграф + w:br) создавал пустые страницы."""
    # Не вставлять если уже есть разрыв перед заголовком
    if _has_page_break_before(paragraph):
        return
    # #v4_REMOVED: _remove_empty_paragraphs_before(paragraph)
    # PRIMUM NON NOCERE: не удаляем параграфы — только ставим page_break_before
    
    # ★ ФИКС 21: page_break_before — не создаёт лишний параграф
    paragraph.paragraph_format.page_break_before = True


def _is_heading_style(style_name):
    """Определяет, является ли стиль заголовком."""
    if not style_name:
        return False
    return style_name.startswith('Heading') or style_name.startswith('Заголовок')


def _heading_level(style_name):
    """Извлечь уровень заголовка (1, 2, 3) из имени стиля. 0 — не заголовок."""
    if not _is_heading_style(style_name):
        return 0
    m = re.search(r'(\d+)', style_name)
    return int(m.group(1)) if m else 1


def _is_toc_style(style_name):
    """Определяет, является ли стиль элементом оглавления."""
    if not style_name:
        return False
    return style_name.startswith('TOC') or style_name.startswith('Содержание')


def _has_math(paragraph):
    """Содержит ли параграф OMML-формулу."""
    return len(paragraph._element.xpath('.//m:oMath')) > 0


def _has_image(paragraph):
    """Содержит ли параграф изображение (drawing/inlineImage)."""
    return (len(paragraph._element.xpath('.//w:drawing')) > 0 or
            len(paragraph._element.xpath('.//w:pict')) > 0)


# ──────────────────────────────────────────────────────────────────
#  Вставка автооглавления
# ──────────────────────────────────────────────────────────────────

def _insert_toc_before_intro(doc):
    """Вставить автооглавление непосредственно перед ВВЕДЕНИЕМ/СОДЕРЖАНИЕМ."""
    if not doc.paragraphs:
        return

    # Ищем целевой параграф (первый структурный элемент)
    target = None
    for p in doc.paragraphs:
        if p.text.strip().upper() in MAIN_ZONE_TRIGGERS:
            target = p
            break
    if target is None:
        # Если не нашли — вставляем в начало
        target = doc.paragraphs[0]

    # Заголовок СОДЕРЖАНИЕ
    ph = target.insert_paragraph_before()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _clear_indents_and_set(ph, align=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(24), space_after=Pt(12))
    rh = ph.add_run('СОДЕРЖАНИЕ')
    rh.bold      = False
    rh.font.name = cfg.FONT_NAME
    rh.font.size = cfg.FONT_SIZE_TOC_TITLE

    # TOC field
    pt = target.insert_paragraph_before()
    pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _clear_indents_and_set(pt, align=WD_ALIGN_PARAGRAPH.LEFT)
    rt = pt.add_run()
    for tag, attr, text in [
        ('w:fldChar', 'begin', None),
        ('w:instrText', None, 'TOC \\o "1-3" \\h \\z \\u'),
        ('w:fldChar', 'separate', None),
    ]:
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = text
        else:
            el.set(qn('w:fldCharType'), attr)
        rt._r.append(el)

    # placeholder
    pr = OxmlElement('w:r')
    pt_el = OxmlElement('w:t')
    pt_el.text = '[Оглавление собрано автоматически]'
    pr.append(pt_el)
    rt._r.append(pr)

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    rt._r.append(end)

    # Разрыв страницы после оглавления
    _insert_page_break_before(target)


# ──────────────────────────────────────────────────────────────────
#  Форматирование формул
# ──────────────────────────────────────────────────────────────────

def _get_section_for_paragraph(p):
    """★ v4.1: Определить секцию, в которой находится параграф.
    Идём по body от начала, подсчитывая sectPr (разрывы секций).
    Параграф принадлежит секции N, если он находится после (N-1)-го разрыва
    и до N-го разрыва.
    """
    try:
        p_elem = p._element
        body = p_elem.getparent()
        while body is not None and body.tag != qn('w:body'):
            body = body.getparent()
        if body is None:
            return None

        doc_part = p.part.document
        sections = doc_part.sections
        if not sections:
            return None

        # Простой подход: считаем sectPr в pPr ДО нашего параграфа
        # Каждый sectPr в pPr означает конец секции
        sect_count = 0
        for child in body:
            if child is p_elem:
                break
            # sectPr внутри pPr параграфа = конец секции
            pPr = child.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                sect_count += 1
            # Таблицы — пропускаем (не параграфы)
        
        # sect_count = количество секций перед нашим параграфом
        # Параграф в секции с индексом sect_count
        if sect_count < len(sections):
            return sections[sect_count]
        return sections[-1]  # fallback
    except Exception:
        return None


def _get_right_tab_position(p):
    """★ ФИКС 27: Вычислить позицию правой табуляции для номера формулы
    по полям секции, в которой находится параграф.
    page_width - left_margin - right_margin = ширина текста
    Формула по центру, номер — по правому краю текста.
    ★ v4.1: Используем _get_section_for_paragraph для точного определения секции.
    """
    try:
        sec = _get_section_for_paragraph(p)
        if sec is None:
            # Fallback: берём последнюю не-титульную секцию
            doc_part = p.part.document
            sections = doc_part.sections
            if len(sections) > 1:
                sec = sections[1]
            elif sections:
                sec = sections[0]
            else:
                return Cm(16.5)
        
        pw = sec.page_width or Mm(210)
        lm = sec.left_margin or Cm(0)
        rm = sec.right_margin or Cm(0)
        right_tab = pw - lm - rm
        return right_tab
    except Exception:
        return Cm(16.5)  # fallback для A4 30+15мм


def _format_formula(p):
    """Выравнивание формулы (центр) и номера (справа) через табуляцию.
    ★ v4.2 PRIMUM NON NOCERE: НЕ трогаем runs вообще — только alignment + tab stops.
    Старые <w:tab/> в runs остаются как есть. Они могут не идеально выравниваться
    с новыми tab stops, но контент НЕ теряется.
    ГОСТ: [формула по центру] TAB [(номер) по правому краю]."""
    
    # Центрирование параграфа + табуляции
    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_ind=Cm(0),
                           space_before=Pt(6), space_after=Pt(6))
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.clear_all()
    right_tab = _get_right_tab_position(p)
    center_tab = int(right_tab / 2)
    tab_stops.add_tab_stop(center_tab, WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
    
    # ★ v4.2: Вся манипуляция runs ОТКЛЮЧЕНА — см. git history v4.1


# ──────────────────────────────────────────────────────────────────
#  Форматирование таблиц
# ──────────────────────────────────────────────────────────────────

def _normalize_tables(doc):
    """Применяет ГОСТ к таблицам: шрифт, интервал, повтор шапки, границы, выравнивание."""
    for table in doc.tables:
        # ★ ФИКС 29: Выравнивание таблиц по центру страницы
        # Без jc таблицы съезжают вправо или влево
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        # Удаляем существующий jc
        for jc in tblPr.findall(qn('w:jc')):
            tblPr.remove(jc)
        # Ставим center
        jc_el = OxmlElement('w:jc')
        jc_el.set(qn('w:val'), 'center')
        tblPr.append(jc_el)

        # Повтор шапки на каждой странице
        if len(table.rows) > 0:
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            if not trPr.xpath('w:tblHeader'):
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        # Границы по ГОСТ (тонкие чёрные линии)
        wu.set_table_border_gost(table)

        # Выравнивание текста и шрифт внутри таблиц
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip() and not _has_math(p):
                        continue
                    _clear_indents_and_set(p, first_line_ind=Cm(0),
                                           space_before=Pt(2), space_after=Pt(2),
                                           line_spacing=1.0)
                    for run in p.runs:
                        _set_run_font(run, font_size=cfg.FONT_SIZE_TABLE)


# ──────────────────────────────────────────────────────────────────
#  Классификатор параграфов
# ──────────────────────────────────────────────────────────────────

class ParagraphType:
    """Типы параграфов для маршрутизации."""
    TITLE_ZONE     = 'title_zone'      # Титульный лист / бланк (иммунитет)
    STRUCTURAL_H1  = 'structural_h1'   # ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и т.п.
    HEADING        = 'heading'         # Заголовок по стилю Heading/Заголовок
    TOC_ENTRY      = 'toc_entry'       # Элемент оглавления
    FORMULA        = 'formula'         # OMML-формула
    MANUAL_FORMULA = 'manual_formula'  # ★ ФИКС 23: Ручная текстовая формула (не OMML)
    FIGURE_CAP     = 'figure_caption'  # Подпись к рисунку
    TABLE_CAP      = 'table_caption'   # Подпись к таблице
    LIST_ITEM      = 'list_item'       # Элемент списка
    WHERE_LINE     = 'where_line'      # «где ...» — расшифровка формулы
    BODY           = 'body'            # Обычный текст
    EMPTY          = 'empty'           # Пустой параграф


def _is_section_heading(text):
    """Определить, похож ли текст на заголовок раздела: '1. ТЕКСТ' или '1.1 Текст'."""
    # Берём только ПЕРВУЮ строку параграфа — остальное может быть обычным текстом
    first_line = text.split('\n')[0].strip()
    if not first_line:
        return False, 0
    
    # ★ ФИКС 8: Паттерн "1.ТЕКСТ" без пробела — тоже заголовок (часто в БР)
    # В оригинале "1.ХАРАКТЕРИСТИКА..." — без пробела после точки
    m = re.match(r'^(\d+(?:\.\d+)*)\s*[.\s]\s*(.+)', first_line)
    if not m:
        # Пробуем без пробела: "1.ТЕКСТ" → num="1", rest="ТЕКСТ"
        m2 = re.match(r'^(\d+)\.([А-ЯЁ]{2,}.*)', first_line)
        if m2:
            m = m2
        else:
            return False, 0
    num_part = m.group(1)  # e.g. "1", "1.1", "1.2.17"
    rest = m.group(2).strip()
    
    # ★ Отсеиваем ложные срабатывания:
    # - Номер с 4+ уровнями (1.2.17.5) — это пункт списка, не заголовок.
    # 3 уровня (1.5.1) допускаем как подраздел уровня 3 (Heading 3).
    if num_part.count('.') >= 3:
        return False, 0
    # - Текст содержит матем. символы/formula: Iном, S=, P=, ≥, ≤
    if re.search(r'[=≥≤<>]', rest):
        return False, 0
    if re.match(r'^[A-Za-zА-Яа-я]{1,4}[ _]', rest) and re.search(r'\d', rest[:10]):
        # Скорее всего "Iном аппарата" — технический текст, не заголовок
        return False, 0
    # - Короткий текст (<10 символов) после номера — скорее пункт списка
    if len(rest) < 10:
        return False, 0
    
    # Если текст ЗАГЛАВНЫМИ — точно заголовок раздела
    upper_ratio = sum(1 for c in rest if c.isupper()) / max(len(rest), 1)
    if upper_ratio > 0.4:
        depth = num_part.count('.')
        # 0 точек — H1, 1 — H2, 2 — H3
        return True, max(1, min(depth + 1, 3))
    
    # ★ ФИКС 17: Если номер с точкой (1.1, 2.3) — это подраздел, порог ниже (>8 символов)
    # ГОСТ-заголовки подразделов бывают короткими: "5.1. Выбор схем" (11 символов)
    has_sub_num = '.' in num_part
    # ★ ФИКС V4.1 (пост-фидбек): заголовки не содержат середину предложения.
    #   Если в rest есть «. » (точка + пробел + буква) до самого конца — это абзац.
    #   Также ограничиваем длину: заголовок редко > 120 символов.
    if re.search(r'\.\s+[А-ЯЁа-яёA-Za-z]', rest[:-1]):
        return False, 0
    if len(rest) > 120:
        return False, 0
    # ★ ФИКС V4.1: без под-номера «1 Текст» / «1. Текст» с нижним регистром —
    # почти всегда пункт списка, а не заголовок главы. Требуем либо под-номер,
    # либо высокий процент заглавных (обработан выше).
    if not has_sub_num:
        return False, 0
    min_len = 8
    if len(rest) > min_len and rest[0].isupper():
        depth = num_part.count('.')
        return True, max(2, min(depth + 1, 3))

    return False, 0


def _is_manual_formula(text, style_name):
    """★ ФИКС 23: Определить, является ли строка ручной (текстовой) формулой.
    
    Признаки:
    1. Стиль "Формула*" — 100% формула
    2. Короткая строка (<100 символов) с = ≥ ≤ где = идёт после переменной
       без разделителя (X=, I=, P=, S=, U=, R= и т.п.)
    3. Строка вида "=0,95∙80=76 кВт" — продолжение формулы
    4. Строка вида "10 ≥ 6 — выполняется" — проверка условия
    
    НЕ формула:
    - "Номинальное напряжение: Uном = 6 кВ" — текстовое описание с параметром
    - "Коэффициент равен 5" — глагол + текст
    - "при t = 0,525 с" — контекстное использование
    """
    if not text:
        return False
    
    # 1. Стиль "Формула*" — однозначно формула
    if style_name and style_name.lower().startswith('формула'):
        return True
    
    # 2. Строка начинается с "=" — продолжение формулы
    if text.startswith('=') and len(text) < 100:
        return True
    
    # 3. Короткая строка с ≥ ≤ и проверкой "выполняется"/"невыполняется"
    if re.search(r'[≥≤]', text) and len(text) < 100:
        return True
    
    # 4. Паттерн "X= значение" или "X = значение" где X — 1-3 буквы (переменная)
    # И строка короткая и не содержит ":" перед "=" (это описание параметра)
    # И нет глаголов/длинного текста перед "="
    if len(text) < 100 and re.search(r'[=]', text):
        # Если перед "=" есть двоеточие — это описание: "Напряжение: U = 6" — НЕ формула
        before_eq = text.split('=')[0]
        if ':' in before_eq and len(before_eq) > 10:
            return False
        # Если перед "=" есть русское слово длиннее 3 символов — скорее описание
        # Но "Рс=" или "Qмр.=" — это переменная, а не слово
        # Проверяем: есть ли русский текст из >3 букв перед "="?
        words_before = re.findall(r'[а-яёА-ЯЁ]{4,}', before_eq)
        if words_before:
            # Исключение: сокращения типа "кВар", "кВт" — это единицы, не текст
            # ★ ФИКС V4: Добавлены индексы формул — "табл", "ном", "доп", "расч" и т.д.
            # Это не русские слова, а условные обозначения в электротехнике
            unit_words = {'квар', 'квт', 'мвт', 'мвар',
                          'табл', 'ном', 'доп', 'макс', 'мин', 'расч',
                          'общ', 'уст', 'раб', 'номи', 'откл'}
            non_unit = [w for w in words_before if w.lower() not in unit_words]
            if non_unit:
                return False
        # Если после "=" есть числа или переменные — похоже на формулу
        after_eq = text.split('=', 1)[1] if '=' in text else ''
        # ★ ФИКС V4: Расширено — не только цифры, но и переменные (tg, sin, cos, буквы)
        if after_eq.strip() and re.search(r'[\d∙×*/]', after_eq):
            return True
        # ★ ФИКС V4: "r = h · tg(α)" — после "=" нет цифр, но есть переменные/функции
        if after_eq.strip() and re.search(r'[a-zA-Zα-ωα-ω]', after_eq):
            return True
    
    return False


def _classify(p, is_main_zone, is_biblio_zone=False, doc=None):
    """Классифицировать параграф по его роли в документе."""
    text = p.text.strip()
    style_name = p.style.name if p.style else 'Normal'

    # Пустой
    if not text and not _has_math(p) and not _has_image(p):
        return ParagraphType.EMPTY

    # ★ Подписи к таблицам/рисункам — однозначны по тексту, проверяем ДО зоны титульника
    # Иначе "Таблица 2.4" в стиле Title в середине документа будет пропущена
    if _RE_TABLE_CAP.match(text):
        return ParagraphType.TABLE_CAP
    if _RE_FIGURE.match(text):
        return ParagraphType.FIGURE_CAP

    # ★ ФИКС 7: Структурные заголовки (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ) даже в стиле Title
    # В оригинальном БР.docx ВВЕДЕНИЕ имеет стиль Title, но должно быть STRUCTURAL_H1
    upper = text.upper().rstrip('.')  # ★ БАГ 14: "ПРИЛОЖЕНИЯ." → "ПРИЛОЖЕНИЯ"
    if upper in STRUCTURAL_KEYWORDS and is_main_zone:
        return ParagraphType.STRUCTURAL_H1

    # Зона титульника — иммунитет
    if not is_main_zone:
        return ParagraphType.TITLE_ZONE

    # Элемент оглавления
    if _is_toc_style(style_name):
        return ParagraphType.TOC_ENTRY

    # Заголовок по СТИЛЮ (не по тексту!)
    if _is_heading_style(style_name):
        level = _heading_level(style_name)
        # Структурные элементы — особый случай H1
        if level == 1 and upper in STRUCTURAL_KEYWORDS:
            return ParagraphType.STRUCTURAL_H1
        # ★ Если номер-префикс не совпадает со стилем — пересчитать уровень
        # (напр. исходник имел Heading 3 у «8.1 Исходные…», должно быть H2).
        is_sh2, sh_level2 = _is_section_heading(text)
        if is_sh2 and doc is not None and sh_level2 != level:
            try:
                p.style = doc.styles[f'Heading {sh_level2}']
            except Exception:
                pass
        return ParagraphType.HEADING

    # ★ Заголовок по ТЕКСТУ (стиль Normal, но выглядит как заголовок раздела)
    # Это ПЕРЕД проверкой на список — иначе "1. ХАРАКТЕРИСТИКА..." станет списком
    is_sh, sh_level = _is_section_heading(text)
    if is_sh and is_main_zone:
        # Присваиваем стиль Heading — чтобы Word видел это как заголовок
        if doc is not None:
            try:
                target_style = f'Heading {sh_level}'
                p.style = doc.styles[target_style]
            except Exception:
                pass
        if sh_level == 1 and upper in STRUCTURAL_KEYWORDS:
            return ParagraphType.STRUCTURAL_H1
        return ParagraphType.HEADING

    # Формула (OMML или номер формулы в конце)
    if _has_math(p) or _RE_FORMULA_NUM.search(text):
        return ParagraphType.FORMULA

    # ★ ФИКС 23: Ручная текстовая формула (не OMML, но по сути формула)
    # Признаки: стиль "Формула*", или короткая строка с = ≥ ≤ > < без глаголов
    # Строка вида "Рс= Pр(л)+Pст.у=18,2+4=22,2 кВт" — это формула, не текст
    # Но "Номинальное напряжение сети: Uном = 6 кВ" — это текст с параметром
    if _is_manual_formula(text, style_name):
        return ParagraphType.MANUAL_FORMULA

    # (TABLE_CAP и FIGURE_CAP уже проверены выше, ДО зоны титульника)

    # ★ ФИКС 15: List Paragraph с ЗАГЛАВНЫМ текстом — это ЗАГОЛОВОК, а не список!
    # В БР.docx "ХАРАКТЕРИСТИКА И АНАЛИЗ ОБЪЕКТА ЭЛЕКТРОСНАБЖЕНИЯ." 
    # в стиле List Paragraph с numPr → Word подставляет "1." через нумерацию
    # Форматер видел numPr и делал LIST_ITEM ("— Текст"), а надо HEADING
    num_prs = p._element.xpath('.//w:numPr')
    if is_main_zone and num_prs and len(text) > 10:
        upper_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
        if upper_ratio > 0.4:
            # Убрать numPr — нумерацию заменит текст заголовка
            pPr = p._element.get_or_add_pPr()
            for np in num_prs:
                pPr.remove(np)
            # Присвоить стиль Heading 1
            if doc is not None:
                try:
                    p.style = doc.styles['Heading 1']
                except Exception:
                    pass
            # Добавить номер раздела в начало текста (numPr давал номер, теперь вручную)
            _lp_heading_counter[0] += 1
            sn = _lp_heading_counter[0]
            for run in p.runs:
                if run.text and run.text.strip():
                    run.text = f"{sn}. {run.text.lstrip()}"
                    break
            return ParagraphType.HEADING

    # Элемент списка
    # ★ В зоне библиографии — НЕ считать нумерованные элементы списком
    num_prs = p._element.xpath('.//w:numPr')
    if is_biblio_zone and num_prs:
        return ParagraphType.BODY  # библиография — нумерованный список, НЕ маркер
    if num_prs or _RE_LIST_PREFIX.match(text) or _RE_NUM_LIST.match(text):
        return ParagraphType.LIST_ITEM

    # Расшифровка формулы («где»)
    if _RE_WHERE_LINE.match(text):
        return ParagraphType.WHERE_LINE

    # Обычный текст
    return ParagraphType.BODY


# ──────────────────────────────────────────────────────────────────
#  Обработчики по типам
# ──────────────────────────────────────────────────────────────────

def _handle_title_zone(p):
    """Титульная зона: ПОЛНЫЙ ИММУНИТЕТ — не трогаем вообще ничего.
    ★ ФИКС G: Титульник + задание выдаются кафедрой — нельзя менять.
    Раньше ставили TNR и меняли размеры — это ломало форматирование.
    Теперь: НИКАКИХ изменений. Только шрифт TNR на runs (ГОСТ требует)."""
    # ★ ФИКС G: Вообще не трогаем отступы, интервалы, размеры
    # Только меняем шрифт на TNR — это минимальное ГОСТ-требование
    for run in p.runs:
        if run.font.name and run.font.name.lower() in ('consolas', 'courier new'):
            continue
        run.font.name = cfg.FONT_NAME
        # НЕ меняем размер, bold, italic — всё от стиля


def _handle_structural_h1(p):
    """Структурный заголовок (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ...): центр, без отступа, нежирный."""
    # ★ ФИКС 25: page_break_before для структурных заголовков (ЗАКЛЮЧЕНИЕ и т.п.)
    # Раньше разрыв страницы ставился только для Heading 1, но не для STRUCTURAL_H1
    _insert_page_break_before(p)
    
    # ★ ФИКС 30: Уменьшены отступы структурных заголовков (было sb=24/sa=12)
    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_ind=Cm(0),
                           space_before=Pt(12), space_after=Pt(6),
                           keep_next=True)
    # Убедимся что стиль — Heading 1
    try:
        p.style = p.part.document.styles['Heading 1']
    except Exception:
        pass
    # #v4_REMOVED: Удаление точки в конце заголовка — PRIMUM NON NOCERE
    # ГОСТ запрещает точку, но форматер не должен менять контент
    # Точку убирает автор документа, не форматер
    
    # #v4_REMOVED: Очистка табуляций из runs — PRIMUM NON NOCERE
    # Табуляции могут быть частью содержимого, не трогаем
    for run in p.runs:
        _set_run_font(run, font_size=cfg.FONT_SIZE_H1, bold=cfg.BOLD_H1)


def _handle_heading(p):
    """Обычный заголовок по стилю. ГОСТ: центр, без отступа, без точки."""
    level = _heading_level(p.style.name)
    is_h1 = (level == 1)

    # #v4_REMOVED: Удаление точки в конце заголовка — PRIMUM NON NOCERE
    # ГОСТ запрещает точку, но форматер не должен менять контент
    
    # #v4_REMOVED: Очистка табуляций из runs — PRIMUM NON NOCERE
    # Табуляции могут быть частью содержимого, не трогаем

    # Разрыв страницы перед H1 (кроме первого заголовка в документе)
    if is_h1:
        _insert_page_break_before(p)

    # Формат абзаца
    # ★ ФИКС 30: Уменьшены space_before/after — раньше было sb=24/sa=12 для H1
    # и sb=18/sa=12 для H2 — визуально огромные дыры между разделами
    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_ind=Cm(0),
                           space_before=Pt(12) if is_h1 else Pt(8),
                           space_after=Pt(6),
                           keep_next=True)

    # ★ ФИКС C: НЕ ставим font_size на runs если он None (наследуется от стиля Heading)
    # Стили Heading уже задают правильный размер через word_config/setup_gost_styles
    bold_map = {1: cfg.BOLD_H1, 2: cfg.BOLD_H2, 3: cfg.BOLD_H3}
    size_map = {1: cfg.FONT_SIZE_H1, 2: cfg.FONT_SIZE_H2, 3: cfg.FONT_SIZE_H3}
    bold_val = bold_map.get(level, False)
    # ★ ФИКС C: Передаём размер в _set_run_font, но он НЕ перезапишет если run.font.size is None
    # Таким образом стиль Heading контролирует размер, а не форматер
    size_val = size_map.get(level, cfg.FONT_SIZE_MAIN)

    for run in p.runs:
        _set_run_font(run, font_size=size_val, bold=bold_val)


def _handle_toc_entry(p):
    """Элемент оглавления: только шрифт, не ломаем структуру TOC."""
    for run in p.runs:
        _set_run_font(run, font_size=cfg.FONT_SIZE_MAIN, bold=False)


def _handle_formula(p):
    """Формула: центровка через табуляции."""
    _format_formula(p)
    for run in p.runs:
        if run.font.name and run.font.name.lower() in ('consolas', 'courier new'):
            continue
        run.font.name = cfg.FONT_NAME
        run.font.size = cfg.FONT_SIZE_MAIN


def _handle_manual_formula(p):
    """★ ФИКС 23: Ручная текстовая формула — по центру, без отступа, без табуляций.
    Отличие от OMML-формул: нет номера формулы, нет табуляций.
    Просто CENTER + нулевой first_line_indent."""
    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_ind=Cm(0),
                           space_before=Pt(6), space_after=Pt(6))
    for run in p.runs:
        if run.font.name and run.font.name.lower() in ('consolas', 'courier new'):
            continue
        run.font.name = cfg.FONT_NAME
        run.font.size = cfg.FONT_SIZE_MAIN


def _handle_figure_caption(p, legacy=False):
    """Подпись к рисунку: по центру, без отступа.
    ★ v4.0 PRIMUM NON NOCERE: НЕ перезаписываем runs — только стилизация.
    Ренейминг "Рис." → "Рисунок" — минимальная правка."""
    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_ind=Cm(0),
                           space_before=Pt(6), space_after=Pt(12))

    # ★ v4.0: НЕ трогаем контент — только стилизация
    # Ренейминг "Рис." → "Рисунок" — минимальная правка в runs
    text = p.text.strip()
    match = _RE_FIGURE.match(text)
    if match and not legacy:
        label = match.group(1)
        # Меняем "Рис." на "Рисунок" — минимальная правка в первом run
        if label.lower().startswith('рис') and label.lower() != 'рисунок':
            for run in p.runs:
                if run.text and 'Рис.' in run.text:
                    run.text = run.text.replace('Рис.', 'Рисунок ')
                    break
                elif run.text and 'рис.' in run.text:
                    run.text = run.text.replace('рис.', 'Рисунок ')
                    break

    for run in p.runs:
        _set_run_font(run, font_size=cfg.FONT_SIZE_MAIN, bold=False)


def _handle_table_caption(p, doc=None, legacy=False):
    """Подпись к таблице: над таблицей.
    ★ ГОСТ: Если есть название (— Название) → левый край, без отступа.
    ★ ГОСТ: Если НЕТ названия (только Таблица N) → правый край."""
    text = p.text.strip()
    match = _RE_TABLE_CAP.match(text)
    has_name = match and match.group(3) and match.group(3).strip()
    
    if has_name:
        # "Таблица N — Название" → левый край, без отступа
        _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                               first_line_ind=Cm(0),
                               space_before=Pt(12), space_after=Pt(6),
                               keep_next=True)
    else:
        # "Таблица N" → правый край
        _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.RIGHT,
                               first_line_ind=Cm(0),
                               space_before=Pt(12), space_after=Pt(6),
                               keep_next=True)

    # ★ v4.0: НЕ трогаем контент — только стилизация
    # Ренейминг "Табл." → "Таблица" — минимальная правка runs (без перезаписи всего)
    text = p.text.strip()
    match = _RE_TABLE_CAP.match(text)
    if match and not legacy:
        label = match.group(1)
        # Меняем "Табл." на "Таблица" — минимальная правка в первом run
        if label.lower().startswith('табл'):
            for run in p.runs:
                if run.text and 'Табл.' in run.text:
                    run.text = run.text.replace('Табл.', 'Таблица ')
                    break
                elif run.text and 'табл.' in run.text:
                    run.text = run.text.replace('табл.', 'Таблица ')
                    break

    for run in p.runs:
        _set_run_font(run, font_size=cfg.FONT_SIZE_MAIN, bold=False)


def _find_next_paragraph(p, doc):
    """Найти следующий НЕПУСТОЙ параграф в документе после данного."""
    found = False
    for dp in doc.paragraphs:
        if found:
            if dp.text.strip():
                return dp
            # пустые — пропускаем
            continue
        if dp._element is p._element:
            found = True
    return None


def _has_math_in_run(run):
    """Проверяет, содержит ли run математические элементы."""
    return len(run._r.xpath('.//m:oMath')) > 0


def _handle_list_item(p):
    """Элемент списка: тире + текст, с абзацным отступом.
    ★ v4.0 PRIMUM NON NOCERE: НЕ перезаписываем runs полностью.
    Только стилизация + минимальная правка маркера (замена • на —)."""
    # Удаляем системный маркер Word (numPr), если есть
    num_prs = p._element.xpath('.//w:numPr')
    if num_prs:
        pPr = p._element.get_or_add_pPr()
        pPr.remove(num_prs[0])

    # ★ v4.0: НЕ перезаписываем runs — только минимальная правка маркера
    # Заменяем • / – на — в первом run (если есть)
    text = p.text.strip()
    if text and text[0] in '-–•':
        for run in p.runs:
            if run.text:
                # Заменяем первый маркер на ГОСТ-тире
                for old_marker in ['•', '–', '-']:
                    if run.text.lstrip().startswith(old_marker):
                        run.text = run.text.replace(old_marker, '—', 1)
                        break
                break

    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           first_line_ind=cfg.FIRST_LINE_INDENT,
                           line_spacing=cfg.LINE_SPACING)
    for run in p.runs:
        _set_run_font(run, font_size=cfg.FONT_SIZE_MAIN)


def _handle_where_line(p):
    """Строка «где ...» — расшифровка формулы. Без абзацного отступа.
    ★ v4.0 PRIMUM NON NOCERE: НЕ удаляем пустые параграфы перед 'где' — только стилизация."""
    # #v4_REMOVED: Удаление пустых параграфов перед "где"
    # PRIMUM NON NOCERE: не удаляем параграфы

    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           first_line_ind=Cm(0),
                           line_spacing=cfg.LINE_SPACING)
    for run in p.runs:
        _set_run_font(run, font_size=cfg.FONT_SIZE_MAIN)


def _handle_body(p):
    """Обычный текст: по ширине, с абзацным отступом."""
    _clear_indents_and_set(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           first_line_ind=cfg.FIRST_LINE_INDENT,
                           line_spacing=cfg.LINE_SPACING)
    for run in p.runs:
        _set_run_font(run, font_size=cfg.FONT_SIZE_MAIN)


def _handle_empty(p):
    """Пустой параграф: минимальные интервалы, чтобы не было дыр.
    ★ v4.0 PRIMUM NON NOCERE: НЕ удаляем page break из пустых абзацев —
    автор мог поставить разрыв намеренно. Только стилизация интервалов."""
    # #v4_REMOVED: Удаление page break из пустых параграфов
    # PRIMUM NON NOCERE: page break мог быть поставлен автором намеренно
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    # ★ ФИКС B: Пустой абзац не должен иметь 1.5 интервал — даёт лишние страницы
    # Ставим 1.0 (одинарный) чтобы визуально не раздувать
    pf.line_spacing = 1.0


# ──────────────────────────────────────────────────────────────────
#  Основная логика
# ──────────────────────────────────────────────────────────────────

def process_document(input_path, output_path=None, fast=False, legacy=False):
    print(f"[~] Читаем: {input_path}")
    doc = Document(input_path)

    # 0. Настройка базовых стилей документа
    wu.setup_gost_styles(doc)

    # 1. Параметры страницы
    # ★ ФИКС 26v2: НЕ трогаем ориентацию и размеры секций ВООБЩЕ!
    # Альбомная секция (LANDSCAPE) должна оставаться — форматер убивал её.
    # Размеры страниц оставляем как в оригинале — они правильные.
    # ГОСТ-поля ставим ТОЛЬКО для секций основной зоны (не первая = титульник).
    for i, section in enumerate(doc.sections):
        if i == 0:
            # Первая секция = титульник — НЕ трогаем вообще
            continue
        # Сохраняем ориентацию и размеры как есть
        # Только устанавливаем ГОСТ-поля
        section.top_margin    = cfg.MARGIN_TOP
        section.bottom_margin = cfg.MARGIN_BOTTOM
        section.left_margin   = cfg.MARGIN_LEFT
        section.right_margin  = cfg.MARGIN_RIGHT

    # 2. Определяем начало основной зоны
    is_main_zone = False

    # Первый проход: находим где начинается основная зона
    for p in doc.paragraphs:
        upper = p.text.strip().upper()
        if upper in MAIN_ZONE_TRIGGERS or _is_toc_style(p.style.name if p.style else ''):
            is_main_zone = True
            break

    # Если нет титульника — вся документ основная зона
    if not is_main_zone:
        is_main_zone = True

    # 3. Обработка параграфов
    print("[*] Анализ и нормализация текста...")

    # ★ ФИКС 11: ОТКЛЮЧЕН — удалял формулы (OMML-параграфы)
    # p.text пуст для OMML, но формулы отображаются в Word.
    # Удаление 104 "пустых" параграфов уничтожило 56% контента!
    # Формулы (m:oMath) — НЕ пустые, они рендерятся в Word/LibreOffice.
    removed_omml = 0
    # for p in list(doc.paragraphs):
    #     has_math = _has_math(p)
    #     text = p.text.strip()
    #     if has_math and not text:
    #         if not _has_image(p):
    #             p._element.getparent().remove(p._element)
    #             removed_omml += 1
    # if removed_omml:
    #     print(f"[ФИКС 11] Удалено {removed_omml} пустых OMML-параграфов")

    # ★ v4.1 ФИКС I (мягкий): схлопывание серий пустых абзацев
    # Старый ФИКС I ломал структуру (удалял пустые перед таблицами).
    # Новый подход: схлопываем серии > 2 пустых, оставляем 2.
    # Также удаляем пустые параграфы с page break (создают пустые страницы).
    # ВАЖНО: сканируем только прямые children body (не внутри таблиц!)
    # ВАЖНО: НЕ модифицируем body во время итерации — сначала собираем, потом удаляем!
    MAX_CONSECUTIVE_EMPTY = 2
    removed_empty = 0
    body = doc.element.body
    consecutive_empty_elems = []  # накапливаем пустые
    to_remove = []  # элементы на удаление
    
    for child in list(body):  # list() — копия, безопасная для итерации
        if child.tag != qn('w:p'):
            # Это таблица или другой элемент — сбрасываем серию
            if len(consecutive_empty_elems) > MAX_CONSECUTIVE_EMPTY:
                to_remove.extend(consecutive_empty_elems[MAX_CONSECUTIVE_EMPTY:])
            for wp_e, has_pb in consecutive_empty_elems[:MAX_CONSECUTIVE_EMPTY]:
                if has_pb:
                    to_remove.append((wp_e, True))
            consecutive_empty_elems = []
            continue
        
        wp = child
        # Определяем, пустой ли параграф
        texts = [t.text for t in wp.findall('.//' + qn('w:t')) if t.text]
        full_text = ''.join(texts).strip()
        has_object = wp.findall('.//' + qn('w:object'))
        has_math = wp.findall('.//' + qn('m:oMath'))
        has_image = wp.findall('.//' + qn('w:drawing'))
        has_page_break = False
        for br in wp.findall('.//' + qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                has_page_break = True
                break
        pPr = wp.find(qn('w:pPr'))
        has_sectPr = pPr is not None and pPr.find(qn('w:sectPr')) is not None
        
        is_empty = (not full_text and not has_object and not has_math 
                    and not has_image and not has_sectPr)
        
        if is_empty:
            consecutive_empty_elems.append((wp, has_page_break))
        else:
            # Конец серии пустых — схлопываем если > MAX
            if len(consecutive_empty_elems) > MAX_CONSECUTIVE_EMPTY:
                to_remove.extend(consecutive_empty_elems[MAX_CONSECUTIVE_EMPTY:])
            for wp_e, has_pb in consecutive_empty_elems[:MAX_CONSECUTIVE_EMPTY]:
                if has_pb:
                    to_remove.append((wp_e, True))
            consecutive_empty_elems = []
    
    # Обработка хвостовой серии
    if len(consecutive_empty_elems) > MAX_CONSECUTIVE_EMPTY:
        to_remove.extend(consecutive_empty_elems[MAX_CONSECUTIVE_EMPTY:])
    for wp_e, has_pb in consecutive_empty_elems:
        if has_pb:
            to_remove.append((wp_e, True))
    
    # Удаление собранных элементов
    for wp_e, _ in to_remove:
        parent_e = wp_e.getparent()
        if parent_e is not None:
            parent_e.remove(wp_e)
            removed_empty += 1
    
    if removed_empty:
        print(f"[ФИКС I] Схлопнуто {removed_empty} лишних пустых абзацев")

    # #v4_REMOVED: ФИКС 14 — заменён на ФИКС I выше (обрабатывает и хвост, и середину)

    # #v4_REMOVED: ФИКС 16 — объединение разделённых заголовков
    # PRIMUM NON NOCERE: не объединяем заголовки и не удаляем параграфы
    # merged_headings = 0
    # ... (код удалён — см. git history v3.8)

    # #v4_REMOVED: ФИКС 32 — удаление пустых Heading-абзацев
    # PRIMUM NON NOCERE: не удаляем пустые Heading-абзацы
    # removed_empty_headings = 0
    # ... (код удалён — см. git history v3.8)

    # Нужно пересчитывать is_main_zone по ходу, т.к. титульник может быть
    is_main_zone = False
    seen_main_trigger = False
    is_biblio_zone = False  # ★ Зона библиографии — не срезать numPr
    # #v4_CHANGED: is_post_conclusion — НЕ пропускаем параграфы после ЗАКЛЮЧЕНИЯ
    # PRIMUM NON NOCERE: форматер должен форматировать ВСЕ параграфы, включая после заключения
    # Но для справки помечаем зону (если нужно для других целей)
    is_post_conclusion = False
    _lp_heading_counter[0] = 0  # ★ Сброс счётчика ФИКС 15

    for p in doc.paragraphs:
        # ★ v4.0: После ЗАКЛЮЧЕНИЯ — форматируем шрифт, но не меняем контент
        upper = p.text.strip().upper().rstrip('.')
        # #v4_REMOVED: continue после заключения — PRIMUM NON NOCERE
        # Форматер должен обрабатывать ВСЕ параграфы, в т.ч. после заключения

        # ★ ФИКС 13: Стили HTML Preformatted / Normal (Web) → нормальный Body
        style_name = p.style.name if p.style else 'Normal'
        if style_name in ('HTML Preformatted', 'Normal (Web)', 'No Spacing'):
            try:
                p.style = doc.styles['Normal']
            except Exception:
                pass

        # Обновляем флаг основной зоны
        if upper in MAIN_ZONE_TRIGGERS or _is_toc_style(p.style.name if p.style else ''):
            if not seen_main_trigger:
                seen_main_trigger = True
                is_main_zone = True

        # ★ Отслеживаем зону библиографии
        if upper in _BIBLIO_KEYWORDS:
            is_biblio_zone = True

        # ★ v4.0: После ЗАКЛЮЧЕНИЯ — форматируем шрифт, но не меняем контент
        if upper == 'ЗАКЛЮЧЕНИЕ':
            is_post_conclusion = True
            # Форматируем сам заголовок ЗАКЛЮЧЕНИЕ и продолжаем
            # (НЕ continue — параграфы после заключения тоже форматировать)

        ptype = _classify(p, is_main_zone, is_biblio_zone=is_biblio_zone, doc=doc)

        # Маршрутизация
        if ptype == ParagraphType.TITLE_ZONE:
            _handle_title_zone(p)
        elif ptype == ParagraphType.STRUCTURAL_H1:
            _handle_structural_h1(p)
        elif ptype == ParagraphType.HEADING:
            _handle_heading(p)
        elif ptype == ParagraphType.TOC_ENTRY:
            _handle_toc_entry(p)
        elif ptype == ParagraphType.FORMULA:
            _handle_formula(p)
        elif ptype == ParagraphType.MANUAL_FORMULA:
            # ★ ФИКС 23: Ручные текстовые формулы — по центру
            _handle_manual_formula(p)
        elif ptype == ParagraphType.FIGURE_CAP:
            _handle_figure_caption(p, legacy=legacy)
        elif ptype == ParagraphType.TABLE_CAP:
            _handle_table_caption(p, doc=doc, legacy=legacy)
        elif ptype == ParagraphType.LIST_ITEM:
            _handle_list_item(p)
        elif ptype == ParagraphType.WHERE_LINE:
            _handle_where_line(p)
        elif ptype == ParagraphType.BODY:
            if is_biblio_zone:
                # ★ ФИКС 24: В зоне библиографии — ТОЛЬКО шрифт TNR
                # НЕ трогаем отступы, выравнивание, интервалы — они свои
                for run in p.runs:
                    _set_run_font(run, font_size=cfg.FONT_SIZE_MAIN)
            else:
                _handle_body(p)
        elif ptype == ParagraphType.EMPTY:
            _handle_empty(p)

    # 4. Нормализация таблиц
    print("[*] Нормализация таблиц...")
    _normalize_tables(doc)

    # 5. Нумерация страниц (внизу по центру, пропуская первую)
    wu.add_page_numbering(doc, smart_skip=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 6. Вставка TOC (если нет)
    doc_xml = doc._element.xml.upper()
    has_toc = ('TOC' in doc_xml or
               any(p.text.strip().upper() in ('СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ')
                   for p in doc.paragraphs))
    if not has_toc:
        print("[+] Вставляем автооглавление...")
        _insert_toc_before_intro(doc)

    # 7. Сохранение
    if not output_path:
        base = os.path.splitext(os.path.basename(input_path))[0]
        out_dir = os.path.join(os.path.dirname(os.path.abspath(input_path)), 'data')
        if not os.path.exists(out_dir):
            out_dir = os.path.dirname(os.path.abspath(input_path))
        output_path = os.path.join(out_dir, f"{base}_GOST.docx")

    wu.save_document_safe(doc, output_path)

    # 7. ★ ФИКС 28: Восстановить титульник и секции из оригинала
    # Форматер ломает титульник через setup_gost_styles (Heading стили наследуются)
    # и может убить альбомную секцию.
    # Решение: 1) заменить параграфы до ВВЕДЕНИЯ на оригинальные
    #          2) восстановить свойства ВСЕХ секций из оригинала
    print("[*] Восстанавливаем титульник из оригинала...")
    title_backup = os.path.join(os.path.dirname(os.path.abspath(input_path)), 'БР_титульник.docx')
    if os.path.exists(title_backup):
        from extract_title import restore_title_section
        restore_title_section(title_backup, output_path, output_path)
    else:
        print("[!] Файл титульника не найден — пропускаем восстановление")
        print(f"    Ожидали: {title_backup}")

    # ★ ФИКС 31: Восстановить свойства ВСЕХ секций из оригинала
    # Форматер НЕ трогает первую секцию, но мог поломать ориентацию других.
    # Восстанавливаем page_width/height/orientation для КАЖДОЙ секции.
    print("[*] Восстанавливаем свойства секций из оригинала...")
    try:
        orig_doc = Document(input_path)
        gost_doc = Document(output_path)
        for i in range(min(len(orig_doc.sections), len(gost_doc.sections))):
            o_sec = orig_doc.sections[i]
            g_sec = gost_doc.sections[i]
            g_sec.page_width = o_sec.page_width
            g_sec.page_height = o_sec.page_height
            g_sec.orientation = o_sec.orientation
            # Для первой секции — полностью восстанавливаем поля из оригинала
            if i == 0:
                g_sec.left_margin = o_sec.left_margin
                g_sec.right_margin = o_sec.right_margin
                g_sec.top_margin = o_sec.top_margin
                g_sec.bottom_margin = o_sec.bottom_margin
        gost_doc.save(output_path)
        print(f"[OK] Свойства секций восстановлены ({min(len(orig_doc.sections), len(gost_doc.sections))} секций)")
    except Exception as e:
        print(f"[!] Ошибка восстановления секций: {e}")

    # 8. COM-обновление (если не --fast)
    if not fast:
        wu.update_document_via_com(output_path)

    print(f"[OK] ГОСТ-форматирование завершено: {output_path}")


# ──────────────────────────────────────────────────────────────────
#  CLI
# ──────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description="ГОСТ-форматер Word-документов v3.2")
    parser.add_argument('-i', '--input', required=True, help="Входной файл или папка")
    parser.add_argument('-o', '--output', help="Выходной файл")
    parser.add_argument('--fast', action='store_true', help="Без обновления полей через MS Word")
    parser.add_argument('--legacy', action='store_true', help="Устаревший стиль оформления (Рис., Таблица 1.)")

    args = parser.parse_args()

    if os.path.isdir(args.input):
        import glob as g
        files = g.glob(os.path.join(args.input, "*.docx"))
        for f in files:
            if "GOST" not in f and not os.path.basename(f).startswith("~"):
                process_document(f, fast=args.fast, legacy=args.legacy)
    else:
        process_document(args.input, args.output, fast=args.fast, legacy=args.legacy)