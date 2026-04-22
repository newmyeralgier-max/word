"""Чистка документа от цветов, лишних точек, дублей пустых абзацев и
остатков чужих документов (например, титульника преддипломной практики)."""
import re
from copy import deepcopy

from docx.oxml.ns import qn

from . import config as cfg
from . import detect
from .utils import p_text, is_empty, strip_trailing_dot


# --- Цвет текста -------------------------------------------------------------

def strip_color(doc) -> int:
    """Убрать явный цвет у ВСЕХ run-ов и стилей — чтобы не было синих/красных
    остатков.

    Чистим:
    * все `<w:color>` в document.xml (runs, pPr, стили внутри рантайма);
    * все `<w:color>` в styles.xml (важно: там сидят шрифтовые стили
      «Heading 3», которые красят заголовки в #4F81BD);
    * `<w:highlight>` — подсветка;
    * `themeColor` атрибуты у оставшихся элементов.
    """
    removed = 0
    elements = [doc.element]
    # styles.xml
    try:
        styles_part = doc.part.package.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles'
        )
        elements.append(styles_part.element)
    except Exception:
        pass
    # Так же надёжнее — через `doc.styles.element`
    try:
        elements.append(doc.styles.element)
    except Exception:
        pass
    seen = set()
    for root in elements:
        if id(root) in seen:
            continue
        seen.add(id(root))
        for r in list(root.iter(qn('w:color'))):
            parent = r.getparent()
            if parent is not None:
                parent.remove(r)
                removed += 1
        for hl in list(root.iter(qn('w:highlight'))):
            parent = hl.getparent()
            if parent is not None:
                parent.remove(hl)
        for u in root.iter(qn('w:u')):
            if u.get(qn('w:color')) is not None:
                u.set(qn('w:color'), 'auto')
            if u.get(qn('w:themeColor')) is not None:
                del u.attrib[qn('w:themeColor')]
    return removed


def strip_underline_squiggle(doc) -> int:
    """Удалить <w:proofErr> — красные/зелёные волны орфографии."""
    removed = 0
    for err in list(doc.element.iter(qn('w:proofErr'))):
        parent = err.getparent()
        if parent is not None:
            parent.remove(err)
            removed += 1
    return removed


# --- Тире/дефисы/пробелы ----------------------------------------------------

_DASH_AFTER_NUM = re.compile(r'(\d)\s*-\s*(\d)')
# «где X - описание», «слово - слово» → em-dash
_WORD_HYPHEN = re.compile(r'(\s)-(\s)')


def normalize_dashes(doc) -> int:
    """В обычном тексте (кроме формул/таблиц-шапок) заменяем одиночный дефис
    с пробелами на «—», и дефис между числами в диапазоне на «–»."""
    changed = 0
    for p in doc.paragraphs:
        if detect.has_math(p):
            continue
        new_runs_text = [r.text or '' for r in p.runs]
        for i, txt in enumerate(new_runs_text):
            updated = txt
            updated = _WORD_HYPHEN.sub(r'\1—\2', updated)
            if updated != txt:
                p.runs[i].text = updated
                changed += 1
    return changed


# --- Nbsp между числом и единицей измерения ---------------------------------

_RE_NUM_UNIT = re.compile(
    r'(\d)\s+(кВт|кВ·А|кВА|кВар|кВАр|кВар|кВ|Вт|В|А|мА|Гц|Ом|м|мм|см|км|%|°C|°)'
)


def add_nbsp_units(doc) -> int:
    changed = 0
    for p in doc.paragraphs:
        if detect.has_math(p):
            continue
        for r in p.runs:
            if not r.text:
                continue
            new = _RE_NUM_UNIT.sub(lambda m: f'{m.group(1)}\u00a0{m.group(2)}', r.text)
            if new != r.text:
                r.text = new
                changed += 1
    return changed


# --- Удаление серий пустых абзацев ------------------------------------------

def collapse_empty_paragraphs(doc, *, max_consec: int = 1) -> int:
    """Оставить не более `max_consec` подряд пустых абзацев в основной зоне.

    Работает только СРЕДИ параграфов в body: каждую серию длиннее допустимой
    усекаем до допустимой длины.  Пустой = ни текста, ни OMML, ни drawing.
    Не трогает абзацы внутри таблиц.

    ВАЖНО: параграфы, содержащие `w:sectPr` в pPr (границы секций), НЕ удаляем —
    иначе схлопнутся секции с разными ориентациями/полями.
    """
    body = doc.element.body
    streak = 0
    to_remove = []
    for child in list(body):
        if child.tag != qn('w:p'):
            streak = 0
            continue
        pPr = child.find(qn('w:pPr'))
        has_sectPr = pPr is not None and pPr.find(qn('w:sectPr')) is not None
        if has_sectPr:
            # Никогда не удаляем параграф с границей секции
            streak = 0
            continue
        # Ручное определение пустого
        has_text = bool(
            ''.join((t.text or '') for t in child.findall('.//' + qn('w:t'))).strip()
        )
        has_img = bool(child.findall('.//' + qn('w:drawing'))) or \
                  bool(child.findall('.//' + qn('w:pict')))
        has_math = bool(child.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'))
        has_break = any(
            br.get(qn('w:type')) == 'page' for br in child.findall('.//' + qn('w:br'))
        )
        has_pb_before = (
            pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None
        )
        if has_text or has_img or has_math or has_break or has_pb_before:
            streak = 0
            continue
        streak += 1
        if streak > max_consec:
            to_remove.append(child)
    for el in to_remove:
        el.getparent().remove(el)
    return len(to_remove)


# --- Удаление точки в конце заголовков --------------------------------------

def strip_trailing_dot_in_headings(doc) -> int:
    """У заголовков (Heading N/Заголовок N и у «ПРИЛОЖЕНИЯ.») убираем финальную
    точку. Делаем на уровне runs — не переписывая p.text."""
    changed = 0
    for p in doc.paragraphs:
        if not detect.is_style_heading(p) and not detect.is_structural_element(p):
            if detect.section_heading_level(p_text(p)) == 0:
                continue
        # Убираем точку в последнем run
        for r in reversed(p.runs):
            if r.text and r.text.rstrip():
                txt = r.text.rstrip()
                new = strip_trailing_dot(txt)
                # Сохраняем хвостовые пробелы как были (в конце заголовка их не надо)
                if new != txt:
                    r.text = new
                    changed += 1
                break
    return changed


# --- Удаление ведущего куска «Отчёт о преддипломной практике» --------------

_MIN_MARKER = 'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО'
_PRACTICE_NEEDLES = ('ПРЕДДИПЛОМНОЙ ПРАКТИК', 'ОТЧЕТ О ПРЕДДИПЛОМНОЙ',
                     'ОТЧЁТ О ПРЕДДИПЛОМНОЙ')


def _para_upper_text(p_el) -> str:
    return ''.join(
        (t.text or '') for t in p_el.findall('.//' + qn('w:t'))
    ).upper()


def remove_foreign_block(doc) -> int:
    """Удалить из ВКР вставленный титульник «Отчёт о преддипломной практике».

    Алгоритм:
    * находим все параграфы, начинающиеся с «МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО» —
      это начало очередного титульного блока;
    * если между i-й и (i+1)-й таким «рамкой» в тексте встречается слово
      «преддипломной практик(и/е)» / «ОТЧЁТ» — весь блок [i..i+1) удаляем,
      включая параграфы с `sectPr` (иначе останется сломанная секция).
    """
    body = doc.element.body
    paragraphs = [(i, c) for i, c in enumerate(body) if c.tag == qn('w:p')]
    # индексы параграфов, которые начинают очередной министерский блок
    frame_starts = []
    for i, p_el in paragraphs:
        if _MIN_MARKER in _para_upper_text(p_el):
            # не считать это «началом» повторно, если предыдущий параграф
            # тоже относится к той же рамке (сравниваем соседние)
            if frame_starts and i - frame_starts[-1] < 4:
                continue
            frame_starts.append(i)
    if len(frame_starts) < 2:
        return 0
    removed_total = 0
    # Идём парами (a, b): если между a и b встречается «преддипломн…» — убираем [a..b).
    to_remove_ranges = []
    for a, b in zip(frame_starts, frame_starts[1:] + [len(body)]):
        chunk_upper = ''.join(
            _para_upper_text(body[k]) for k in range(a, b) if body[k].tag == qn('w:p')
        )
        if any(n in chunk_upper for n in _PRACTICE_NEEDLES):
            to_remove_ranges.append((a, b))
    # Удаляем с конца, чтобы индексы не поехали
    for a, b in reversed(to_remove_ranges):
        # Перед удалением найдём первый параграф ПОСЛЕ диапазона и поставим
        # ему pageBreakBefore — чтобы соседние блоки не склеивались на одной
        # странице. (Иначе после удаления вставного титульника ВКР-титульник
        # и Задание срастаются.)
        body_list = list(body)
        if b < len(body_list):
            after = body_list[b]
            if after.tag == qn('w:p'):
                _ensure_page_break_before(after)
        for el in body_list[a:b]:
            body.remove(el)
            removed_total += 1
    return removed_total


def _ensure_page_break_before(p_el):
    from docx.oxml import OxmlElement
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_el.insert(0, pPr)
    if pPr.find(qn('w:pageBreakBefore')) is None:
        pb = OxmlElement('w:pageBreakBefore')
        pPr.insert(0, pb)


# --- Склейка подписей таблиц ------------------------------------------------

def merge_table_captions(doc) -> int:
    """Если идёт «Таблица X.Y» (номер без названия) + следующий параграф с
    названием таблицы — склеиваем в одну строку «Таблица X.Y — Название»."""
    changed = 0
    paras = list(doc.paragraphs)
    for idx, p in enumerate(paras):
        t = p_text(p)
        m = detect.match_table_num_only(t)
        if not m:
            continue
        # Ищем ближайший непустой следующий параграф, который не начинается
        # с таблицы и не является страничным разрывом.
        nxt = None
        for q in paras[idx + 1:idx + 4]:
            qt = p_text(q)
            if not qt:
                continue
            # Если это уже начало другой таблицы/структурного — не склеиваем.
            if detect.match_table_caption(qt) or detect.is_structural_element(q):
                break
            nxt = q
            break
        if nxt is None:
            continue
        num = m.group(2)
        cap = p_text(nxt)
        new_text = f'Таблица {num} — {cap}'
        # Переписываем текст первого параграфа
        _replace_paragraph_text(p, new_text)
        # Удаляем следующий параграф
        nxt._element.getparent().remove(nxt._element)
        changed += 1
    return changed


def _replace_paragraph_text(p, new_text: str):
    """Полностью заменить текст параграфа на `new_text`, сохранив pPr и
    используя шрифт первого существующего run (если есть)."""
    first_rPr = None
    for r in p.runs:
        rPr = r._element.find(qn('w:rPr'))
        if rPr is not None:
            first_rPr = deepcopy(rPr)
            break
    # Удаляем все старые runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # Создаём новый run с текстом
    new_run = p.add_run(new_text)
    if first_rPr is not None:
        # Вставляем rPr в начало нового run
        r_elem = new_run._element
        r_elem.insert(0, first_rPr)
