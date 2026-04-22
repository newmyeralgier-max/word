"""Унификация страничных секций под ГОСТ.

Главная задача — привести ВСЕ portrait-секции документа к единым полям
30 / 15 / 20 / 20 мм, чтобы текст и таблицы не «гуляли». Landscape-секции
сохраняют ориентацию, но тоже получают унифицированные поля (с поворотом).
"""
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import config as cfg


def _set_orient(sec, orient_value: str):
    """Жёстко зафиксировать атрибут `w:orient` в pgSz (иначе Word/LibreOffice
    может «переориентировать» страницу, если W>H при orient=portrait)."""
    sectPr = sec._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.set(qn('w:orient'), orient_value)


def kill_mirror_margins(doc) -> bool:
    """Убрать `<w:mirrorMargins/>` из settings.xml.

    Главный виновник «гуляющего» текста: при mirror margins чётные страницы
    получают зеркальные поля (L=15 / R=30 вместо L=30 / R=15). Word применяет
    это автоматически, даже если `sectPr` одинаковы. Единственное место, где
    включается эта галка — `settings.xml`.

    Возвращает True, если что-то удалили.
    """
    try:
        settings = doc.settings.element
    except Exception:
        return False
    removed = False
    for el in list(settings.iter(qn('w:mirrorMargins'))):
        el.getparent().remove(el)
        removed = True
    # Заодно выключим evenAndOddHeaders — иначе тоже могут быть разные нижние
    # колонтитулы на чётных / нечётных, что визуально воспринимается как съезд.
    for el in list(settings.iter(qn('w:evenAndOddHeaders'))):
        el.getparent().remove(el)
        removed = True
    # И gutterAtTop
    for el in list(settings.iter(qn('w:gutterAtTop'))):
        el.getparent().remove(el)
        removed = True
    return removed


def unify_section_geometry(doc) -> int:
    """Выставить единые поля для всех секций.

    Для PORTRAIT: L=30, R=15, T=20, B=20 мм, размер A4 210×297.
    Для LANDSCAPE: меняем местами W/H, поля L=20 R=20 T=30 B=15 мм (чтобы
    область печати совпадала с portrait-аналогом при визуальном сравнении).
    """
    n = 0
    for sec in doc.sections:
        # Определяем ориентацию по явному атрибуту (W/H могут быть мусором)
        is_landscape = sec.orientation == WD_ORIENTATION.LANDSCAPE
        if is_landscape:
            sec.page_width = cfg.PAGE_HEIGHT   # 297
            sec.page_height = cfg.PAGE_WIDTH   # 210
            sec.top_margin = cfg.MARGIN_LEFT
            sec.bottom_margin = cfg.MARGIN_RIGHT
            sec.left_margin = cfg.MARGIN_TOP
            sec.right_margin = cfg.MARGIN_BOTTOM
            _set_orient(sec, 'landscape')
        else:
            sec.page_width = cfg.PAGE_WIDTH
            sec.page_height = cfg.PAGE_HEIGHT
            sec.left_margin = cfg.MARGIN_LEFT
            sec.right_margin = cfg.MARGIN_RIGHT
            sec.top_margin = cfg.MARGIN_TOP
            sec.bottom_margin = cfg.MARGIN_BOTTOM
            _set_orient(sec, 'portrait')
        # На всякий случай сбрасываем gutter/header/footer distance
        sec.gutter = 0
        sec.header_distance = cfg.MARGIN_TOP // 2 or cfg.MARGIN_TOP
        sec.footer_distance = cfg.MARGIN_BOTTOM // 2 or cfg.MARGIN_BOTTOM
        n += 1
    return n


# --- Нумерация страниц ------------------------------------------------------

def add_page_numbers(doc, *, skip_first: bool = True):
    """Поставить в нижний колонтитул номер страницы (по правому краю).

    `skip_first=True` — на первой странице секции (титуле) номер не печатать.
    """
    for idx, sec in enumerate(doc.sections):
        sec.different_first_page_header_footer = skip_first and idx == 0
        footer = sec.footer
        footer.is_linked_to_previous = False
        # Полностью очищаем колонтитул — в т.ч. `<w:sdt>`-обёртки Word
        # (Page Numbers Gallery), которые python-docx не видит в .paragraphs.
        ftr_el = footer._element  # w:ftr
        for child in list(ftr_el):
            ftr_el.remove(child)
        p = footer.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        # Выравнивание по правому краю
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)
        _add_page_number_field(p)


def _add_page_number_field(p):
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE \\* MERGEFORMAT'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_elem = run._element
    for el in (fld_begin, instr, fld_sep, t, fld_end):
        r_elem.append(el)
