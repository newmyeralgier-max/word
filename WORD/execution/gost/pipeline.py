"""Оркестратор ГОСТ-пайплайна v5."""
from collections import Counter

from docx import Document

from . import config as cfg
from . import cleanup, detect, titles, headings, paragraphs, lists as lst, \
              formulas, tables, figures, toc, page, content
from .utils import p_text, is_empty


def _dispatch_paragraph(p, in_title_zone: bool):
    """Привести одиночный параграф к нужному формату. Возвращает метку типа."""
    if in_title_zone:
        return detect.PType.TITLE_ZONE

    text = p_text(p)

    # Пустой параграф
    if is_empty(p):
        paragraphs.format_empty(p)
        return detect.PType.EMPTY

    # OMML-формула
    if detect.is_formula_paragraph(p):
        formulas.format_formula(p)
        return detect.PType.FORMULA

    # Where-строка
    if detect.is_where_line(p):
        formulas.format_where(p)
        return detect.PType.WHERE

    # Рисунок
    if detect.has_image(p):
        figures.format_figure_paragraph(p)
        return detect.PType.FIGURE_IMG

    # Подпись к рисунку
    if detect.match_figure_caption(text):
        figures.format_figure_caption(p)
        return detect.PType.FIGURE_CAP

    # Подпись к таблице (полная)
    if detect.match_table_caption(text):
        tables.format_table_caption(p)
        return detect.PType.TABLE_CAP

    # Структурный элемент ГОСТ
    if detect.is_structural_element(p):
        headings.format_structural(p)
        return detect.PType.STRUCTURAL

    # Заголовок по стилю
    lvl = detect.is_style_heading(p)
    if not lvl:
        # Заголовок по тексту («1. Характеристика...»)
        lvl = detect.section_heading_level(text)
    if lvl:
        headings.format_heading(p, lvl, page_break=(lvl == 1))
        return detect.PType.HEADING

    # Элементы оглавления оставляем TOC-стилю
    if detect.is_toc_paragraph(p):
        return detect.PType.TOC

    # Списки
    if detect.list_marker_kind(text):
        lst.format_list_item(p)
        return detect.PType.LIST

    # Обычный текст
    paragraphs.format_body(p)
    return detect.PType.BODY


def run(doc_path: str, out_path: str, *, verbose: bool = True):
    """Прогнать полный пайплайн над `doc_path` и сохранить в `out_path`."""
    doc = Document(doc_path)
    stats = Counter()

    # 0a. Убить mirror margins в settings.xml — главный виновник «гуляющего»
    #     текста: иначе чётные страницы получают зеркальные поля.
    stats['mirror_killed'] = int(page.kill_mirror_margins(doc))

    # 0b. Unify sections — одинаковые поля в каждом `sectPr`.
    page.unify_section_geometry(doc)

    # 1. Чистка — цвет, waves, удаление foreign block
    stats['color_removed'] = cleanup.strip_color(doc)
    stats['proof_err_removed'] = cleanup.strip_underline_squiggle(doc)
    stats['foreign_block_removed'] = cleanup.remove_foreign_block(doc)

    # 1a. Контент: починить даты в ЗАДАНИИ, снести рукописный TOC,
    #     переписать ВВЕДЕНИЕ. РЕФЕРАТ и TOC вставляются позже,
    #     чтобы порядок был: титул → РЕФЕРАТ → СОДЕРЖАНИЕ → ВВЕДЕНИЕ.
    stats['dates_fixed'] = content.fix_task_dates(doc)
    stats['manual_toc_removed'] = content.remove_manual_toc(doc)
    stats['intro_rewritten'] = int(content.replace_intro(doc))

    # 2. Склейка подписей таблиц «Таблица X.Y» + «Название» → одна строка
    stats['captions_merged'] = cleanup.merge_table_captions(doc)

    # 3. Коллапс серий пустых параграфов. ВАЖНО: слишком агрессивный коллапс
    #    приводит к тому, что LibreOffice при конвертации в PDF теряет
    #    финальные разделы (ЗАКЛЮЧЕНИЕ, СПИСОК ЛИТЕРАТУРЫ, ПРИЛОЖЕНИЯ) — часть
    #    контента «прилипает» к плавающим OLE-объектам из ЗАДАНИЯ и уходит
    #    за границу страницы. Поэтому оставляем до 3 подряд пустых.
    stats['empty_collapsed'] = cleanup.collapse_empty_paragraphs(doc, max_consec=3)

    # 4. Определяем границу title zone
    title_end_idx = titles.find_title_zone_end(doc)
    if verbose:
        print(f'[gost] title zone = paragraphs [0..{title_end_idx})')
    # Защищаем титульник (меняем только шрифт и цвет)
    titles.normalize_title_zone(doc, title_end_idx)

    # 5. Форматируем все параграфы вне title zone
    for idx, p in enumerate(list(doc.paragraphs)):
        in_title = idx < title_end_idx
        kind = _dispatch_paragraph(p, in_title)
        stats[kind] += 1

    # 6. Таблицы (границы, центр, ширина)
    for t in doc.tables:
        tables.format_table(t)
    stats['tables_formatted'] = len(doc.tables)

    # 7. Тире и неразрывные пробелы — после форматирования
    stats['dashes_fixed'] = cleanup.normalize_dashes(doc)
    stats['nbsp_added'] = cleanup.add_nbsp_units(doc)

    # 8. Убрать точки в конце заголовков (на уровне runs)
    stats['dot_stripped'] = cleanup.strip_trailing_dot_in_headings(doc)

    # 9. TOC вставить, если нет. Сначала TOC (СОДЕРЖАНИЕ), потом РЕФЕРАТ
    #    перед ним — итоговый порядок: титул → РЕФЕРАТ → СОДЕРЖАНИЕ → ВВЕДЕНИЕ.
    inserted = toc.insert_toc_before(doc)
    stats['toc_inserted'] = int(inserted)
    stats['referat_inserted'] = int(content.insert_referat(doc))

    # 10. Повторяем unify — после вставок параграфов могут добавиться секции
    page.unify_section_geometry(doc)

    # 10a. Чистка хвоста — удалить пустые абзацы после последнего содержательного.
    stats['tail_blank_removed'] = content.remove_blank_tail(doc)

    # 11. Номера страниц
    page.add_page_numbers(doc, skip_first=True)

    # 12. Сохраняем
    doc.save(out_path)
    if verbose:
        for k, v in sorted(stats.items()):
            print(f'[gost] {k:<25} {v}')
    return stats
