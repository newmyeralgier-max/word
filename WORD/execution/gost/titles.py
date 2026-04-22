"""Определение и защита «title zone» — всего, что идёт ДО первого ВВЕДЕНИЕ/
РЕФЕРАТА. Внутри зоны мы меняем только шрифт на TNR и цвет на чёрный —
не ломая авторскую вёрстку кафедры.
"""
from docx.oxml.ns import qn

from . import config as cfg
from . import detect
from .utils import p_text, set_run_font


def find_title_zone_end(doc):
    """Вернуть индекс первого параграфа основной зоны. Если не найдено —
    возвращаем len(doc.paragraphs) (тогда всё станет title zone)."""
    for i, p in enumerate(doc.paragraphs):
        t = p_text(p).upper().strip().strip('.').strip()
        if not t:
            continue
        if t in cfg.MAIN_ZONE_TRIGGERS:
            return i
    return len(doc.paragraphs)


def normalize_title_zone(doc, end_idx: int) -> int:
    """Только шрифт и цвет — отступы/размеры/выравнивание не трогаем."""
    changed = 0
    for i, p in enumerate(doc.paragraphs):
        if i >= end_idx:
            break
        for r in p.runs:
            set_run_font(r, font_name=cfg.FONT_NAME, color=cfg.COLOR_BLACK)
            changed += 1
    return changed
