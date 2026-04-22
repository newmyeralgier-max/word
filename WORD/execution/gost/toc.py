"""Вставка автооглавления СОДЕРЖАНИЕ перед первым структурным элементом."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import config as cfg
from . import detect
from .utils import apply_p_format, set_run_font, p_text


def has_toc(doc) -> bool:
    """Определить, есть ли уже TOC-поле."""
    for instr in doc.element.iter(qn('w:instrText')):
        if instr.text and 'TOC' in instr.text.upper():
            return True
    for p in doc.paragraphs:
        name = (p.style.name or '').lower()
        if name.startswith('toc'):
            return True
    return False


def _find_first_main_paragraph(doc):
    """Первый параграф основной зоны (ВВЕДЕНИЕ/РЕФЕРАТ)."""
    for p in doc.paragraphs:
        t = p_text(p).upper().strip().strip('.').strip()
        if t in cfg.MAIN_ZONE_TRIGGERS and t != 'СОДЕРЖАНИЕ' and t != 'ОГЛАВЛЕНИЕ':
            return p
    return None


def insert_toc_before(doc):
    if has_toc(doc):
        return False
    target = _find_first_main_paragraph(doc)
    if target is None:
        return False

    # Заголовок «СОДЕРЖАНИЕ»
    ph = target.insert_paragraph_before()
    apply_p_format(
        ph,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_ind=Cm(0),
        space_before=Pt(0),
        space_after=Pt(12),
        line_spacing=cfg.LINE_SPACING,
        page_break_before=True,
    )
    run = ph.add_run('СОДЕРЖАНИЕ')
    set_run_font(run, font_name=cfg.FONT_NAME, font_size=cfg.FONT_SIZE_H1,
                 bold=False, italic=False, color=cfg.COLOR_BLACK, force_size=True)

    # Поле TOC
    pt = target.insert_paragraph_before()
    apply_p_format(pt, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_ind=Cm(0),
                   line_spacing=cfg.LINE_SPACING)
    rt = pt.add_run()
    r_elem = rt._element

    def make_fld(tag, **kwargs):
        el = OxmlElement(tag)
        for k, v in kwargs.items():
            el.set(qn(k), v)
        return el

    fld_begin = make_fld('w:fldChar', **{'w:fldCharType': 'begin', 'w:dirty': 'true'})
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = make_fld('w:fldChar', **{'w:fldCharType': 'separate'})
    # placeholder run
    t_el = OxmlElement('w:t')
    t_el.text = 'Оглавление будет обновлено в Word (F9).'
    fld_end = make_fld('w:fldChar', **{'w:fldCharType': 'end'})
    for el in (fld_begin, instr, fld_sep, t_el, fld_end):
        r_elem.append(el)

    # У следующего параграфа (ВВЕДЕНИЕ) — page-break-before
    target.paragraph_format.page_break_before = True
    return True
