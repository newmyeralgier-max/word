#!/usr/bin/env python3
"""validate_formatter.py — Проверка что форматер не потерял контент.

Запуск:
    python validate_formatter.py original.docx formatted.docx
    python validate_formatter.py original.docx formatted.docx --verbose

Возвращает:
    0 — всё ок (0 потерь)
    1 — есть потери контента

Проверки:
1. Все непустые параграфы из оригинала присутствуют в GOST
2. Кол-во таблиц совпадает
3. Кол-во изображений совпадает
4. Кол-во OMML-формул совпадает
5. Нет параграфов с текстом, ставших пустыми
6. Отчёт: сколько/какие параграфы потеряны, какие появились
"""

import sys
import argparse
from collections import Counter
from docx import Document
from docx.oxml.ns import qn


def count_tables(doc):
    """Подсчёт таблиц в документе."""
    return len(doc.tables)


def count_drawings(doc):
    """Подсчёт изображений (w:drawing) в документе."""
    count = 0
    for p in doc.paragraphs:
        count += len(p._element.findall('.//' + qn('w:drawing')))
    # Также в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    count += len(p._element.findall('.//' + qn('w:drawing')))
    return count


def count_omml(doc):
    """Подсчёт OMML-формул (m:oMath) в документе."""
    count = 0
    body = doc.element.body
    count += len(body.findall('.//' + qn('m:oMath')))
    return count


def get_nonempty_paragraphs(doc):
    """Получить список непустых параграфов (текст)."""
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def validate(original_path, formatted_path, verbose=False):
    """Валидация форматера — сравнение оригинала и результата."""

    issues = []
    warnings = []

    print(f"Загрузка: {original_path}")
    orig_doc = Document(original_path)
    print(f"Загрузка: {formatted_path}")
    fmt_doc = Document(formatted_path)

    # === 1. Параграфы ===
    orig_paras = get_nonempty_paragraphs(orig_doc)
    fmt_paras = get_nonempty_paragraphs(fmt_doc)

    orig_set = Counter(orig_paras)
    fmt_set = Counter(fmt_paras)

    lost_paras = []
    for text, count in orig_set.items():
        fmt_count = fmt_set.get(text, 0)
        diff = count - fmt_count
        if diff > 0:
            lost_paras.append((text[:80], diff))

    new_paras = []
    for text, count in fmt_set.items():
        orig_count = orig_set.get(text, 0)
        diff = count - orig_count
        if diff > 0:
            new_paras.append((text[:80], diff))

    print(f"\n{'='*60}")
    print(f"ПАРАГРАФЫ")
    print(f"{'='*60}")
    print(f"  Оригинал: {len(orig_paras)} непустых ({len(orig_set)} уникальных)")
    print(f"  GOST:     {len(fmt_paras)} непустых ({len(fmt_set)} уникальных)")

    if lost_paras:
        print(f"  ❌ ПОТЕРЯНО: {sum(c for _, c in lost_paras)} параграфов")
        if verbose:
            for text, count in lost_paras:
                print(f"    - [{count}x] {text}")
        issues.append(f"Потеряно {sum(c for _, c in lost_paras)} параграфов")
    else:
        print(f"  ✅ Потерь параграфов нет")

    if new_paras:
        print(f"  ⚠️  НОВЫХ: {sum(c for _, c in new_paras)} параграфов")
        if verbose:
            for text, count in new_paras:
                print(f"    + [{count}x] {text}")
        warnings.append(f"Появилось {sum(c for _, c in new_paras)} новых параграфов")

    # === 2. Таблицы ===
    orig_tables = count_tables(orig_doc)
    fmt_tables = count_tables(fmt_doc)
    print(f"\n{'='*60}")
    print(f"ТАБЛИЦЫ")
    print(f"{'='*60}")
    print(f"  Оригинал: {orig_tables}")
    print(f"  GOST:     {fmt_tables}")

    if orig_tables != fmt_tables:
        diff = fmt_tables - orig_tables
        sign = "+" if diff > 0 else ""
        print(f"  ❌ Разница: {sign}{diff}")
        issues.append(f"Таблицы: {orig_tables} → {fmt_tables} ({sign}{diff})")
    else:
        print(f"  ✅ Совпадает")

    # === 3. Изображения ===
    orig_imgs = count_drawings(orig_doc)
    fmt_imgs = count_drawings(fmt_doc)
    print(f"\n{'='*60}")
    print(f"ИЗОБРАЖЕНИЯ (w:drawing)")
    print(f"{'='*60}")
    print(f"  Оригинал: {orig_imgs}")
    print(f"  GOST:     {fmt_imgs}")

    if orig_imgs != fmt_imgs:
        diff = fmt_imgs - orig_imgs
        sign = "+" if diff > 0 else ""
        print(f"  ❌ Разница: {sign}{diff}")
        issues.append(f"Изображения: {orig_imgs} → {fmt_imgs} ({sign}{diff})")
    else:
        print(f"  ✅ Совпадает")

    # === 4. OMML-формулы ===
    orig_omml = count_omml(orig_doc)
    fmt_omml = count_omml(fmt_doc)
    print(f"\n{'='*60}")
    print(f"OMML-ФОРМУЛЫ (m:oMath)")
    print(f"{'='*60}")
    print(f"  Оригинал: {orig_omml}")
    print(f"  GOST:     {fmt_omml}")

    if orig_omml != fmt_omml:
        diff = fmt_omml - orig_omml
        sign = "+" if diff > 0 else ""
        print(f"  ❌ Разница: {sign}{diff}")
        issues.append(f"OMML-формулы: {orig_omml} → {fmt_omml} ({sign}{diff})")
    else:
        print(f"  ✅ Совпадает")

    # === 5. Проверка пустых параграфов с контентом ===
    # Параграфы которые имели текст в оригинале, но стали пустыми в GOST
    orig_texts = {p.text.strip() for p in orig_doc.paragraphs if p.text.strip()}
    fmt_texts = {p.text.strip() for p in fmt_doc.paragraphs if p.text.strip()}
    emptied = orig_texts - fmt_texts
    if emptied:
        print(f"\n{'='*60}")
        print(f"❌ ПАРАГРАФЫ СТАВШИЕ ПУСТЫМИ: {len(emptied)}")
        print(f"{'='*60}")
        if verbose:
            for t in sorted(emptied)[:20]:
                print(f"  - {t[:80]}")
        issues.append(f"{len(emptied)} параграфов стали пустыми")

    # === ИТОГ ===
    print(f"\n{'='*60}")
    print(f"ИТОГ")
    print(f"{'='*60}")

    if issues:
        print(f"  ❌ ПРОБЛЕМЫ ({len(issues)}):")
        for issue in issues:
            print(f"    - {issue}")
        if warnings:
            print(f"  ⚠️  Предупреждения ({len(warnings)}):")
            for w in warnings:
                print(f"    - {w}")
        return 1
    else:
        print(f"  ✅ ВСЕ ПРОВЕРКИ ПРОЙДЕНЫ — 0 потерь контента")
        if warnings:
            print(f"  ⚠️  Предупреждения ({len(warnings)}):")
            for w in warnings:
                print(f"    - {w}")
        return 0


def main():
    parser = argparse.ArgumentParser(
        description="Проверка что ГОСТ-форматер не потерял контент"
    )
    parser.add_argument("original", help="Путь к оригинальному DOCX")
    parser.add_argument("formatted", help="Путь к отформатированному DOCX")
    parser.add_argument("-v", "--verbose", action="store_true",
                        help="Показать детали потерь")
    args = parser.parse_args()

    rc = validate(args.original, args.formatted, args.verbose)
    sys.exit(rc)


if __name__ == "__main__":
    main()
