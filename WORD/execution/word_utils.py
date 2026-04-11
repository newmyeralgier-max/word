"""
word_utils.py -- Общие утилиты для ГОСТ-форматирования Word-документов.

Используется и build_docx.py, и format_docx.py.
Содержит: стилизацию, нумерацию страниц, MML2OMML, безопасное сохранение, COM-обновление.
"""

import os
import sys
import re
import glob
import shutil
from datetime import datetime
from pathlib import Path

from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# Загружаем конфиг (файл лежит рядом)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import word_config as cfg


# ??????????????????????????????????????????????????????????????????
#  ГОСТ-стили
# ??????????????????????????????????????????????????????????????????

def setup_gost_styles(doc):
    """Настроить встроенные стили Word под ГОСТ 7.32 (Normal, Heading 1-3, TOC 1-3)."""

    # ?? Normal ??????????????????????????????????????????????????
    sn = doc.styles['Normal']
    sn.font.name      = cfg.FONT_NAME
    sn.font.size      = cfg.FONT_SIZE_MAIN
    sn.font.color.rgb  = cfg.COLOR_BLACK
    pf = sn.paragraph_format
    pf.space_before       = Pt(0)
    pf.space_after        = Pt(0)
    pf.line_spacing       = cfg.LINE_SPACING
    pf.first_line_indent  = cfg.FIRST_LINE_INDENT
    pf.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control      = True

    # ?? Heading 1-3 (английское + русское имя) ??????????????????
    bold_map = {1: cfg.BOLD_H1, 2: cfg.BOLD_H2, 3: cfg.BOLD_H3}
    size_map = {1: cfg.FONT_SIZE_H1, 2: cfg.FONT_SIZE_H2, 3: cfg.FONT_SIZE_H3}

    for lvl in range(1, 4):
        for name in (f'Heading {lvl}', f'Заголовок {lvl}'):
            try:
                hs = doc.styles[name]
            except KeyError:
                continue
            hs.font.name       = cfg.FONT_NAME
            hs.font.color.rgb  = cfg.COLOR_BLACK
            hs.font.bold       = bold_map[lvl]
            hs.font.size       = size_map[lvl]
            hs.paragraph_format.space_before      = Pt(24) if lvl == 1 else Pt(18)
            hs.paragraph_format.space_after        = Pt(12)
            hs.paragraph_format.line_spacing       = cfg.LINE_SPACING
            hs.paragraph_format.keep_with_next     = True
            hs.paragraph_format.alignment          = WD_ALIGN_PARAGRAPH.CENTER
            hs.paragraph_format.first_line_indent  = Cm(0)

    # ?? TOC 1-3 ????????????????????????????????????????????????
    for i in range(1, 4):
        try:
            ts = doc.styles[f'TOC {i}']
        except KeyError:
            continue
        ts.font.name  = cfg.FONT_NAME
        ts.font.size  = cfg.FONT_SIZE_MAIN
        ts.font.bold  = False
        ts.paragraph_format.space_before  = Pt(0)
        ts.paragraph_format.space_after   = Pt(0)
        ts.paragraph_format.line_spacing  = cfg.LINE_SPACING


# ??????????????????????????????????????????????????????????????????
#  MML2OMML -- расположение XSL-файла и кэш
# ??????????????????????????????????????????????????????????????????

_xsl_transform = None          # кэш XSLT-объекта

def _local_xsl_path():
    return os.path.join(os.path.dirname(__file__), 'MML2OMML.XSL')


def get_omml_xslt():
    """Вернуть lxml.etree.XSLT или False, если XSLT не найден."""
    global _xsl_transform
    if _xsl_transform is not None:
        return _xsl_transform

    local = _local_xsl_path()

    # 1) Нет локальной копии -- ищем в Office и копируем
    if not os.path.exists(local):
        pf = os.environ.get("PROGRAMFILES", r"C:\Program Files")
        pf86 = os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)")
        patterns = [
            os.path.join(pf, r"Microsoft Office\root\Office*\MML2OMML.XSL"),
            os.path.join(pf86, r"Microsoft Office\root\Office*\MML2OMML.XSL"),
            os.path.join(pf, r"Microsoft Office\Office*\MML2OMML.XSL"),
            os.path.join(pf86, r"Microsoft Office\Office*\MML2OMML.XSL"),
        ]
        source = None
        for pat in patterns:
            matches = glob.glob(pat)
            if matches:
                source = matches[0]
                break
        if source:
            print(f"  [+] Копируем MML2OMML.XSL ? {local}")
            shutil.copy(source, local)
        else:
            print("  [!] MML2OMML.XSL не найден -- формулы будут текстовыми")
            _xsl_transform = False
            return False

    # 2) Парсим
    try:
        _xsl_transform = etree.XSLT(etree.parse(local))
    except Exception as e:
        print(f"  [!] Ошибка загрузки XSLT: {e}")
        _xsl_transform = False

    return _xsl_transform


def latex_to_omml(latex_str):
    """Конвертировать LaTeX ? OMML XML-элемент (или None)."""
    import latex2mathml.converter

    xslt = get_omml_xslt()
    if not xslt:
        return None
    try:
        mathml_str  = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        return xslt(mathml_tree).getroot()
    except Exception as e:
        print(f"  [!] Формула '{latex_str[:40]}...': {e}")
        return None


# ??????????????????????????????????????????????????????????????????
#  Нумерация страниц
# ??????????????????????????????????????????????????????????????????

def add_page_field(run):
    """Вставить поле PAGE в данный run."""
    for tag, attr in [
        ('w:fldChar', 'begin'),
        ('w:instrText', None),
        ('w:fldChar', 'separate'),
        ('w:fldChar', 'end'),
    ]:
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = 'PAGE'
        else:
            el.set(qn('w:fldCharType'), attr)
        run._r.append(el)


def add_seq_field(run, seq_name="Рисунок"):
    """Вставить поле SEQ (автонумерация) в данный run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {seq_name} \\* ARABIC '
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    t = OxmlElement('w:t')
    t.text = '1'
    run._r.append(t)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)



def add_page_numbering(doc, *, smart_skip=True, align=WD_ALIGN_PARAGRAPH.RIGHT):
    """
    Добавить номера страниц в футер.
    smart_skip=True: пропускает секции, где нумерация уже есть.
    """
    for section in doc.sections:
        footer = section.footer

        if smart_skip:
            f_xml = footer._element.xml.upper() if footer else ''
            h_xml = section.header._element.xml.upper() if section.header else ''
            if (re.search(r'INSTRTEXT.*?(PAGE|NUMPAGES)', f_xml)
                    or re.search(r'INSTRTEXT.*?(PAGE|NUMPAGES)', h_xml)):
                continue
            # цифры в footer
            if footer and any(re.search(r'\d', p.text) for p in footer.paragraphs):
                continue

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = align
        p.clear()
        run = p.add_run()
        run.font.name = cfg.FONT_NAME
        run.font.size = cfg.FONT_SIZE_MAIN
        add_page_field(run)


def add_update_fields_setting(doc):
    """Добавляет флаг принудительного обновления полей (TOC) при открытии документа в Word."""
    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

# ??????????????????????????????????????????????????????????????????
#  Безопасное сохранение
# ??????????????????????????????????????????????????????????????????

def save_document_safe(doc, output_path):
    """Сохранить с перехватом PermissionError и добавлением флага обновления."""
    import time
    
    # Автоматически добавляем флаг обновления полей при каждом сохранении
    add_update_fields_setting(doc)

    max_retries = 3
    base_name, ext = os.path.splitext(output_path)
    
    for attempt in range(1, max_retries + 1):
        try:
            doc.save(output_path)
            print(f"\n[OK] Документ сохранён: {os.path.abspath(output_path)}")
            return True
        except PermissionError:
            print(f"\n[!] Файл '{os.path.basename(output_path)}' открыт или занят. Попытка {attempt}/{max_retries}...")
            if attempt < max_retries:
                time.sleep(2)
            else:
                fallback_path = f"{base_name}_copy{ext}"
                try:
                    doc.save(fallback_path)
                    print(f"\n[!] Удалось сохранить аварийную копию: {os.path.abspath(fallback_path)}")
                    return True
                except Exception as e:
                    print(f"\n[!] Ошибка записи аварийной копии: {e}")
                return False


# ??????????????????????????????????????????????????????????????????
#  Продвинутое форматирование таблиц и объектов
# ??????????????????????????????????????????????????????????????????

def set_table_border_gost(table):
    """
    Установить границы таблицы по ГОСТ (тонкие черные линии 0.5pt).
    """
    tbl = table._element
    tblPr = tbl.xpath('w:tblPr')
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        tblPr = tblPr[0]

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 4 * 1/8 pt = 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def add_gost_caption(paragraph, label="Рисунок", text=""):
    """
    Добавить подпись по ГОСТ: 'Рисунок 1 — Название' или 'Таблица 1 — Название'.
    Автоматически выставляет выравнивание и вставляет SEQ поле.
    """
    paragraph.clear()
    
    if label.lower() == "рисунок":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)

    run = paragraph.add_run(f"{label} ")
    run.font.name = cfg.FONT_NAME
    run.font.size = cfg.FONT_SIZE_MAIN
    
    add_seq_field(paragraph.add_run(), seq_name=label)
    
    suffix = f" — {text}" if text else ""
    run_text = paragraph.add_run(suffix)
    run_text.font.name = cfg.FONT_NAME
    run_text.font.size = cfg.FONT_SIZE_MAIN


def set_section_landscape(section):
    """Переключить ориентацию секции в альбомную."""
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height


def remove_first_page_numbering(doc):
    """Скрыть номер страницы на первом листе (титульнике)."""
    for section in doc.sections:
        section.different_first_page_header_footer = True
        # Очищаем футер первой страницы, если там что-то было
        section.first_page_footer.is_linked_to_previous = False
        for p in section.first_page_footer.paragraphs:
            p.text = ""


def fix_list_indents(doc):
    """
    Исправить отступы всех списков на ГОСТовские (1.25см).
    Word по умолчанию делает слишком большие отступы.
    """
    for para in doc.paragraphs:
        if para.style.name.lower().startswith(('list', 'список')):
            para.paragraph_format.left_indent = Cm(1.25)
            para.paragraph_format.first_line_indent = Cm(-0.63) # Стандартный выступ для маркера



# ??????????????????????????????????????????????????????????????????
#  Word COM: обновление TOC + страницы
# ??????????????????????????????????????????????????????????????????

def update_document_via_com(output_path, *, visible=False):
    """Открывает документ через COM, обновляет поля, TOC, убирает жирность из TOC, сохраняет."""
    import win32com.client

    print("\n[...] Запускаем MS Word для обновления оглавления...")
    word = None
    doc_com = None
    try:
        # Используем DispatchEx для запуска независимого процесса Word
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = visible
        word.DisplayAlerts = 0          # wdAlertsNone

        abs_path = os.path.abspath(output_path)
        doc_com = word.Documents.Open(abs_path)

        doc_com.Fields.Update()
        for toc in doc_com.TablesOfContents:
            toc.Update()
            toc.Range.Font.Bold = False     # содержание без жирного

        doc_com.Save()
        print("[OK] Оглавление и нумерация обновлены MS Word!")

    except Exception as e:
        print(f"[!] Ошибка Word COM: {e}")
    finally:
        if doc_com:
            try:
                doc_com.Close(False)
            except Exception:
                pass
        if word:
            try:
                word.DisplayAlerts = -1     # wdAlertsAll
                word.Quit()
            except Exception:
                pass


# ??????????????????????????????????????????????????????????????????
#  Файловые операции и логирование
# ??????????????????????????????????????????????????????????????????

def create_backup(filepath: str) -> str:
    """Создаёт резервную копию файла с временной меткой."""
    path = Path(filepath)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"backup_{timestamp}_{path.name}"
    backup_path = path.parent / backup_name
    shutil.copy2(filepath, backup_path)
    return str(backup_path)


def log_operation(operation: str, details: str, log_dir: str = ".tmp"):
    """Логирует операцию в файл."""
    os.makedirs(log_dir, exist_ok=True)
    log_path = os.path.join(log_dir, "operations.log")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {operation}: {details}\n")

