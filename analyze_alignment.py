from docx import Document
from docx.oxml.ns import qn

doc = Document('/mnt/d/1. Project/Word/data/БР.docx')

print('=== PARAGRAPHS WITH ALIGNMENT ISSUES ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style_name = p.style.name if p.style else 'Normal'
    align = p.alignment
    pf = p.paragraph_format
    li = pf.left_indent
    ri = pf.right_indent
    fi = pf.first_line_indent
    
    # Проверяем аномалии: LEFT align в основной зоне, странные отступы
    if align and str(align) != 'JUSTIFY (3)' and str(align) != 'CENTER (1)':
        # Не заголовок, не подпись
        if not text.upper() in {'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СОДЕРЖАНИЕ'}:
            if not text.startswith('Рис') and not text.startswith('Таблица'):
                if len(text) > 30:  # длинный текст — явно body
                    print(f'P{i}: align={align}, style={style_name}, li={li}, ri={ri}, fi={fi}')
                    print(f'  text=[{text[:80]}]')
    
    # Проверяем параграфы с левым отступом (могут съехать)
    if li and str(li) != '0' and str(li) != 'None':
        li_cm = li / 360000 if li else 0
        if li_cm > 0.5:  # больше 0.5см — подозрительно
            print(f'P{i}: LEFT_INDENT={li_cm:.2f}cm, style={style_name}, align={align}')
            print(f'  text=[{text[:80]}]')

print()
print('=== CHECKING PAGES 31-32 (approx paragraphs 450-500) ===')
total = len(doc.paragraphs)
print(f'Total paragraphs: {total}')
# Примерно середина документа
start = max(0, total // 2 - 30)
for i in range(start, min(start + 60, total)):
    p = doc.paragraphs[i]
    text = p.text.strip()
    if not text:
        continue
    style_name = p.style.name if p.style else 'Normal'
    align = p.alignment
    pf = p.paragraph_format
    li = pf.left_indent
    print(f'P{i}: style={style_name}, align={align}, li={li}, text=[{text[:60]}]')
