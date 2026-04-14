from docx import Document

doc = Document('/mnt/d/1. Project/Word/data/БР_GOST.docx')

print('=== SECTIONS GOST ===')
for i, s in enumerate(doc.sections):
    w_mm = s.page_width / 36000 if s.page_width else 0
    h_mm = s.page_height / 36000 if s.page_height else 0
    lm_mm = s.left_margin / 36000 if s.left_margin else 0
    rm_mm = s.right_margin / 36000 if s.right_margin else 0
    tm_mm = s.top_margin / 36000 if s.top_margin else 0
    bm_mm = s.bottom_margin / 36000 if s.bottom_margin else 0
    print(f'S{i+1}: {w_mm:.1f}x{h_mm:.1f}mm, orient={s.orientation}')
    print(f'  L={lm_mm:.1f} R={rm_mm:.1f} T={tm_mm:.1f} B={bm_mm:.1f}mm')

print()
print('=== TABLES GOST ===')
for i, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.columns) if t.columns else 0
    first_cell = t.rows[0].cells[0].text[:50] if rows > 0 else ''
    print(f'T{i+1}: {rows}x{cols}, first=[{first_cell}]')

print()
print('=== TABLE ALIGNMENT (jc) ===')
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for i, t in enumerate(doc.tables):
    tbl = t._element
    tblPr = tbl.find(f'{{{ns}}}tblPr')
    jc = None
    if tblPr is not None:
        jc_el = tblPr.find(f'{{{ns}}}jc')
        if jc_el is not None:
            jc = jc_el.get(f'{{{ns}}}val')
    rows = len(t.rows)
    first_cell = t.rows[0].cells[0].text[:30] if rows > 0 else ''
    print(f'T{i+1}: jc={jc}, rows={rows}, first=[{first_cell}]')
