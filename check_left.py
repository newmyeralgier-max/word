from docx import Document
from docx.oxml.ns import qn

doc = Document('/mnt/d/1. Project/Word/data/БР_GOST.docx')

print('=== PARAGRAPHS WITH LEFT INDENT (potential left-shift) ===')
count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else 'Normal'
    pf = p.paragraph_format
    li = pf.left_indent
    fi = pf.first_line_indent
    
    li_cm = li / 360000 if li else 0
    fi_cm = fi / 360000 if fi else 0
    
    if li_cm > 0.3 or fi_cm > 2.0:
        count += 1
        print(f'P{i}: style={style}, li={li_cm:.2f}cm, fi={fi_cm:.2f}cm, align={p.alignment}')
        print(f'  text=[{text[:70]}]')
        if count > 30:
            break

print()
print('=== Normal (Web) STYLE PARAGRAPHS ===')
nw_count = 0
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ''
    if 'Web' in style or 'Preformatted' in style:
        text = p.text.strip()[:60] if p.text.strip() else '[empty]'
        print(f'P{i}: style={style}, align={p.alignment}, text=[{text}]')
        nw_count += 1
        if nw_count > 10:
            break

if nw_count == 0:
    print('  Нет параграфов в стиле Normal (Web) — все сконвертированы!')

print()
print('=== CHECK PAGE 31-32 AREA (body paragraphs with LEFT align) ===')
body_left = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text or len(text) < 20:
        continue
    align = p.alignment
    style = p.style.name if p.style else ''
    if str(align) == 'LEFT (0)' and not style.startswith('Heading') and not text.startswith('Рис') and not text.startswith('Таблица'):
        body_left += 1
        if body_left <= 10:
            print(f'P{i}: style={style}, align={align}, text=[{text[:60]}]')
if body_left == 0:
    print('  Нет body-параграфов с LEFT align!')
else:
    print(f'  Всего body с LEFT: {body_left}')
