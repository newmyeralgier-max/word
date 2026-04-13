from docx import Document
doc = Document('/mnt/d/1. Project/Word/data/БР.docx')

# Find "Таблица 2.4" etc and check context
targets = ['Таблица 2.4', 'Таблица 2.5', 'Таблица 4.1', 'Таблица 5.2']
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    for tgt in targets:
        if t.startswith(tgt):
            style = p.style.name if p.style else 'None'
            print(f'P{i} style={style} text=[{t[:120]}]')
            for j in range(1, 4):
                if i+j < len(doc.paragraphs):
                    np = doc.paragraphs[i+j]
                    nt = np.text.strip()
                    ns = np.style.name if np.style else 'None'
                    print(f'  +{j} style={ns} text=[{nt[:80]}]')
