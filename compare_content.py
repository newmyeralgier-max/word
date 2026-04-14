"""Сравнить текстовое наполнение оригинала и GOST"""
from docx import Document

doc_orig = Document('/mnt/d/1. Project/Word/data/БР.docx')
doc_gost = Document('/mnt/d/1. Project/Word/data/БР_GOST_v2.docx')

# Собираем весь текст из параграфов
orig_texts = []
for p in doc_orig.paragraphs:
    t = p.text.strip()
    if t:
        orig_texts.append(t)

gost_texts = []
for p in doc_gost.paragraphs:
    t = p.text.strip()
    if t:
        gost_texts.append(t)

print(f'Оригинал: {len(orig_texts)} непустых параграфов')
print(f'GOST:     {len(gost_texts)} непустых параграфов')
print(f'Разница: {len(orig_texts) - len(gost_texts)} параграфов потеряно')

# Ищем параграфы которые есть в оригинале но нет в GOST
print('\n=== ПАРАГРАФЫ ИЗ ОРИГИНАЛА, ОТСУТСТВУЮЩИЕ В GOST ===')
gost_set = set(gost_texts)
missing = []
for i, t in enumerate(orig_texts):
    if t not in gost_set and len(t) > 10:  # короткие могут повторяться
        missing.append((i, t))

print(f'Найдено {len(missing)} уникальных параграфов из оригинала, отсутствующих в GOST:')
for idx, (i, t) in enumerate(missing[:40]):
    print(f'  ORIG_P{i}: [{t[:80]}]')

# И наоборот — что появилось в GOST но нет в оригинале
orig_set = set(orig_texts)
new_in_gost = []
for i, t in enumerate(gost_texts):
    if t not in orig_set and len(t) > 10:
        new_in_gost.append((i, t))

print(f'\nНайдено {len(new_in_gost)} новых параграфов в GOST:')
for idx, (i, t) in enumerate(new_in_gost[:20]):
    print(f'  GOST_P{i}: [{t[:80]}]')

# Проверяем таблицы
print(f'\n=== ТАБЛИЦЫ ===')
print(f'Оригинал: {len(doc_orig.tables)} таблиц')
print(f'GOST:     {len(doc_gost.tables)} таблиц')

# Проверяем изображения
orig_imgs = 0
for p in doc_orig.paragraphs:
    if p._element.xpath('.//w:drawing') or p._element.xpath('.//w:pict'):
        orig_imgs += 1

gost_imgs = 0
for p in doc_gost.paragraphs:
    if p._element.xpath('.//w:drawing') or p._element.xpath('.//w:pict'):
        gost_imgs += 1

print(f'Оригинал: {orig_imgs} параграфов с изображениями')
print(f'GOST:     {gost_imgs} параграфов с изображениями')

# OMML формулы
orig_math = 0
for p in doc_orig.paragraphs:
    if p._element.xpath('.//m:oMath') or p._element.xpath('.//m:oMathPara'):
        orig_math += 1

gost_math = 0
for p in doc_gost.paragraphs:
    if p._element.xpath('.//m:oMath') or p._element.xpath('.//m:oMathPara'):
        gost_math += 1

print(f'Оригинал: {orig_math} параграфов с OMML-формулами')
print(f'GOST:     {gost_math} параграфов с OMML-формулами')
