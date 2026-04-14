"""Детально проверить что пропало: OMML формулы и изображения"""
from docx import Document
from docx.oxml.ns import qn

doc_orig = Document('/mnt/d/1. Project/Word/data/БР.docx')
doc_gost = Document('/mnt/d/1. Project/Word/data/БР_GOST_v2.docx')

# 1. Сравнить OMML формулы — найти потерянную
print('=== OMML ФОРМУЛЫ: сравнение ===')
orig_formulas = []
for i, p in enumerate(doc_orig.paragraphs):
    if p._element.xpath('.//m:oMath') or p._element.xpath('.//m:oMathPara'):
        text = p.text.strip()[:60] if p.text.strip() else '[no text]'
        orig_formulas.append((i, text))

gost_formulas = []
for i, p in enumerate(doc_gost.paragraphs):
    if p._element.xpath('.//m:oMath') or p._element.xpath('.//m:oMathPara'):
        text = p.text.strip()[:60] if p.text.strip() else '[no text]'
        gost_formulas.append((i, text))

print(f'Оригинал: {len(orig_formulas)} OMML-формул')
print(f'GOST:     {len(gost_formulas)} OMML-формул')

# 2. Изображения — какие потеряны
print('\n=== ИЗОБРАЖЕНИЯ ===')
orig_img_descs = []
for i, p in enumerate(doc_orig.paragraphs):
    has_drawing = len(p._element.xpath('.//w:drawing')) > 0
    has_pict = len(p._element.xpath('.//w:pict')) > 0
    if has_drawing or has_pict:
        text = p.text.strip()[:60] if p.text.strip() else '[no caption]'
        style = p.style.name if p.style else ''
        orig_img_descs.append((i, text, style))
        print(f'ORIG P{i}: drawing={has_drawing}, pict={has_pict}, style={style}, text=[{text}]')

gost_img_descs = []
for i, p in enumerate(doc_gost.paragraphs):
    has_drawing = len(p._element.xpath('.//w:drawing')) > 0
    has_pict = len(p._element.xpath('.//w:pict')) > 0
    if has_drawing or has_pict:
        text = p.text.strip()[:60] if p.text.strip() else '[no caption]'
        style = p.style.name if p.style else ''
        gost_img_descs.append((i, text, style))
        print(f'GOST P{i}: drawing={has_drawing}, pict={has_pict}, style={style}, text=[{text}]')

# 3. Проверить удалённые пустые параграфы — может быть между рисунком и капшоном
print('\n=== РИСУНОК 4.2 в оригинале (детально) ===')
for i, p in enumerate(doc_orig.paragraphs):
    text = p.text.strip()
    if '4.2' in text and ('Рис' in text or 'рис' in text):
        # Показываем 5 параграфов вокруг
        for j in range(max(0, i-2), min(len(doc_orig.paragraphs), i+5)):
            pp = doc_orig.paragraphs[j]
            has_d = len(pp._element.xpath('.//w:drawing')) > 0
            has_p = len(pp._element.xpath('.//w:pict')) > 0
            tt = pp.text.strip()[:60] if pp.text.strip() else '[empty]'
            st = pp.style.name if pp.style else ''
            marker = ' ← HAS IMAGE' if (has_d or has_p) else ''
            print(f'  ORIG P{j}: style={st}, text=[{tt}]{marker}')

print('\n=== РИСУНОК 4.2 в GOST (детально) ===')
for i, p in enumerate(doc_gost.paragraphs):
    text = p.text.strip()
    if '4.2' in text and ('Рис' in text or 'рис' in text):
        for j in range(max(0, i-2), min(len(doc_gost.paragraphs), i+5)):
            pp = doc_gost.paragraphs[j]
            has_d = len(pp._element.xpath('.//w:drawing')) > 0
            has_p = len(pp._element.xpath('.//w:pict')) > 0
            tt = pp.text.strip()[:60] if pp.text.strip() else '[empty]'
            st = pp.style.name if pp.style else ''
            marker = ' ← HAS IMAGE' if (has_d or has_p) else ''
            print(f'  GOST P{j}: style={st}, text=[{tt}]{marker}')

# 4. Проверить "задание растянулось" — почему страницы так далеко
print('\n=== СТРУКТУРА ТИТУЛЬНИКА В GOST (первые 100 параграфов) ===')
for i in range(min(100, len(doc_gost.paragraphs))):
    p = doc_gost.paragraphs[i]
    text = p.text.strip()[:50] if p.text.strip() else '[empty]'
    style = p.style.name if p.style else ''
    # Только важные параграфы
    if text != '[empty]' or i < 5:
        print(f'P{i}: style={style}, text=[{text}]')
