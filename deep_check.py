from docx import Document
from docx.oxml.ns import qn
import re

doc_orig = Document('/mnt/d/1. Project/Word/data/БР.docx')
doc_gost = Document('/mnt/d/1. Project/Word/data/БР_GOST.docx')

# === 1. ТИТУЛЬНИК: сколько параграфов до СОДЕРЖАНИЯ в оригинале? ===
print('=== ТИТУЛЬНИК: параграфы до СОДЕРЖАНИЯ в оригинале ===')
title_end = 0
for i, p in enumerate(doc_orig.paragraphs):
    upper = p.text.strip().upper()
    print(f'P{i}: style={p.style.name if p.style else ""}, text=[{upper[:60]}]')
    if upper in ('СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ'):
        title_end = i
        break

print(f'\nТитульник = {title_end} параграфов\n')

# === 2. ТАБЛИЦА 2.5 и 4.4 — найти в оригинале ===
print('=== ТАБЛИЦЫ В ОРИГИНАЛЕ ===')
for i, p in enumerate(doc_orig.paragraphs):
    text = p.text.strip()
    if re.match(r'^(Табл\.|Таблица)\s*(2\.5|4\.4)', text, re.I):
        style = p.style.name if p.style else ''
        print(f'P{i}: style={style}, text=[{text[:80]}]')
        # Проверяем что идёт после
        for j in range(1, 6):
            if i+j < len(doc_orig.paragraphs):
                np = doc_orig.paragraphs[i+j]
                nt = np.text.strip()[:60] if np.text.strip() else '[empty]'
                print(f'  +{j}: style={np.style.name if np.style else ""}, text=[{nt}]')

# === 3. ТАБЛИЦА 2.5 и 4.4 — найти в GOST ===
print('\n=== ТАБЛИЦЫ В GOST ===')
for i, p in enumerate(doc_gost.paragraphs):
    text = p.text.strip()
    if re.match(r'^(Табл\.|Таблица)\s*(2\.5|4\.4)', text, re.I):
        style = p.style.name if p.style else ''
        print(f'P{i}: style={style}, text=[{text[:80]}]')
        for j in range(1, 6):
            if i+j < len(doc_gost.paragraphs):
                np = doc_gost.paragraphs[i+j]
                nt = np.text.strip()[:60] if np.text.strip() else '[empty]'
                print(f'  +{j}: style={np.style.name if np.style else ""}, text=[{nt}]')

# === 4. РИСУНОК 4.2 ===
print('\n=== РИСУНОК 4.2 в оригинале ===')
for i, p in enumerate(doc_orig.paragraphs):
    text = p.text.strip()
    if re.match(r'^(Рис\.|Рисунок)\s*4\.2', text, re.I):
        print(f'P{i}: style={p.style.name if p.style else ""}, text=[{text[:80]}]')
        # Есть ли изображение в этом или следующем параграфе?
        has_img = len(p._element.xpath('.//w:drawing')) > 0 or len(p._element.xpath('.//w:pict')) > 0
        print(f'  has_image={has_img}')
        if i+1 < len(doc_orig.paragraphs):
            np = doc_orig.paragraphs[i+1]
            has_img2 = len(np._element.xpath('.//w:drawing')) > 0 or len(np._element.xpath('.//w:pict')) > 0
            print(f'  next has_image={has_img2}, text=[{np.text.strip()[:60]}]')

print('\n=== РИСУНОК 4.2 в GOST ===')
found_42 = False
for i, p in enumerate(doc_gost.paragraphs):
    text = p.text.strip()
    if re.match(r'^(Рис\.|Рисунок)\s*4\.2', text, re.I):
        found_42 = True
        print(f'P{i}: style={p.style.name if p.style else ""}, text=[{text[:80]}]')
if not found_42:
    print('НЕ НАЙДЕН!')

# === 5. ПОСЛЕ ЗАКЛЮЧЕНИЯ ===
print('\n=== ВСЁ ПОСЛЕ ЗАКЛЮЧЕНИЯ в GOST ===')
after_conclusion = False
for i, p in enumerate(doc_gost.paragraphs):
    text = p.text.strip().upper()
    if 'ЗАКЛЮЧЕНИЕ' in text and len(text) < 20:
        after_conclusion = True
    if after_conclusion:
        t = p.text.strip()[:60] if p.text.strip() else '[empty]'
        style = p.style.name if p.style else ''
        print(f'P{i}: style={style}, text=[{t}]')
