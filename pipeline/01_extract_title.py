"""Этап 1: Вырезать титульную зону (до первого структурного элемента).

Сохраняет её в отдельный .docx файл, который в конце пайплайна «вклеится»
обратно. Тело документа (начиная с первого структурного элемента) сохраняется
в отдельный .docx для дальнейших правок.

Использование:
    python 01_extract_title.py <input.docx> <title_out.docx> <body_out.docx>
"""
from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent))
from common import (
    iter_body_elements,
    find_first_structural_idx,
    paragraph_tag,
    table_tag,
)


def extract(input_path: str, title_out: str, body_out: str) -> None:
    doc = Document(input_path)
    elems = iter_body_elements(doc)
    split_idx = find_first_structural_idx(elems)
    if split_idx < 0:
        print("[!] Не найден структурный элемент — падаем.")
        sys.exit(2)
    print(f"[OK] Точка разреза: elem #{split_idx} — «{elems[split_idx].text.strip()[:80]}»")

    body_children = [c for c in doc.element.body if c.tag != qn("w:sectPr")]
    sectPr = doc.element.body.find(qn("w:sectPr"))

    # Границы в XML (индексы в body_children)
    title_cutoff = elems[split_idx].idx_in_body  # всё ДО этого индекса — титульник

    # --- Титульник ---
    title_doc = Document(input_path)  # копия
    tbody = title_doc.element.body
    for child in list(tbody):
        if child.tag == qn("w:sectPr"):
            continue
        tbody.remove(child)
    for i, child in enumerate(body_children):
        if i >= title_cutoff:
            break
        # возвращаем перед sectPr
        if sectPr is not None and tbody.find(qn("w:sectPr")) is not None:
            tbody.insert(list(tbody).index(tbody.find(qn("w:sectPr"))), copy.deepcopy(child))
        else:
            tbody.append(copy.deepcopy(child))
    title_doc.save(title_out)
    print(f"[OK] Титульник сохранён: {title_out}")

    # --- Тело ---
    body_doc = Document(input_path)
    bbody = body_doc.element.body
    for child in list(bbody):
        if child.tag == qn("w:sectPr"):
            continue
        bbody.remove(child)
    for i, child in enumerate(body_children):
        if i < title_cutoff:
            continue
        if sectPr is not None and bbody.find(qn("w:sectPr")) is not None:
            bbody.insert(list(bbody).index(bbody.find(qn("w:sectPr"))), copy.deepcopy(child))
        else:
            bbody.append(copy.deepcopy(child))
    body_doc.save(body_out)
    print(f"[OK] Тело сохранено: {body_out}")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python 01_extract_title.py <input.docx> <title_out.docx> <body_out.docx>")
        sys.exit(1)
    extract(sys.argv[1], sys.argv[2], sys.argv[3])
