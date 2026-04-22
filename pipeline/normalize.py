"""Нормализация контента: опечатки, слитные слова, десятичные разделители,
перенумерация глав.

НЕ удаляет параграфы. НЕ трогает OMML-формулы. НЕ трогает титульную зону
(pipeline отдаёт нам body без титульника).
"""
from __future__ import annotations

import re
from typing import Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn

from common import (
    iter_body_elements,
    paragraph_text,
    paragraph_has_omml,
    paragraph_has_drawing,
    replace_in_paragraph,
    set_paragraph_text,
    BodyElem,
)


# ─── Словарь замен опечаток / слитных слов ───────────────────────────────
# Порядок важен: более длинные варианты заменяются раньше коротких.
TYPOGRAPHY_REPLACEMENTS: List[Tuple[str, str]] = [
    # Кириллические опечатки
    ("Воздущные", "Воздушные"),
    ("Воздущн", "Воздушн"),
    # Слитное «Блок» + английское имя
    ("БлокControlledVoltageSource", "Блок Controlled Voltage Source"),
    ("БлокACVoltageSource", "Блок AC Voltage Source"),
    ("БлокThree-PhaseProgrammableVoltageSource", "Блок Three-Phase Programmable Voltage Source"),
    ("БлокThree-PhaseSource", "Блок Three-Phase Source"),
    ("БлокThree-PhaseBreaker", "Блок Three-Phase Breaker"),
    ("БлокThree-PhasePISectionLine", "Блок Three-Phase PI Section Line"),
    ("БлокPiSectionLine", "Блок Pi Section Line"),
    ("БлокSaturableTransformer", "Блок Saturable Transformer"),
    ("БлокLinearTransformer", "Блок Linear Transformer"),
    ("БлокIdealSwitch", "Блок Ideal Switch"),
    ("БлокBreaker", "Блок Breaker"),
    ("БлокDisplay", "Блок Display"),
    ("БлокScope", "Блок Scope"),
    ("БлокMultimeter", "Блок Multimeter"),
    # Слитные английские идентификаторы
    ("Three-PhaseProgrammableVoltageSource", "Three-Phase Programmable Voltage Source"),
    ("Three-PhaseSource", "Three-Phase Source"),
    ("Three-PhaseBreaker", "Three-Phase Breaker"),
    ("Three-PhasePISectionLine", "Three-Phase PI Section Line"),
    ("Three-PhaseV-IMeasurement", "Three-Phase V-I Measurement"),
    ("PiSectionLine", "Pi Section Line"),
    ("AC Voltage Source", "AC Voltage Source"),
    ("SimscapeElectrical", "Simscape Electrical"),
    ("MATLABv2020b", "MATLAB v2020b"),
    ("MATLABv2019b", "MATLAB v2019b"),
    ("MATLABv", "MATLAB v"),
    ("Simulationtime", "Simulation time"),
    ("Solveroptions", "Solver options"),
    ("InitialStateSetting", "Initial State Setting"),
    ("Tozero", "To zero"),
    ("ACVoltageSource", "AC Voltage Source"),
    ("ControlledVoltageSource", "Controlled Voltage Source"),
    ("LinearTransformer", "Linear Transformer"),
    ("SaturableTransformer", "Saturable Transformer"),
    ("IdealSwitch", "Ideal Switch"),
    ("CurrentMeasurement", "Current Measurement"),
    ("VoltageMeasurement", "Voltage Measurement"),
    # «программыMATLAB» и похожие слитные случаи с кириллицей
    ("программыMATLAB", "программы MATLAB"),
    ("программыMatlab", "программы MATLAB"),
    ("программMATLAB", "программ MATLAB"),
    ("ПРОГРАММЫMATLAB", "ПРОГРАММЫ MATLAB"),
    ("БЛОКОВПРОГРАММЫ", "БЛОКОВ ПРОГРАММЫ"),
    ("ОБЗОР ИЗМЕРИТЕЛЬНЫХ БЛОКОВПРОГРАММЫMATLAB", "ОБЗОР ИЗМЕРИТЕЛЬНЫХ БЛОКОВ ПРОГРАММЫ MATLAB"),
    ("ОБЗОР КОММУТАЦИОННЫХ БЛОКОВ ПРОГРАММЫMATLAB", "ОБЗОР КОММУТАЦИОННЫХ БЛОКОВ ПРОГРАММЫ MATLAB"),
    # кавычки/разделители
    ("  ", " "),
    ("МИНИСТЕРСТВО  НАУКИИВЫСШЕГО  ОБРАЗОВАНИЯ", "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ"),
    ("НАУКИИВЫСШЕГО", "НАУКИ И ВЫСШЕГО"),
    # Даля слипшие
    ("© ФГБОУ ВО“Луганский", "© ФГБОУ ВО «Луганский"),
    ("имени Владимира Даля ”", "имени Владимира Даля»"),
    ("имени Владимира Даля”", "имени Владимира Даля»"),
    ("“Луганский", "«Луганский"),
    # ★ V4.1 доп. фиксы из скриншотов
    ("Three-PhaseTransformer", "Three-Phase Transformer"),
    ("Three-PhaseTranformer", "Three-Phase Transformer"),
    ("ThreeWindings", "Three Windings"),
    ("ТрёхобмоточныйТР", "Трёхобмоточный ТР"),
    ("ТРЕХОБМОТОЧНЫЙТР", "ТРЕХОБМОТОЧНЫЙ ТР"),
    ("WindingParameters", "Winding Parameters"),
    ("фрагментокна", "фрагмент окна"),
    ("окнакнастроек", "окна к настроек"),
    ("настроекблока", "настроек блока"),
    ("настроекоткорректированных", "настроек откорректированных"),
    ("Окнопараметровблока", "Окно параметров блока "),
    ("Окнонастроек", "Окно настроек "),
    ("откорректированныхпараметров", "откорректированных параметров "),
    ("параметровблока", "параметров блока "),
    ("блокаThree-Phase", "блока Three-Phase "),
    ("блокаThreePhase", "блока Three-Phase "),
    ("блокаTransformer", "блока Transformer"),
    ("блокаTransform", "блока Transform"),
    ("кпримеру", "к примеру"),
    ("попримеру", "по примеру"),
    ("спримера", "с примера"),
    ("Например", "Например"),
    ("впримере", "в примере"),
    ("типапримера", "типа примера"),
    ("каквидно", "как видно"),
    ("рассматриваемойсхемы", "рассматриваемой схемы"),
    ("пренебрежимомалыми", "пренебрежимо малыми"),
    # ★ V4.2: слипшие английские имена после ревью
    ("MATLABR2020b", "MATLAB R2020b"),
    ("MATLABR2019b", "MATLAB R2019b"),
    ("MATLABR", "MATLAB R"),
    ("SpecializedPowerSystems", "Specialized Power Systems"),
    ("SimPowerSystems", "SimPowerSystems"),  # keep — official name
    ("MatrixLaboratory", "Matrix Laboratory"),
    ("MathWorks", "MathWorks"),  # keep — company
    ("CommandWindow", "Command Window"),
    ("SimulinkStartPage", "Simulink Start Page"),
    ("NewModel", "New Model"),
    ("OpenModel", "Open Model"),
    ("SaveModel", "Save Model"),
    ("PrintModel", "Print Model"),
    ("SimulationStopTime", "Simulation Stop Time"),
    ("LibraryBrowser", "Library Browser"),
    ("ModelExplorer", "Model Explorer"),
    ("SaveAs", "Save As"),
    ("ElectricalSources", "Electrical Sources"),
    ("SeriesRLCBranch", "Series RLC Branch"),
    ("Three-PhaseSeriesRLCBranch", "Three-Phase Series RLC Branch"),
    ("Three-PhaseSeriesRLC", "Three-Phase Series RLC"),
    ("ACCurrentSource", "AC Current Source"),
    ("DCCurrentSource", "DC Current Source"),
    # склейки: англ.+рус.
    ("программеMATLAB", "программе MATLAB"),
    ("моделейс", "моделей с"),
    ("переменныеNone", "переменные None"),
    # ОкноХХХ + block name в подписях
    ("ВкладкиMain", "Вкладки Main"),
    ("вменюParameters", "в меню Parameters"),
    ("Окнанастроекконфигурации", "Окна настроек конфигурации "),
    ("Окнонастроеквкладки", "Окно настроек вкладки "),
    ("Окнонастроеквкладкиаа", "Окно настроек вкладки "),
    ("ОкнонастроеквкладкиParameter", "Окно настроек вкладки Parameter "),
    ("ОкнонастроеквкладкиConfiguration", "Окно настроек вкладки Configuration "),
    ("Окнонастроек", "Окно настроек "),
    ("измеряемыепеременные", "измеряемые переменные"),
    ("параметровзагрузки", "параметров загрузки"),
    # ★ V4.3: глубокие склейки из Session 5 ревью
    # ex: "FrequencyusedforRLCspecification"
    ("FrequencyusedforRLCspecification", "Frequency used for RLC specification"),
    ("usedforRLCspecification", "used for RLC specification"),
    ("usedfor", "used for "),
    ("forRLCspecification", "for RLC specification"),
    ("RLCspecification", "RLC specification"),
    # слипшаяся кириллица без границ CamelCase — добиваем словарём
    ("Окнанастроекконфигурации", "Окна настроек конфигурации "),
    ("Окнанастроекпараметров", "Окна настроек параметров "),
    ("Фрагментокнанастроек", "Фрагмент окна настроек "),
    ("Фрагментокна", "Фрагмент окна "),
    ("окнанастроекпараметров", "окна настроек параметров "),
    ("окнанастроекконфигурации", "окна настроек конфигурации "),
    ("окнанастроек", "окна настроек "),
    ("настроекпараметровблока", "настроек параметров блока "),
    ("настроекпараметров", "настроек параметров "),
    ("настроекблокаThree", "настроек блока Three"),
    ("параметровблокаThree", "параметров блока Three"),
    ("конфигурацииSaturable", "конфигурации Saturable"),
    ("конфигурацииLinear", "конфигурации Linear"),
    ("параметровблокаSaturable", "параметров блока Saturable"),
    ("параметровблокаLinear", "параметров блока Linear"),
    ("блокаThree-Phase Transformer", "блока Three-Phase Transformer"),
    # «а.ОкнанастроекконфигурацииSaturable Transformer» — варианты
    ("а.Окнанастроек", "а. Окно настроек "),
    ("б.Окнанастроек", "б. Окно настроек "),
    ("в.Окнанастроек", "в. Окно настроек "),
    # Session 6 добивка после pipeline-прогона
    ("фрагментокнанастроек", "фрагмент окна настроек "),
    ("ОкнапараметровблоковSeriesRLCLoad", "Окна параметров блоков Series RLC Load"),
    ("иParallelRLCLoad", " и Parallel RLC Load"),
    ("ParallelRLCLoad", "Parallel RLC Load"),
    ("SeriesRLCLoad", "Series RLC Load"),
    ("БлокThree-PhaseFault", "Блок Three-Phase Fault "),
    ("Three-PhaseFault", "Three-Phase Fault"),
    ("представляетсобой", " представляет собой "),
    ("трёхфазныйкороткозамыкатель", "трёхфазный короткозамыкатель"),
    ("короткозамыкатель", "короткозамыкатель"),
    ("PhaseFaultпредставляетсобойтрёхфазный", "Phase Fault представляет собой трёхфазный"),
    ("а.Окна настроек", "а. Окно настроек"),
    ("б.Окна настроек", "б. Окно настроек"),
    ("в.Окна настроек", "в. Окно настроек"),
    # «ОкнанастроекконфигурацииSaturable Transformer»
    ("Окна настроек конфигурации Saturable", "Окно настроек конфигурации Saturable"),
    ("Окна настроек конфигурации Linear", "Окно настроек конфигурации Linear"),
]


# ★ V4.1: общее правило «русская буква в конце + англ. слово сразу после»
# Для того чтобы не перечислять все возможные случаи вручную.
_RU_EN_STICK = re.compile(r"([а-яёА-ЯЁ])([A-Z][a-z]{2,})")
_EN_RU_STICK = re.compile(r"([a-z]{2,})([А-ЯЁ][а-яё]+)")


def _auto_split_glued(text: str) -> str:
    # кириллица + англ. CamelCase слово → вставляем пробел
    out = _RU_EN_STICK.sub(r"\1 \2", text)
    out = _EN_RU_STICK.sub(r"\1 \2", out)
    return out


def apply_typography(doc_path_in: str, doc_path_out: str) -> dict:
    """Пройти по ВСЕМ параграфам (включая вложенные в таблицы) и применить
    словарные замены + авто-разделитель склеенных слов."""
    doc = Document(doc_path_in)
    counts: Dict[str, int] = {}
    total = 0
    auto_split = 0
    body = doc.element.body
    omml_tag = qn("m:oMath")
    for p_elem in body.iter(qn("w:p")):
        # Типография: безопасна внутри oMath — replace_in_paragraph
        # обходит w:t внутри m:oMath.
        for old, new in TYPOGRAPHY_REPLACEMENTS:
            n = replace_in_paragraph(p_elem, old, new)
            if n:
                counts[old] = counts.get(old, 0) + n
                total += n
        # ★ авто-разделитель слипшихся слов (rus+Eng / eng+Rus)
        for tn in p_elem.findall(".//" + qn("w:t")):
            if not tn.text:
                continue
            # не трогать w:t внутри oMath
            anc = tn.getparent()
            inside_math = False
            while anc is not None and anc is not p_elem:
                if anc.tag == omml_tag:
                    inside_math = True
                    break
                anc = anc.getparent()
            if inside_math:
                continue
            new_text = _auto_split_glued(tn.text)
            if new_text != tn.text:
                auto_split += 1
                tn.text = new_text
                tn.set(qn("xml:space"), "preserve")
    doc.save(doc_path_out)
    return {
        "total_replacements": total,
        "per_pattern": counts,
        "auto_split_runs": auto_split,
    }


# ─── Десятичные разделители ────────────────────────────────────────────────

_DECIMAL_RE = re.compile(r"(?<!\w)(\d+)\.(\d+)(?!\w)")
# Параграф-«заголовок секции» — начинается с N., N.N, N.N.N и т.п.
_SECTION_PREFIX_RE = re.compile(r"^\s*\d+(\.\d+){0,3}\.?\s")
# Подписи рисунков/таблиц и формулы — целиком пропускаем.
_CAPTION_PREFIX_RE = re.compile(r"^\s*(Рис(унок|\.)|Табл(ица|\.)|Формула|Уравнение)\b", re.I)
# Диапазоны/версии: «АПвП(3х70)», «2020b», «1e-12», «v2020b», «x10^3»
# Не трогаем контексты: 2020b, 1e-12, v2020b уже остаются из-за границ \w.
# Но нужно защитить подписи рисунков и формульные вставки (ссылки) вида «Рисунок 6.1» — уже покрыто _CAPTION_PREFIX_RE.
# Внутри обычного текста встречается «по формуле (4.3)» и «на рисунке 6.2» — эти тоже НЕ трогаем.
_INLINE_REF_RE = re.compile(r"\b(рисун|таблиц|рис\.|табл\.|формул|выраж|уравн|пункт|разд|глав)", re.I)


def fix_decimals(doc_path_in: str, doc_path_out: str) -> dict:
    """Заменить точку на запятую в числах вида 32.09 → 32,09.

    Не трогаем:
    - OMML-параграфы
    - Заголовки (toc*, Heading*, структурные элементы)
    - Параграфы, начинающиеся с «N.», «N.N», «N.N.N» — это тоже заголовки
      или списки «1. ...».
    - В оставшихся параграфах — ДОПОЛНИТЕЛЬНО НЕ ТРОГАЕМ первый числовой
      токен, если он в начале строки (защита от строк вида «5.1 Что-то»
      в стиле Normal).
    """
    doc = Document(doc_path_in)
    total = 0
    skipped_headings = 0
    for p in doc.paragraphs:
        if paragraph_has_omml(p._element):
            continue
        style = (p.style.name if p.style else "") or ""
        if style.startswith("toc") or style.startswith("Heading") or style == "List Number":
            skipped_headings += 1
            continue
        t_nodes = p._element.findall(".//" + qn("w:t"))
        if not t_nodes:
            continue
        buf = "".join(n.text or "" for n in t_nodes)
        # Параграф-заголовок по форме текста
        if _SECTION_PREFIX_RE.match(buf):
            skipped_headings += 1
            continue
        # Подписи рисунков / таблиц — пропускаем целиком
        if _CAPTION_PREFIX_RE.match(buf):
            skipped_headings += 1
            continue

        def _repl(m):
            nonlocal total
            start = m.start()
            end = m.end()
            prefix = buf[:start]
            # 1) В самом начале — оставляем как есть (секция).
            if prefix.strip() == "":
                return m.group(0)
            # 2) Референс: «рисунок 3.1», «табл. 4.2», «формуле (4.5)», «главе 2.1».
            window = buf[max(0, start - 40):start].lower()
            if re.search(r"(рисун|рис\.|таблиц|табл\.|формул|выраж|уравн|пункт|п\.\s|разд|глав|главе|пп?\.)\s*$", window):
                return m.group(0)
            # 3) В круглых скобках: «(4.5)» — это номер формулы, не десятичная.
            if start > 0 and buf[start - 1] == "(" and end < len(buf) and buf[end] == ")":
                return m.group(0)
            total += 1
            return f"{m.group(1)},{m.group(2)}"

        new_buf = _DECIMAL_RE.sub(_repl, buf)
        if new_buf != buf:
            t_nodes[0].text = new_buf
            t_nodes[0].set(qn("xml:space"), "preserve")
            for tn in t_nodes[1:]:
                tn.text = ""
    doc.save(doc_path_out)
    return {"fixed": total, "skipped_headings": skipped_headings}


# ─── Перенумерация разделов ───────────────────────────────────────────────
# Мапа «старый ключ заголовка (после удаления номера)» → «новый номер».
# Ключ — нормализованный заголовок (lower, убраны знаки препинания в конце).

# ★ ФИКС V4.1: пустые. Агрессивная перенумерация в прошлый раз сломала
# иерархию (пользователь пожаловался: «идёт 7 раздел, сразу стал 3.4»,
# «9. Моделирование кольцевой — по идее только 5»). До тех пор пока нет
# ручной карты, оставляем оригинальные номера из исходника нетронутыми.
SECTION_MAP: List[Tuple[str, str]] = []
INNER_RENUMBER: List[Tuple[re.Pattern, str]] = []


def _renumber_chapter_title(text: str) -> str:
    """Если текст — заголовок главы (CAPS), найти в SECTION_MAP и поставить новый номер."""
    # Убираем существующий номер, если есть
    m = re.match(r"^\s*(\d+(?:\.\d+)*)\.?\s+(.+?)\s*$", text)
    body = m.group(2) if m else text.strip()
    upper = body.upper()
    for key, new_num in SECTION_MAP:
        if key in upper:
            return f"{new_num}. {body}"
    return text


def _renumber_subtitle(text: str) -> str:
    for pat, repl in INNER_RENUMBER:
        if pat.match(text):
            return pat.sub(repl, text)
    return text


def renumber_sections(doc_path_in: str, doc_path_out: str) -> dict:
    """Перенумеровать главы и подглавы."""
    doc = Document(doc_path_in)
    changes = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or paragraph_has_omml(p._element):
            continue
        new = t
        # H1 CAPS — только если почти весь текст в верхнем регистре и короткий
        letters = [c for c in t if c.isalpha()]
        upper_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
        if upper_ratio > 0.7 and len(t) < 200:
            new2 = _renumber_chapter_title(t)
            if new2 != t:
                new = new2
        # Внутренняя перенумерация
        if new == t:
            new2 = _renumber_subtitle(t)
            if new2 != t:
                new = new2
        if new != t:
            # применяем через set_paragraph_text (теряем мелкое форматирование
            # внутри заголовка — ок, заголовок всё равно перестилизуется)
            set_paragraph_text(p._element, new)
            changes.append((t[:60], new[:60]))
    doc.save(doc_path_out)
    return {"changes": len(changes), "items": changes}


# ─── Срезать точку в конце заголовков ──────────────────────────────────────

_HEAD_RE = re.compile(r"^\s*\d+(\.\d+){0,3}\.?\s+\S")


def replace_bullets_and_unbold(doc_path_in: str, doc_path_out: str) -> dict:
    """★ V4.1:
    1) «• » в начале параграфа → «— » (по ГОСТ).
    2) Снимаем bold с абзацев, у которых ВСЕ runs — bold, текст длиннее 60
       символов, стиль Normal/Body. Это «Этап 2. ...» и прочие случайно
       полностью выделенные жирным абзацы.
    """
    doc = Document(doc_path_in)
    bullets_replaced = 0
    unbold_count = 0

    from docx.oxml.ns import qn as _qn

    for p in doc.paragraphs:
        # 1) Bullet replacement
        t_nodes = p._element.findall(".//" + _qn("w:t"))
        for i, tn in enumerate(t_nodes):
            if not tn.text:
                continue
            # обработать только первый непустой text-узел параграфа
            stripped = tn.text.lstrip()
            if stripped.startswith("•"):
                # заменить «• » / «•» на «— »
                new_text = tn.text.replace("•", "—", 1)
                # убрать лишний пробел сразу после «—»
                new_text = re.sub(r"—\s+", "— ", new_text, count=1)
                if not new_text.lstrip().startswith("— "):
                    new_text = new_text.replace("—", "— ", 1)
                tn.text = new_text
                tn.set(_qn("xml:space"), "preserve")
                bullets_replaced += 1
                break

        # 2) Unbold длинных абзацев body text
        full = p.text
        if len(full.strip()) < 60:
            continue
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading") or style.startswith("Title") or style.startswith("toc"):
            continue
        # Пропускаем подписи рисунков/таблиц
        if re.match(r"^\s*(Рисунок|Таблица)\b", full):
            continue
        runs = list(p.runs)
        if not runs:
            continue
        # все runs c непустым текстом должны быть bold=True
        text_runs = [r for r in runs if r.text and r.text.strip()]
        if not text_runs:
            continue
        if all(r.bold for r in text_runs):
            for r in text_runs:
                r.bold = False
            unbold_count += 1

    doc.save(doc_path_out)
    return {"bullets_replaced": bullets_replaced, "unbolded_paragraphs": unbold_count}


def strip_trailing_periods_in_headings(doc_path_in: str, doc_path_out: str) -> dict:
    doc = Document(doc_path_in)
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or len(t) > 200:
            continue
        if not _HEAD_RE.match(t):
            # структурные элементы (CAPS, короткие)
            letters = [c for c in t if c.isalpha()]
            upper_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
            if upper_ratio < 0.7 or len(t) > 80:
                continue
        if t.endswith("."):
            new = t[:-1]
            set_paragraph_text(p._element, new)
            n += 1
    doc.save(doc_path_out)
    return {"trimmed": n}


if __name__ == "__main__":
    # smoke test
    import sys, os
    if len(sys.argv) != 3:
        print("Usage: python normalize.py <in.docx> <out.docx>")
        sys.exit(1)
    inp, out = sys.argv[1], sys.argv[2]
    tmp1 = out + ".1.tmp"
    tmp2 = out + ".2.tmp"
    tmp3 = out + ".3.tmp"
    print("[1/4] typography ...", apply_typography(inp, tmp1))
    print("[2/4] decimals ...", fix_decimals(tmp1, tmp2))
    print("[3/4] renumber ...", renumber_sections(tmp2, tmp3))
    print("[4/4] strip periods ...", strip_trailing_periods_in_headings(tmp3, out))
    for t in (tmp1, tmp2, tmp3):
        try:
            os.remove(t)
        except OSError:
            pass
    print(f"[OK] saved: {out}")
