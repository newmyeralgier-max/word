import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

def check_styles(file_path):
    doc = docx.Document(file_path)
    
    print("--- STYLES USED ---")
    styles = set()
    for para in doc.paragraphs:
        styles.add(para.style.name)
    for style in sorted(list(styles)):
        print(style)
        
    print("\n--- SAMPLE PARAGRAPH FONT ---")
    if doc.paragraphs:
        para = doc.paragraphs[0]
        print(f"Text: {para.text[:50]}...")
        print(f"Style: {para.style.name}")
        # Check font of first run if available
        if para.runs:
            run = para.runs[0]
            print(f"Font Name: {run.font.name}")
            print(f"Font Size: {run.font.size}")

if __name__ == "__main__":
    check_styles(r"d:\1. Project\Word\data\Уч пособ Матлаб ПРАВИЛЬНО2.docx")
