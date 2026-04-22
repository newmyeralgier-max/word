import docx
import sys
import glob
import os

sys.stdout.reconfigure(encoding='utf-8')

def check_tables():
    files = glob.glob(r"d:\1. Project\Word\data\Уч пособ Матлаб ПРАВИЛЬНО2.docx")
    if not files:
        print("File not found via glob")
        return
    
    file_path = files[0]
    print(f"Opening: {file_path}")
    doc = docx.Document(file_path)
    print(f"Total tables: {len(doc.tables)}")
    
    if doc.tables:
        table = doc.tables[0]
        print(f"First table style: {table.style.name}")
        # Print first row to see content
        for cell in table.rows[0].cells:
            print(f"Cell: {cell.text}")

if __name__ == "__main__":
    check_tables()
