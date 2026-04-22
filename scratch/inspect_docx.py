import docx
import sys

# Force UTF-8 for stdout
sys.stdout.reconfigure(encoding='utf-8')

def extract_docx_info(file_path):
    doc = docx.Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    print("\n--- HEADINGS ---")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            print(f"{para.style.name}: {para.text}")
            
    print("\n--- LAST 50 PARAGRAPHS ---")
    # Get last 50 non-empty paragraphs
    last_paras = [p.text for p in doc.paragraphs if p.text.strip()]
    for p in last_paras[-50:]:
        print(p)
        print("-" * 20)

if __name__ == "__main__":
    extract_docx_info(r"d:\1. Project\Word\data\Уч пособ Матлаб ПРАВИЛЬНО2.docx")
