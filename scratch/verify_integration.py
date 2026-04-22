import docx
import sys
import glob

sys.stdout.reconfigure(encoding='utf-8')

def verify():
    file_path = r"d:\1. Project\Word\data\Уч пособ Матлаб ПРАВИЛЬНО2_MODIFIED.docx"
    doc = docx.Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    print("\n--- NEW HEADINGS ---")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            print(f"{para.style.name}: {para.text}")
            
    print("\n--- LAST 10 PARAGRAPHS ---")
    paras = [p for p in doc.paragraphs if p.text.strip()]
    for p in paras[-10:]:
        print(f"[{p.style.name}] {p.text}")

if __name__ == "__main__":
    verify()
