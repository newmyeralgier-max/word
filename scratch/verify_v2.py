import docx
import sys
import glob

sys.stdout.reconfigure(encoding='utf-8')

def verify():
    file_path = r"d:\1. Project\Word\data\Уч пособ Матлаб ПРАВИЛЬНО2_V2.docx"
    doc = docx.Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Count images (inline shapes)
    img_count = 0
    for p in doc.paragraphs:
        if 'Graphic' in p._element.xml or 'pic:pic' in p._element.xml:
            img_count += 1
            
    print(f"Total figures (approx): {img_count}")

    print("\n--- CHAPTER 6 HEADINGS ---")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and para.text.startswith('6.'):
            print(f"{para.style.name}: {para.text}")
            
    print("\n--- LAST 5 PARAGRAPHS ---")
    paras = [p for p in doc.paragraphs if p.text.strip()]
    for p in paras[-5:]:
        print(f"[{p.style.name}] {p.text}")

if __name__ == "__main__":
    verify()
