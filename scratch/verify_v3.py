import docx
import sys
import glob

sys.stdout.reconfigure(encoding='utf-8')

def verify_v3():
    file_path = r"d:\1. Project\Word\data\Уч пособ Матлаб ПРАВИЛЬНО2_V3.docx"
    doc = docx.Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Check for the error justification keywords
    found_justification = False
    for p in doc.paragraphs:
        if "Обоснование возникших несоответствий" in p.text:
            found_justification = True
            break
    print(f"Found Justification section: {found_justification}")
    
    # Check for methodology keywords
    found_methodology = False
    for p in doc.paragraphs:
        if "Как правильно перенести данные в Simulink?" in p.text:
            found_methodology = True
            break
    print(f"Found Methodology section: {found_methodology}")

if __name__ == "__main__":
    verify_v3()
