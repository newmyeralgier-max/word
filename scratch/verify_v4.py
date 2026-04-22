import docx
import sys
import glob

sys.stdout.reconfigure(encoding='utf-8')

def verify_v4():
    file_path = r"d:\1. Project\Word\data\Уч пособ Матлаб ПРАВИЛЬНО2_V4.docx"
    doc = docx.Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Check last 5 paragraphs to ensure Chapter 6 is there
    paras = [p for p in doc.paragraphs if p.text.strip()]
    for p in paras[-5:]:
        print(f"[{p.style.name}] {p.text}")

if __name__ == "__main__":
    verify_v4()
