from docx import Document
from docx.oxml.ns import qn

doc = Document('/mnt/d/1. Project/Word/data/БР_GOST.docx')

print('=== SECTIONS ===')
for i, s in enumerate(doc.sections):
    w_mm = s.page_width / 36000 if s.page_width else 0
    h_mm = s.page_height / 36000 if s.page_height else 0
    lm_mm = s.left_margin / 36000 if s.left_margin else 0
    rm_mm = s.right_margin / 36000 if s.right_margin else 0
    tm_mm = s.top_margin / 36000 if s.top_margin else 0
    bm_mm = s.bottom_margin / 36000 if s.bottom_margin else 0
    print(f'S{i+1}: {w_mm:.1f}x{h_mm:.1f}mm, orient={s.orientation}')
    print(f'  L={lm_mm:.1f} R={rm_mm:.1f} T={tm_mm:.1f} B={bm_mm:.1f}mm')

print()
print('=== TABLE ALIGNMENT ===')
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for i, t in enumerate(doc.tables):
    tbl = t._element
    tblPr = tbl.find(f'{{{ns}}}tblPr')
    jc = None
    if tblPr is not None:
        jc_el = tblPr.find(f'{{{ns}}}jc')
        if jc_el is not None:
            jc = jc_el.get(f'{{{ns}}}val')
    rows = len(t.rows)
    first_cell = t.rows[0].cells[0].text[:40] if rows > 0 else ''
    print(f'T{i+1}: jc={jc}, {rows}rows, [{first_cell}]')

print()
print('=== FIRST 5 PARAGRAPHS ===')
for i in range(min(5, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    text = p.text.strip()[:60] if p.text.strip() else '[empty]'
    style = p.style.name if p.style else 'None'
    print(f'P{i}: style={style}, text=[{text}]')

print()
print('=== CHECK HEADINGS spacing ===')
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ''
    if style.startswith('Heading') or style.startswith('Заголовок'):
        pf = p.paragraph_format
        sb = pf.space_before
        sa = pf.space_after
        text = p.text.strip()[:50] if p.text.strip() else '[empty]'
        sb_pt = sb / 12700 if sb else 0
        sa_pt = sa / 12700 if sa else 0
        print(f'P{i}: style={style}, sb={sb_pt:.0f}pt, sa={sa_pt:.0f}pt, text=[{text}]')
        if i > 200:
            break

print()
print('=== FORMULA TAB STOPS ===')
count = 0
for i, p in enumerate(doc.paragraphs):
    if p._element.xpath('.//m:oMath'):
        count += 1
        ts = p.paragraph_format.tab_stops
        tab_pos = 'none'
        for tab in ts:
            tab_pos = f'{tab.position/36000:.1f}mm'
        align = p.alignment
        print(f'F{count}: align={align}, tab_stop={tab_pos}')
        if count >= 5:
            break
